
import os, psycopg
from psycopg.rows import dict_row
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

def _conninfo() -> str:
    dsn = os.getenv("DATABASE_URL") or os.getenv("PG_DSN")
    if not dsn:
        raise RuntimeError("Set DATABASE_URL or PG_DSN for Postgres connection")
    return dsn

def _add_title(doc: Document, text: str):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(18)
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT

def _add_kv_table(doc: Document, rows):
    t = doc.add_table(rows=0, cols=2)
    for k, v in rows:
        r = t.add_row().cells
        r[0].text = k; r[1].text = str(v or "")

def generate_qs_report(snapshot_id: str, out_path: str, template_path: str | None = None) -> str:
    doc = Document(template_path) if template_path else Document()
    _add_title(doc, "Qlik Sense Technical Overview")
    doc.add_paragraph(f"Snapshot: {snapshot_id}")

    with psycopg.connect(_conninfo(), row_factory=dict_row) as conn:
        env = conn.execute("SELECT * FROM repmeta_qs.v_environment_overview WHERE snapshot_id = %s", (snapshot_id,)).fetchone()
        _add_title(doc, "Environment")
        if env:
            _add_kv_table(doc, [
                ("Product", env.get("product_name")),
                ("Product Version", env.get("product_version")),
                ("Build Version", env.get("build_version")),
                ("Build Date", env.get("build_date")),
                ("Single Node Only", env.get("single_node_only")),
            ])
        else:
            doc.add_paragraph("No environment data available.")

        lic = conn.execute("SELECT * FROM repmeta_qs.v_license_summary WHERE snapshot_id = %s", (snapshot_id,)).fetchone()
        _add_title(doc, "License Summary")
        if lic:
            _add_kv_table(doc, [
                ("License Number", lic.get("license_number")),
                ("Control Number", lic.get("control_number")),
                ("Expiration", lic.get("expiration")),
                ("Professional Assignees", lic.get("professional_assignees")),
                ("Analyzer Time Assignees", lic.get("analyzer_time_assignees")),
            ])
        else:
            doc.add_paragraph("No license data.")

        _add_title(doc, "Apps by Stream")
        rows = conn.execute("""
            SELECT COALESCE(stream, 'Unassigned') AS stream, count(*) AS apps
            FROM repmeta_qs.v_apps WHERE snapshot_id = %s
            GROUP BY 1 ORDER BY 2 DESC
        """, (snapshot_id,)).fetchall()
        t = doc.add_table(rows=1, cols=2)
        hdr = t.rows[0].cells; hdr[0].text = "Stream"; hdr[1].text = "Apps"
        for r in rows:
            row = t.add_row().cells
            row[0].text = r["stream"]; row[1].text = str(r["apps"])

        _add_title(doc, "Reload Tasks")
        rows = conn.execute("""
            SELECT enabled, count(*) AS ct FROM repmeta_qs.v_reload_tasks
            WHERE snapshot_id = %s GROUP BY enabled ORDER BY enabled DESC
        """, (snapshot_id,)).fetchall()
        t = doc.add_table(rows=1, cols=2)
        hdr = t.rows[0].cells; hdr[0].text = "Enabled"; hdr[1].text = "Count"
        for r in rows:
            row = t.add_row().cells
            row[0].text = "Yes" if r["enabled"] else "No"
            row[1].text = str(r["ct"])

        _add_title(doc, "Users (by Directory)")
        rows = conn.execute("""
            SELECT user_directory, count(*) AS ct FROM repmeta_qs.v_users
            WHERE snapshot_id = %s GROUP BY 1 ORDER BY 2 DESC
        """, (snapshot_id,)).fetchall()
        t = doc.add_table(rows=1, cols=2)
        hdr = t.rows[0].cells; hdr[0].text = "User Directory"; hdr[1].text = "Users"
        for r in rows:
            row = t.add_row().cells
            row[0].text = str(r["user_directory"]); row[1].text = str(r["ct"])

        _add_title(doc, "Governance Checks")
        row = conn.execute("SELECT * FROM repmeta_qs.v_governance_checks WHERE snapshot_id = %s", (snapshot_id,)).fetchone()
        if row:
            _add_kv_table(doc, [
                ("Apps Without Reload Tasks", row.get("apps_without_tasks")),
                ("Disabled Reload Tasks", row.get("disabled_tasks_count")),
            ])

    doc.save(out_path)
    return out_path
