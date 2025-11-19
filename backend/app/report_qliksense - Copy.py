from __future__ import annotations

import os
import tempfile
from datetime import datetime
from typing import Iterable, Optional, Sequence, Tuple, List

import psycopg
from psycopg.rows import dict_row

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, Inches, RGBColor
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

# =========================
# Qlik brand palette (HEX + RGBColor)
# =========================
QLIK_HEX = {
    "green":  "009845",  # Qlik Green
    "green2": "68BD45",
    "blue":   "00A3E0",
    "teal":   "00BFA5",
    "gray9":  "212529",
    "gray6":  "636E72",
    "gray3":  "C8CDD2",
    "gray1":  "F4F6F8",
    "danger": "EF4444",
    "warn":   "F59E0B",
}

def _hex_to_rgbcolor(hexstr: str) -> RGBColor:
    hexstr = hexstr.strip().lstrip("#")
    r = int(hexstr[0:2], 16)
    g = int(hexstr[2:4], 16)
    b = int(hexstr[4:6], 16)
    return RGBColor(r, g, b)

# Pre-baked RGBColor for font colors
QLIK_RGB = {k: _hex_to_rgbcolor(v) for k, v in QLIK_HEX.items()}

FONT_FAMILY = os.getenv("QS_REPORT_FONT") or "Segoe UI"

# =========================
# Low‑level helpers
# =========================

def _conninfo() -> str:
    dsn = os.getenv("DATABASE_URL") or os.getenv("PG_DSN")
    if dsn:
        return dsn
    host = os.getenv("PGHOST", "localhost")
    port = os.getenv("PGPORT", "5432")
    user = os.getenv("PGUSER", "postgres")
    password = os.getenv("PGPASSWORD", "")
    dbname = os.getenv("PGDATABASE", os.getenv("DB_NAME", "postgres"))
    return f"host={host} port={port} user={user} password={password} dbname={dbname}"

def _set_cell_bg(cell, hex_color: str):
    """Set cell background using HEX (no '#')."""
    hex_color = hex_color.strip().lstrip("#").upper()
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)

def _para(doc: Document, text: str = "", size: int = 11, bold: bool = False,
          color: Optional[RGBColor] = None, align=WD_ALIGN_PARAGRAPH.LEFT):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.font.size = Pt(size)
    r.font.bold = bold
    r.font.name = FONT_FAMILY
    if color:
        r.font.color.rgb = color
    p.alignment = align
    return p

def _h1(doc: Document, text: str):
    p = _para(doc, text, size=20, bold=True, color=QLIK_RGB["green"])
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(4)
    return p

def _h2(doc: Document, text: str):
    p = _para(doc, text, size=14, bold=True, color=QLIK_RGB["blue"])
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(2)
    return p

def _footer_with_page_numbers(section):
    footer = section.footer
    p = footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("Page ")
    r.font.name = FONT_FAMILY
    r.font.size = Pt(9)
    fld = OxmlElement("w:fldSimple")
    fld.set(qn("w:instr"), "PAGE \\* MERGEFORMAT")
    p._p.append(fld)
    r2 = p.add_run(" of ")
    r2.font.name = FONT_FAMILY
    r2.font.size = Pt(9)
    fld2 = OxmlElement("w:fldSimple")
    fld2.set(qn("w:instr"), "NUMPAGES \\* MERGEFORMAT")
    p._p.append(fld2)

def _hr(doc: Document, color="A3A3A3"):
    t = doc.add_table(rows=1, cols=1)
    cell = t.rows[0].cells[0]
    tcPr = cell._tc.get_or_add_tcPr()
    borders = OxmlElement("w:tcBorders")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "8")
    bottom.set(qn("w:color"), color)
    borders.append(bottom)
    tcPr.append(borders)

def _kpi_cards(doc: Document, items: Sequence[Tuple[str, str, str]]):
    """
    items: list of (label, value, tone) tone in {'ok','warn','bad','info'}
    """
    tone_map = {
        "ok":   QLIK_RGB["green"],
        "warn": QLIK_RGB["warn"],
        "bad":  QLIK_RGB["danger"],
        "info": QLIK_RGB["gray6"],
    }
    cols = 4 if len(items) >= 4 else max(2, len(items))
    t = doc.add_table(rows=1, cols=cols)
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    for idx, it in enumerate(items):
        label, value, tone = it
        cell = t.rows[0].cells[idx]
        _set_cell_bg(cell, QLIK_HEX["gray1"])
        p_val = cell.paragraphs[0]
        run = p_val.add_run(str(value))
        run.bold = True
        run.font.size = Pt(18)
        run.font.name = FONT_FAMILY
        run.font.color.rgb = tone_map.get(tone, QLIK_RGB["gray9"])
        p_val.alignment = WD_ALIGN_PARAGRAPH.CENTER

        p_lab = cell.add_paragraph()
        r2 = p_lab.add_run(label)
        r2.font.size = Pt(9)
        r2.font.color.rgb = QLIK_RGB["gray6"]
        r2.font.name = FONT_FAMILY
        p_lab.alignment = WD_ALIGN_PARAGRAPH.CENTER

def _table_2col(doc: Document, title_left: str, title_right: str, rows: Iterable[Tuple[str, str]]):
    t = doc.add_table(rows=1, cols=2)
    h = t.rows[0].cells
    h[0].text = title_left
    h[1].text = title_right
    for a, b in rows:
        r = t.add_row().cells
        r[0].text = str(a or "")
        r[1].text = str(b or "")

def _add_toc(doc: Document):
    p = doc.add_paragraph()
    run = p.add_run("Table of Contents")
    run.bold = True
    run.font.size = Pt(14)
    p = doc.add_paragraph()
    fld = OxmlElement("w:fldSimple")
    fld.set(qn("w:instr"), r'TOC \o "1-3" \h \z \u')
    p._p.append(fld)
    _para(doc, "(Right‑click → Update Field in Word to refresh TOC)", size=9, color=QLIK_RGB["gray6"])

def _cover_page(doc: Document, title: str, subtitle_lines: List[str], logo_path: Optional[str] = None):
    # Full‑bleed left rail (Qlik green) using a 1x2 table
    section = doc.sections[0]
    section.left_margin = Inches(0.6)
    section.right_margin = Inches(0.6)
    section.top_margin = Inches(0.6)
    section.bottom_margin = Inches(0.6)

    table = doc.add_table(rows=1, cols=2)
    c0, c1 = table.rows[0].cells
    _set_cell_bg(c0, QLIK_HEX["green"])

    # Left: logo + big accent band
    if logo_path and os.path.exists(logo_path):
        p = c0.paragraphs[0]
        run = p.add_run()
        try:
            run.add_picture(logo_path, height=Inches(0.5))
        except Exception:
            pass
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    c0.add_paragraph("")

    # Right: hero title
    p = c1.paragraphs[0]
    r = p.add_run(title)
    r.font.size = Pt(26)
    r.font.bold = True
    r.font.name = FONT_FAMILY
    r.font.color.rgb = QLIK_RGB["green"]
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT

    for line in subtitle_lines:
        q = c1.add_paragraph()
        rr = q.add_run(line)
        rr.font.size = Pt(11)
        rr.font.name = FONT_FAMILY
        rr.font.color.rgb = QLIK_RGB["gray6"]

    # Add a visual separator
    doc.add_paragraph("")
    _hr(doc, color="CCCCCC")

def _rgb01_from_hex(hexstr: str):
    hexstr = hexstr.strip().lstrip("#")
    r = int(hexstr[0:2], 16) / 255.0
    g = int(hexstr[2:4], 16) / 255.0
    b = int(hexstr[4:6], 16) / 255.0
    return (r, g, b)

def _insert_chart(doc: Document, title: str, series: List[Tuple[str, float]], width_in=6.0):
    """Create a bar chart image with matplotlib if available; fallback to a table."""
    _h2(doc, title)
    try:
        import matplotlib.pyplot as plt
        labels = [a for a, _ in series]
        values = [float(b) for _, b in series]
        fig, ax = plt.subplots(figsize=(width_in, 2.6))
        color = _rgb01_from_hex(QLIK_HEX["green"])
        bars = ax.barh(range(len(labels)), values, color=[color])
        ax.set_yticks(range(len(labels)))
        ax.set_yticklabels(labels, fontsize=9)
        ax.invert_yaxis()
        ax.grid(axis='x', color='#E5E7EB', linewidth=0.8, alpha=0.6)
        ax.set_xlabel("Count", fontsize=9)
        ax.set_title(title, fontsize=11, color='#009845')
        for i, b in enumerate(bars):
            ax.text(b.get_width()+0.1, b.get_y()+b.get_height()/2, str(values[i]), va='center', fontsize=9)
        fig.tight_layout()
        tmp = tempfile.NamedTemporaryFile(delete=False, suffix=".png")
        fig.savefig(tmp.name, dpi=200)
        plt.close(fig)
        doc.add_picture(tmp.name, width=Inches(width_in))
    except Exception:
        _table_2col(doc, "Item", "Value", series)

# =========================
# Data access
# =========================

def _fetch_one(cur, sql: str, *args):
    return cur.execute(sql, args).fetchone()

def _fetch_all(cur, sql: str, *args):
    return cur.execute(sql, args).fetchall()

def _get_customer_name(cur, customer_id) -> Optional[str]:
    if customer_id is None:
        return None
    candidates = [
        "SELECT name FROM repmeta.customer WHERE customer_id=%s",
        "SELECT customer_name FROM repmeta.customer WHERE customer_id=%s",
        "SELECT name FROM customers WHERE id=%s",
    ]
    for q in candidates:
        try:
            row = cur.execute(q, (customer_id,)).fetchone()
            if row:
                return list(row.values())[0]
        except Exception:
            continue
    return None

# =========================
# Main API
# =========================

def generate_qs_report(snapshot_id: int, out_path: str, logo_path: Optional[str] = None) -> str:
    doc = Document()

    # Cover page
    created = datetime.now().strftime("%b %d, %Y %H:%M")
    dsn = _conninfo()
    with psycopg.connect(dsn, row_factory=dict_row) as conn:
        cur = conn.cursor()
        env = _fetch_one(cur, "SELECT * FROM repmeta_qs.v_environment_overview WHERE snapshot_id=%s", snapshot_id) or {}
        customer = _get_customer_name(cur, env.get("customer_id")) or f"Customer #{env.get('customer_id', 'N/A')}"

    _cover_page(
        doc,
        "Qlik Sense — Executive Technical Overview",
        [
            customer,
            f"Snapshot #{snapshot_id}",
            f"Generated {created}",
        ],
        logo_path=logo_path,
    )

    # New section for content with footer
    section = doc.add_section()
    section.left_margin = Inches(0.6)
    section.right_margin = Inches(0.6)
    section.top_margin = Inches(0.5)
    section.bottom_margin = Inches(0.5)
    _footer_with_page_numbers(section)

    # Pull all data we need
    with psycopg.connect(dsn, row_factory=dict_row) as conn:
        cur = conn.cursor()

        lic = _fetch_one(cur, "SELECT * FROM repmeta_qs.v_license_summary WHERE snapshot_id=%s", snapshot_id) or {}
        lic30 = _fetch_one(cur, "SELECT * FROM repmeta_qs.v_license_usage_30d WHERE snapshot_id=%s", snapshot_id) or {}
        apps = _fetch_one(cur, "SELECT * FROM repmeta_qs.v_app_summary WHERE snapshot_id=%s", snapshot_id) or {}
        objs = _fetch_one(cur, "SELECT * FROM repmeta_qs.v_app_objects_summary WHERE snapshot_id=%s", snapshot_id) or {}
        exts = _fetch_one(cur, "SELECT * FROM repmeta_qs.v_extension_summary WHERE snapshot_id=%s", snapshot_id) or {}
        rules = _fetch_one(cur, "SELECT * FROM repmeta_qs.v_security_rule_summary WHERE snapshot_id=%s", snapshot_id) or {}
        rts = _fetch_one(cur, "SELECT * FROM repmeta_qs.v_reload_task_summary WHERE snapshot_id=%s", snapshot_id) or {}
        gov = _fetch_one(cur, "SELECT * FROM repmeta_qs.v_governance_checks WHERE snapshot_id=%s", snapshot_id) or {}

        apps_by_stream = _fetch_all(cur, """
            SELECT COALESCE(stream,'(Unassigned)') AS k, COUNT(*) AS v
            FROM repmeta_qs.v_apps WHERE snapshot_id=%s GROUP BY 1 ORDER BY v DESC, k
        """, snapshot_id)

        reload_enabled = _fetch_all(cur, """
            SELECT CASE WHEN enabled THEN 'Yes' ELSE 'No' END AS k, COUNT(*) AS v
            FROM repmeta_qs.v_reload_tasks WHERE snapshot_id=%s GROUP BY 1 ORDER BY 1
        """, snapshot_id)

        users_by_dir = _fetch_all(cur, """
            SELECT COALESCE(user_directory,'(unknown)') AS k, COUNT(*) AS v
            FROM repmeta_qs.v_users WHERE snapshot_id=%s GROUP BY 1 ORDER BY v DESC, k
        """, snapshot_id)

        reload_activity = _fetch_one(cur, """
            WITH last_by_app AS (
              SELECT snapshot_id, app_id, MAX(last_stop_time) AS last_stop_time
              FROM repmeta_qs.v_reload_tasks WHERE snapshot_id=%s GROUP BY 1,2
            )
            SELECT
              COUNT(*) FILTER (WHERE last_stop_time >= now() - interval '30 days') AS apps_reloaded_30d,
              COUNT(*) FILTER (WHERE last_stop_time >= now() - interval '90 days') AS apps_reloaded_90d
            FROM last_by_app
        """, snapshot_id) or {}

    # Executive blocks
    _h1(doc, "Executive Summary")
    _hr(doc)

    _h2(doc, "Estate Overview")
    _kpi_cards(doc, [
        ("Total Apps", apps.get("total_apps", 0), "info"),
        ("Published Apps", apps.get("published_apps", 0), "ok"),
        ("Streams", apps.get("streams", 0), "info"),
        ("Streams w/ Apps", apps.get("streams_with_apps", 0), "info"),
    ])
    _kpi_cards(doc, [
        ("Users", env.get("user_count", 0), "info"),
        ("Nodes", env.get("node_count", 0), "info"),
        ("Extensions", env.get("extension_count", 0), "info"),
        ("Reload Tasks", env.get("reload_task_count", 0), "info"),
    ])

    _h2(doc, "Reload Health")
    _kpi_cards(doc, [
        ("Apps reloaded (30d)", reload_activity.get("apps_reloaded_30d", 0), "ok"),
        ("Apps reloaded (90d)", reload_activity.get("apps_reloaded_90d", 0), "ok"),
        ("Failed tasks (last run)", rts.get("failed", 0), "bad"),
        ("Tasks > 3h (last run)", rts.get("over_3h", 0), "warn"),
    ])

    _h2(doc, "License")
    _kpi_cards(doc, [
        ("Professional alloc", lic.get("professional_allocations", 0), "info"),
        ("Analyzer alloc", lic.get("analyzer_allocations", 0), "info"),
        ("Valid to", lic.get("expiration", "") or "—", "info"),
        ("License #", lic.get("license_number", "") or "—", "info"),
    ])
    _kpi_cards(doc, [
        ("Analyzer — Used 30d", lic30.get("analyzer_used_30d", 0), "ok"),
        ("Analyzer — Not used 30d", lic30.get("analyzer_not_used_30d", 0), "warn"),
        ("Analyzer — Never used", lic30.get("analyzer_never_used", 0), "bad"),
        ("Professional — Used 30d", lic30.get("professional_used_30d", 0), "ok"),
    ])
    _kpi_cards(doc, [
        ("Professional — Not used 30d", lic30.get("professional_not_used_30d", 0), "warn"),
        ("Professional — Never used", lic30.get("professional_never_used", 0), "bad"),
        ("", "", "info"),
        ("", "", "info"),
    ])

    _h2(doc, "Governance")
    _kpi_cards(doc, [
        ("Apps without reload tasks", gov.get("apps_without_tasks", 0), "bad"),
        ("Disabled reload tasks", gov.get("disabled_tasks_count", 0), "warn"),
        ("Custom security rules", rules.get("custom_rules", 0), "info"),
        ("Disabled security rules", rules.get("disabled_rules", 0), "warn"),
    ])

    # Visuals (charts w/ fallback)
    top_streams = [(r["k"], r["v"]) for r in apps_by_stream[:8]]
    _insert_chart(doc, "Top Streams by App Count", top_streams, width_in=6.0)

    _h1(doc, "Detailed Insights")
    _hr(doc)
    _h2(doc, "Apps by Stream")
    _table_2col(doc, "Stream", "Apps", [(r["k"], r["v"]) for r in apps_by_stream])

    _h2(doc, "Reload Tasks — Enabled split")
    _table_2col(doc, "Enabled", "Count", [(r["k"], r["v"]) for r in reload_enabled])

    _h2(doc, "Users by Directory")
    _table_2col(doc, "Directory", "Users", [(r["k"], r["v"]) for r in users_by_dir])

    _h1(doc, "Environment")
    _hr(doc)
    _table_2col(doc, "Key", "Value", [
        ("Product", env.get("product_name")),
        ("Version", env.get("product_version")),
        ("Build", env.get("build_version")),
        ("Build date", env.get("build_date")),
        ("Single node only", "Yes" if env.get("single_node_only") else "No"),
        ("Streams", env.get("stream_count")),
        ("Users", env.get("user_count")),
        ("Extensions", env.get("extension_count")),
    ])

    # TOC
    _h1(doc, "Contents")
    _hr(doc)
    _add_toc(doc)

    doc.save(out_path)
    return out_path

# ---------------- CLI ----------------

def _parse_args(argv):
    import argparse
    p = argparse.ArgumentParser(description="Generate Qlik‑themed consulting report")
    p.add_argument("snapshot_id", type=int)
    p.add_argument("output", help="Output .docx path")
    p.add_argument("--logo", help="Optional logo image path", default=None)
    return p.parse_args(argv)

def _main():
    import sys
    try:
        args = _parse_args(sys.argv[1:])
        path = generate_qs_report(args.snapshot_id, args.output, logo_path=args.logo)
        print(f"Wrote: {path}")
    except Exception:
        import traceback
        traceback.print_exc()
        raise

if __name__ == "__main__":
    _main()
