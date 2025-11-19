
def _fmt_duration_t90(mins):
    """Smart units for median run: <60 => min, <1440 => hours, else days."""
    try:
        m = float(mins or 0.0)
    except Exception:
        m = 0.0
    if m < 60:
        return f"{m:.2f} min"
    h = m / 60.0
    if m < 1440:
        return f"{h:.2f} h"
    d = h / 24.0
    return f"{d:.2f} d"


import io
import os
import re
import json
import logging
from collections import defaultdict
from datetime import datetime, timezone, date
from typing import Any, Dict, List, Tuple, Optional, NamedTuple

import httpx  # used to fetch latest GA train from GitHub

try:
    # Required for report generation
    from docx import Document
    from docx.shared import Pt
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
except ImportError as e:
    # Make failure obvious and stop early if python-docx isn't present
    raise RuntimeError(
        "python-docx is required for export_report.py. "
        "Install it with:  pip install python-docx"
    ) from e

def _apply_global_styles(doc):
    """Base styles: Calibri 11, tighter spacing, consistent H1–H3."""
    try:
        styles = doc.styles
        normal = styles["Normal"]
        normal.font.name = "Calibri"
        normal.font.size = Pt(11)
        pf = normal.paragraph_format
        pf.space_before = Pt(3)
        pf.space_after = Pt(3)
        pf.line_spacing = 1.15

        for name, size in (("Heading 1", 16), ("Heading 2", 13), ("Heading 3", 12)):
            try:
                st = styles[name]
                st.font.name = "Calibri"
                st.font.size = Pt(size)
            except Exception:
                # If the style doesn't exist, skip silently
                pass
    except Exception:
        # Never block report generation due to styling
        pass


def _apply_footer(doc, footer_text: str = "Confidential — Customer Use Only"):
    """Adds 'Page X of Y' and optional left text; graceful if styles missing."""
    try:
        section = doc.sections[0]
        footer = section.footer
        p = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
        if footer_text:
            run_left = p.add_run(footer_text + "    ")
            run_left.font.size = Pt(8)

        def _add_field(run, instr):
            fld = OxmlElement("w:fldSimple")
            from docx.oxml.ns import qn as _qn
            fld.set(_qn("w:instr"), instr)
            run._r.addnext(fld)

        run = p.add_run()
        _add_field(run, "PAGE")
        p.add_run(" of ")
        run2 = p.add_run()
        _add_field(run2, "NUMPAGES")

        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT

    except Exception:
        pass

from .db import connection

SCHEMA = os.getenv("REPMETA_SCHEMA", "repmeta")
log = logging.getLogger("export_report")


# ============================================================
# Latest Release Fixes (reads from {SCHEMA}.replicate_release_issue)
# ============================================================
from collections import OrderedDict as _OD

def _rowget(row, key, default=None):
    """Return row[key] for asyncpg.Record, dict, or sequence; graceful fallback."""
    try:
        if hasattr(row, "get"):
            return row.get(key, default)
        if key in row:
            return row[key]
    except Exception:
        pass
    return default


async def _load_and_group_latest_release_issues(conn):
    """
    Returns: (latest_label: str, groups: OrderedDict[str, list[dict]])
    Groups latest-train rows by endpoint from {SCHEMA}.replicate_release_issue.
    Uses the passed `conn` if it's open; otherwise opens a short-lived async connection.
    """
    import os, re

    sql = f"""
        SELECT version, issue_date, title, url, jira, endpoints, buckets, text
        FROM {SCHEMA}.replicate_release_issue
        ORDER BY COALESCE(issue_date, DATE '1900-01-01') DESC, version DESC, jira NULLS LAST
    """

    # --- helper to read via your app's wrappers (for the in-scope report connection) ---
    async def _fetch_via_wrappers(c):
        try:
            return await _try_all(c, sql, ())
        except Exception:
            pass
        try:
            return await _all(c, sql, ())
        except Exception:
            pass
        return None

    rows = None

    # 1) Try the incoming connection if it looks open/usable
    try:
        is_closed_flag = getattr(conn, "closed", None)
        is_open = (conn is not None) and (is_closed_flag is False or is_closed_flag == 0)
    except Exception:
        is_open = False

    if is_open:
        rows = await _fetch_via_wrappers(conn)

    # 2) If wrappers didn’t work or conn is closed, open our own async connection and fetch directly
    if rows is None:
        import psycopg
        from psycopg.rows import dict_row

        dsn = os.getenv("DATABASE_URL") or os.getenv("REPMETA_PG_DSN")
        if not dsn:
            log.error("Latest-fixes: no DATABASE_URL/REPMETA_PG_DSN set")
            return "Latest", _OD()

        # IMPORTANT: await the coroutine to get the connection, then use it
        ac = await psycopg.AsyncConnection.connect(dsn)
        try:
            async with ac.cursor(row_factory=dict_row) as cur:
                await cur.execute(sql)
                rows = await cur.fetchall()
        finally:
            try:
                await ac.close()
            except Exception:
                pass

    if not rows:
        return "Latest", _OD()

    # ---- row accessor that works with dict / asyncpg.Record-like / tuple-ish ----
    def _rowget(row, key, default=None):
        try:
            if hasattr(row, "get"):
                return row.get(key, default)
            if key in row:
                return row[key]
        except Exception:
            pass
        return default

    # Prefer latest month by issue_date; fallback to version parsing
    def month_key(r):
        d = _rowget(r, "issue_date")
        return (d.year, d.month) if d else (0, 0)

    dated = [r for r in rows if _rowget(r, "issue_date") is not None]
    latest_label = "Latest"

    if dated:
        y, m = max((month_key(r) for r in dated))
        month_names = [None,"January","February","March","April","May","June","July","August","September","October","November","December"]
        latest_label = f"{month_names[m]} {y}"
        def latest_filter(r):
            d = _rowget(r, "issue_date")
            return d and d.year == y and d.month == m
    else:
        month_map = {
            "january":1,"february":2,"march":3,"april":4,"may":5,"june":6,
            "july":7,"august":8,"september":9,"october":10,"november":11,"december":12
        }
        def parse_vm(v):
            if not v: return (0,0)
            vlow = v.lower()
            y = 0
            m = 0
            m_year = re.search(r"(20\d{2})", vlow)
            if m_year: y = int(m_year.group(1))
            for nm, code in month_map.items():
                if nm in vlow:
                    m = code
                    break
            return (y, m)
        y, m = (0,0)
        for r in rows:
            vy, vm = parse_vm(_rowget(r, "version"))
            if (vy, vm) > (y, m): y, m = vy, vm
        month_names = [None,"January","February","March","April","May","June","July","August","September","October","November","December"]
        latest_label = f"{month_names[m]} {y}" if (y and m) else "Latest"
        def latest_filter(r):
            vy, vm = parse_vm(_rowget(r, "version"))
            return (vy, vm) == (y, m) if (y and m) else True

    latest_rows = [r for r in rows if latest_filter(r)]
    if not latest_rows:
        return latest_label, _OD()

    # Group by endpoints (default "General")
    def endpoints_of(r):
        eps = _rowget(r, "endpoints") or ["General"]
        try:
            return list(eps) if eps else ["General"]
        except Exception:
            return ["General"]

    all_eps = set()
    for r in latest_rows:
        for ep in endpoints_of(r):
            all_eps.add(ep or "General")

    def ep_key(name: str):
        return (name == "General", (name or "").lower())

    groups = _OD()
    for ep in sorted(all_eps, key=ep_key):
        bucket = []
        for r in latest_rows:
            if ep in endpoints_of(r):
                bucket.append({
                    "jira": _rowget(r, "jira") or "-",
                    "title": _rowget(r, "title") or (_rowget(r, "text") or "-"),
                    "buckets": _rowget(r, "buckets") or [],
                    "url": _rowget(r, "url") or "-",
                })
        if bucket:
            groups[ep] = bucket

    return latest_label, groups


def _render_latest_release_fixes_section(doc, latest_label: str, groups):
    """Render the 'Latest Release Fixes' section using groups from DB (no URL column)."""
    _add_heading(doc, "Latest Release Fixes", 1)
    _add_text(
        doc,
        "Because at least one of your servers may not be on the latest GA release, "
        "here is the full set of fixes in the latest release across all endpoints.",
        size=9, italic=True,
    )
    if not groups:
        _add_text(doc, "Release issues are not available in the database.", size=10, italic=True)
        return

    doc.add_paragraph()
    _add_text(doc, f"All Fixes in Latest Release ({latest_label})", size=11, bold=True)

    for ep, items in groups.items():
        doc.add_paragraph()
        _add_text(doc, ep, size=10, bold=True)

        # No URL column anymore
        headers = ["JIRA", "Summary", "Buckets"]
        rows = []
        for r in items:
            # items are dicts from the loader
            jira = (r.get("jira") or "-")
            title = r.get("title") or (r.get("text") or "-")
            if title and len(title) > 200:
                title = title[:197] + "..."
            buckets = r.get("buckets") or []
            if isinstance(buckets, (list, tuple)):
                buckets = ", ".join(buckets)
            elif not buckets:
                buckets = "-"
            rows.append((jira, title, buckets))

        _add_table(doc, headers=headers, rows=rows, style="Light Shading Accent 1")



# ============================================================
# Built-in master lists (fallback if DB tables are absent/empty)
# ============================================================
BUILTIN_MASTER_SOURCE_ENDPOINTS = [
    "Log Stream",  # Added Log Stream as a valid source endpoint type
    "Amazon Aurora (MySQL)",
    "Amazon Aurora (PostgreSQL)",
    "Amazon RDS for MySQL",
    "Amazon RDS for MariaDB",
    "Amazon RDS for PostgreSQL",
    "Amazon RDS for SQL Server",
    "Amazon RDS for Oracle",
    "Google Cloud SQL for MySQL",
    "Google Cloud SQL for PostgreSQL",
    "Google Cloud SQL for SQL Server",
    "Google Cloud AlloyDB for PostgreSQL",
    "Microsoft Azure SQL Database (MS-CDC)",
    "Microsoft Azure SQL Managed Instance",
    "Microsoft Azure Database for MySQL",
    "Microsoft Azure Database for PostgreSQL",
    "MongoDB Atlas",
    "Oracle (on-prem / Oracle Cloud)",
    "Teradata Vantage",
    "IBM DB2 for LUW",
    "IBM DB2 for z/OS",
    "IBM DB2 for iSeries",
    "IBM Informix",
    "Microsoft SQL Server",
    "MySQL",
    "MariaDB",
    "Percona (via MySQL endpoint)",
    "SAP Sybase ASE",
    "SAP HANA 2.0",
    "File endpoint",
    "File Channel endpoint",
    "IBM IMS (ARC)",
    "IBM VSAM Batch (ARC)",
    "Salesforce (Streaming CDC / Incremental Load)",
]

BUILTIN_MASTER_TARGET_ENDPOINTS = [
    "Amazon Aurora (MySQL)",
    "Amazon Aurora (PostgreSQL)",
    "Amazon RDS for MySQL",
    "Amazon RDS for MariaDB",
    "Amazon RDS for PostgreSQL",
    "Amazon RDS for SQL Server",
    "Amazon RDS for Oracle",
    "Google Cloud SQL for MySQL",
    "Google Cloud SQL for PostgreSQL",
    "Google Cloud SQL for SQL Server",
    "Google Cloud AlloyDB for PostgreSQL",
    "Microsoft Azure SQL Database",
    "Microsoft Azure SQL Managed Instance",
    "Microsoft Azure Database for MySQL",
    "Microsoft Azure Database for PostgreSQL",
    "Oracle (on-prem / Oracle Cloud)",
    "IBM DB2 for LUW",
    "IBM DB2 for z/OS",
    "IBM DB2 for iSeries",
    "Microsoft SQL Server",
    "MySQL",
    "PostgreSQL",
    "MariaDB",
    "SAP Sybase ASE",
    "File endpoint",
    "File Channel endpoint",
    "Amazon Redshift",
    "Amazon S3",
    "Amazon MSK",
    "Kafka",
    "Log Stream",
    "Google BigQuery",
    "Google Cloud Storage",
    "Google Dataproc",
    "Microsoft Azure Data Lake / ADLS Gen2 / Blob",
    "Databricks (SQL Warehouse / Lakehouse)",
    "Snowflake",
    "Oracle Autonomous Data Warehouse",
]

# Will be filled at runtime (from DB if available, else fall back to BUILTIN_*)
MASTER_SOURCE_ENDPOINTS: List[str] = []
MASTER_TARGET_ENDPOINTS: List[str] = []
ALIAS_TO_CANON: Dict[str, str] = {}
MASTER_NORM: Dict[str, str] = {}  # set at runtime after loading masters

# ============================================================
# Alias map (fallback) — built from Replicate license tickers and common variants
# ============================================================
def _n(s: str) -> str:
    return re.sub(r"[^a-z0-9]+", "", s.lower())

DEFAULT_ALIAS_TO_CANON: Dict[str, str] = {
    # ---------- Sources ----------
    _n("IMS"): "IBM IMS (ARC)",
    _n("VSAM"): "IBM VSAM Batch (ARC)",
    _n("MySQL"): "MySQL",
    _n("RDSMySQL"): "Amazon RDS for MySQL",
    _n("Oracle"): "Oracle (on-prem / Oracle Cloud)",
    _n("RDSPostgreSQL"): "Amazon RDS for PostgreSQL",
    _n("RDSSQLServer"): "Amazon RDS for SQL Server",
    _n("AWSAuroraPostgreSQL"): "Amazon Aurora (PostgreSQL)",
    _n("AzureSQLServerM"): "Microsoft Azure SQL Database (MS-CDC)",
    _n("File"): "File endpoint",
    _n("FileChannel"): "File Channel endpoint",
    _n("GoogleAlloyDBPostgreSQL"): "Google Cloud AlloyDB for PostgreSQL",
    _n("GoogleCloudMySQL"): "Google Cloud SQL for MySQL",
    _n("GoogleCloudPostgreSQL"): "Google Cloud SQL for PostgreSQL",
    _n("GoogleCloudSQLServer"): "Google Cloud SQL for SQL Server",
    _n("DB2LUW"): "IBM DB2 for LUW",
    _n("DB2iSeries"): "IBM DB2 for iSeries",
    _n("DB2zOS"): "IBM DB2 for z/OS",
    _n("Informix"): "IBM Informix",
    _n("SQLServer"): "Microsoft SQL Server",
    _n("AzureMySQL"): "Microsoft Azure Database for MySQL",
    _n("AzureSQL"): "Microsoft Azure SQL Database (MS-CDC)",
    _n("MongoDBAtlas"): "MongoDB Atlas",
    _n("MongoDB"): "MongoDB Atlas",
    _n("PostgreSQL"): "PostgreSQL",
    _n("SAP HANA"): "SAP HANA 2.0",
    _n("SybaseASE"): "SAP Sybase ASE",
    _n("Teradata"): "Teradata Vantage",
    _n("Salesforce"): "Salesforce (Streaming CDC / Incremental Load)",

    # Common free-form variants seen in QEM/JSON:
    _n("mssql"): "Microsoft SQL Server",
    _n("sql server"): "Microsoft SQL Server",
    _n("azuresqldatabase"): "Microsoft Azure SQL Database",
    _n("alloydb"): "Google Cloud AlloyDB for PostgreSQL",
    _n("gcp alloydb"): "Google Cloud AlloyDB for PostgreSQL",
    _n("percona"): "Percona (via MySQL endpoint)",
    _n("mariadb"): "MariaDB",

    # ---------- Targets ----------
    _n("MySQL"): "MySQL",
    _n("PostgreSQL"): "PostgreSQL",
    _n("Oracle"): "Oracle (on-prem / Oracle Cloud)",
    _n("SQLServer"): "Microsoft SQL Server",
    _n("S3"): "Amazon S3",
    _n("AmazonMSK"): "Amazon MSK",
    _n("Kafka"): "Kafka",
    _n("LogStream"): "Log Stream",
    _n("File"): "File endpoint",
    _n("FileChannel"): "File Channel endpoint",
    _n("GoogleAlloyDBPostgreSQL"): "Google Cloud AlloyDB for PostgreSQL",
    _n("BigQuery"): "Google BigQuery",
    _n("GoogleStorage"): "Google Cloud Storage",
    _n("Dataproc"): "Google Dataproc",
    _n("ADLS"): "Microsoft Azure Data Lake / ADLS Gen2 / Blob",
    _n("AzureSQLServer"): "Microsoft Azure SQL Database",
    _n("AzureMySQL"): "Microsoft Azure Database for MySQL",
    _n("AzurePostgreSQL"): "Microsoft Azure Database for PostgreSQL",
    _n("Redshift"): "Amazon Redshift",
    _n("DB2zOS"): "IBM DB2 for z/OS",
    _n("GCPDatabricksDelta"): "Databricks (SQL Warehouse / Lakehouse)",
    _n("AWSDatabricksDelta"): "Databricks (SQL Warehouse / Lakehouse)",
    _n("DatabricksAWS"): "Databricks (SQL Warehouse / Lakehouse)",
    _n("DatabricksAzure"): "Databricks (SQL Warehouse / Lakehouse)",
    _n("DatabricksGoogleCloud"): "Databricks (SQL Warehouse / Lakehouse)",
    _n("AzureDatabricksDelta"): "Databricks (SQL Warehouse / Lakehouse)",
    _n("SnowflakeAWS"): "Snowflake",
    _n("SnowflakeAzure"): "Snowflake",
    _n("SnowflakeGoogle"): "Snowflake",
    _n("Snowflake"): "Snowflake",

    # more forgiving common strings
    _n("googlegcs"): "Google Cloud Storage",
    _n("gcs"): "Google Cloud Storage",
    _n("ms adls"): "Microsoft Azure Data Lake / ADLS Gen2 / Blob",
    _n("databricks"): "Databricks (SQL Warehouse / Lakehouse)",
}

# Extra permissive synonyms used across both roles
DEFAULT_ALIAS_TO_CANON.update({
    _n("microsoftsqlserver"): "Microsoft SQL Server",
    _n("sqlserver_mscdc"): "Microsoft SQL Server",
    _n("postgres"): "PostgreSQL",
    _n("postgresql"): "PostgreSQL",
    _n("googlebigquery"): "Google BigQuery",
})

# ============================================================
# Utilities: normalization & noise filtering
# ============================================================
def _normalize_token(s: str) -> str:
    return re.sub(r"[^a-z0-9]+", "", (s or "").lower())

_STRIP_TAILS = [" (ms-cdc)"]

def _filter_noise_token(s: str) -> bool:
    return _normalize_token(s) in {"na", "n/a", "null", "nulltarget", "unknown", "(unknown)"}

def _build_master_norm() -> Dict[str, str]:
    return { _normalize_token(name): name for name in (MASTER_SOURCE_ENDPOINTS + MASTER_TARGET_ENDPOINTS) }

def canonize_to_master(name: str, is_source: bool) -> str:
    if not name:
        return "Unknown"
    n = str(name).strip()

    # strip tails
    low = n.lower()
    for t in _STRIP_TAILS:
        if low.endswith(t) and _normalize_token(n) not in MASTER_NORM:
            n = n[: -len(t)]
            break

    key = _normalize_token(n)

    # direct master hit
    if key in MASTER_NORM:
        return MASTER_NORM[key]

    # alias map hit
    if key in ALIAS_TO_CANON:
        return ALIAS_TO_CANON[key]

    # last-ditch: case-insensitive compare within relevant universe
    master = MASTER_SOURCE_ENDPOINTS if is_source else MASTER_TARGET_ENDPOINTS
    for m in master:
        if m.lower() == n.lower():
            return m

    return n

# ============================================================
# DB helpers & dynamic config loaders
# ============================================================
async def _set_row_factory(conn):
    try:
        from psycopg.rows import dict_row  # type: ignore
        await conn.set_row_factory(dict_row)  # psycopg3 async
    except Exception:
        try:
            conn.row_factory = dict_row  # type: ignore[attr-defined]
        except Exception:
            pass

async def _one(conn, sql: str, params: Tuple[Any, ...]) -> Dict[str, Any]:
    cur = await conn.execute(sql, params)
    row = await cur.fetchone()
    return row or {}

async def _all(conn, sql: str, params: Tuple[Any, ...]) -> List[Dict[str, Any]]:
    cur = await conn.execute(sql, params)
    rows = await cur.fetchall()
    return list(rows or [])

async def _try_all(conn, sql: str, params: Tuple[Any, ...]) -> Optional[List[Dict[str, Any]]]:
    try:
        return await _all(conn, sql, params)
    except Exception as e:
        log.debug("optional query failed; rolling back: %s", e)
        try:
            await conn.rollback()
        except Exception:
            pass
        return None

async def _load_master_and_alias_from_db(conn):
    """Populate MASTER_* and ALIAS_TO_CANON from DB if the config tables exist; else fall back."""
    global MASTER_SOURCE_ENDPOINTS, MASTER_TARGET_ENDPOINTS, ALIAS_TO_CANON, MASTER_NORM

    sources = await _try_all(conn, f"SELECT name FROM {SCHEMA}.endpoint_master_sources ORDER BY name", ())
    targets = await _try_all(conn, f"SELECT name FROM {SCHEMA}.endpoint_master_targets ORDER BY name", ())
    MASTER_SOURCE_ENDPOINTS = [r["name"] for r in sources] if sources else BUILTIN_MASTER_SOURCE_ENDPOINTS
    MASTER_TARGET_ENDPOINTS = [r["name"] for r in targets] if targets else BUILTIN_MASTER_TARGET_ENDPOINTS

    alias_rows = await _try_all(conn, f"SELECT alias, canonical FROM {SCHEMA}.endpoint_alias_map", ())
    if alias_rows:
        ALIAS_TO_CANON = { _normalize_token(r["alias"]): r["canonical"] for r in alias_rows }
    else:
        ALIAS_TO_CANON = dict(DEFAULT_ALIAS_TO_CANON)

    MASTER_NORM = _build_master_norm()

# ============================================================
# docx helpers
# ============================================================
def _add_title(doc: Document, text: str):
    p = doc.add_paragraph()
    p.style = doc.styles["Title"]
    r = p.add_run(text)
    r.font.size = Pt(24)
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    return p

def _add_heading(doc: Document, text: str, level: int = 1):
    return doc.add_heading(text, level=level)

def _add_text(doc: Document, text: str, size: int = 11, bold: bool = False, italic: bool = False):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.font.size = Pt(size)
    r.bold = bold
    r.italic = italic
    return p

def _add_toc(doc: Document):
    p = doc.add_paragraph()
    fld = OxmlElement("w:fldSimple")
    fld.set(qn("w:instr"), r'TOC \o "1-3" \h \z \u')
    p._p.append(fld)  # type: ignore[attr-defined]

def _set_cell_shading(cell, fill_hex: str = "EDF2FF"):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), fill_hex)
    tc_pr.append(shd)


def _add_note_panel(doc, title=None, hint=None):
    """
    Render a simple two-row table (header + body) for customer-facing notes.
    Kept intentionally simple so it's fully editable in Word.
    """
    if not title:
        title = "📝 Notes & Observations"
    if not hint:
        hint = "Capture key takeaways, risks, decisions, and next steps…"

    t = doc.add_table(rows=2, cols=1)
    try:
        t.style = "Light Shading Accent 2"
    except Exception:
        pass

    hdr = t.rows[0].cells[0]
    _set_cell_shading(hdr, "E8F5E9")  # pale green
    p = hdr.paragraphs[0] if hdr.paragraphs else hdr.add_paragraph()
    r = p.add_run(title)
    try:
        r.bold = True
    except Exception:
        pass

    body = t.rows[1].cells[0]
    p2 = body.paragraphs[0] if body.paragraphs else body.add_paragraph()
    r2 = p2.add_run(hint)
    try:
        r2.italic = True
    except Exception:
        pass
    # Compact scaffold bullets instead of many blank lines
    for _b in ['• Key takeaway …', '• Risks / blockers …', '• Decisions …', '• Next steps …']:
        ptmp = body.add_paragraph(_b)
        try:
            ptmp.paragraph_format.space_before = Pt(2)
            ptmp.paragraph_format.space_after = Pt(2)
        except Exception:
            pass
    doc.add_paragraph()

# ============================
# Optional Qlik branding (header only)
# ============================
DEFAULT_BRAND_LOGO_URL = os.getenv(
    "REPMETA_BRAND_FIRSTPAGE_LOGO_URL",
    "https://github.com/tejqliky/qlik-repmeta/blob/main/Qlik_logo.png",
)
DEFAULT_BRAND_BANNER_URL = os.getenv(
    "REPMETA_BRAND_BANNER_URL",
    "https://github.com/tejqliky/qlik-repmeta/blob/main/Qlik_Banner.png",
)

def _to_raw_github(url: str) -> str:
    try:
        import re as _re
        m = _re.match(r"https?://github\.com/([^/]+)/([^/]+)/blob/([^/]+)/(.*)", url or "", flags=_re.I)
        if m:
            user, repo, branch, rest = m.groups()
            return f"https://raw.githubusercontent.com/{user}/{repo}/{branch}/{rest}"
    except Exception:
        pass
    return url

def _try_fetch_bytes(url: str):
    try:
        import urllib.request
        req = urllib.request.Request(url, headers={"User-Agent": "repmeta-report/1.0"})
        with urllib.request.urlopen(req, timeout=20) as r:
            return r.read()
    except Exception:
        return None

def _read_brand_asset(url: str, path_env: str):
    # URL takes precedence; path is fallback
    u = _to_raw_github(url) if url else None
    if u:
        data = _try_fetch_bytes(u)
        if data:
            return data
    p = os.getenv(path_env)
    if p and os.path.exists(p):
        try:
            with open(p, "rb") as f:
                return f.read()
        except Exception:
            return None
    return None

def _apply_branding(doc: Document):
    """
    Adds Qlik logo on the FIRST page header and a green banner in the header of ALL pages.
    Skips silently if assets are unavailable. Does NOT alter body styles/colors/logic.
    """
    from docx.shared import Inches
    import io as _io

    if not getattr(doc, "sections", None):
        return

    logo_bytes = _read_brand_asset(DEFAULT_BRAND_LOGO_URL, "REPMETA_BRAND_FIRSTPAGE_LOGO_PATH")
    banner_bytes = _read_brand_asset(DEFAULT_BRAND_BANNER_URL, "REPMETA_BRAND_BANNER_PATH")

    if not logo_bytes and not banner_bytes:
        return

    sec = doc.sections[0]
    try:
        sec.different_first_page_header_footer = True
    except Exception:
        pass

    def _place_image_in_header(header_obj, data: bytes, width_in: float, align_left: bool = True):
        if not data:
            return
        p = header_obj.paragraphs[0] if header_obj.paragraphs else header_obj.add_paragraph()
        r = p.add_run()
        try:
            r.add_picture(_io.BytesIO(data), width=Inches(width_in))
            try:
                p.alignment = WD_ALIGN_PARAGRAPH.LEFT if align_left else WD_ALIGN_PARAGRAPH.CENTER
            except Exception:
                pass
        except Exception:
            # If picture decoding fails, skip silently
            return

    # Choose reasonable widths
    banner_width = 6.5

    # First page header
    h_first = getattr(sec, "first_page_header", None) or sec.header
    if h_first:
        if logo_bytes:
            _place_image_in_header(h_first, logo_bytes, width_in=1.8, align_left=True)
        if banner_bytes:
            _place_image_in_header(h_first, banner_bytes, width_in=banner_width, align_left=True)

    # Default header for pages 2+
    h_std = sec.header
    if h_std and banner_bytes:
        _place_image_in_header(h_std, banner_bytes, width_in=banner_width, align_left=True)


def _cell_bold(cell, size: int = 12):
    for p in cell.paragraphs:
        for r in p.runs:
            r.bold = True
            r.font.size = Pt(size)

def _add_table(doc: Document, headers: List[str], rows: List[Tuple[Any, ...]], style: str = "Light Shading Accent 1"):
    t = doc.add_table(rows=1, cols=len(headers))
    t.style = style
    hdr = t.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = str(h)
        _cell_bold(hdr[i], 10)
    for r in rows:
        row = t.add_row().cells
        for i, v in enumerate(r):
            row[i].text = "" if v is None else str(v)
            try:
                s_val = ('' if v is None else str(v)).strip()
                if s_val.endswith('%'):
                    s_num = s_val[:-1]
                else:
                    s_num = s_val
                float(s_num.replace(',', ''))
                row[i].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
            except Exception:
                pass
    return t

def _kpi_cards(doc: Document, cards: List[Tuple[str, Any, str]]):
    t = doc.add_table(rows=1, cols=len(cards))
    for i, (label, value, color) in enumerate(cards):
        c = t.rows[0].cells[i]
        _set_cell_shading(c, color)
        p1 = c.paragraphs[0]
        p1.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r1 = p1.add_run(str(value))
        r1.font.size = Pt(18)
        r1.bold = True
        p2 = c.add_paragraph()
        p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r2 = p2.add_run(label)
        r2.font.size = Pt(9)
        r2.italic = True
    return t

def _fmt_int(x: Any) -> str:
    try:
        return f"{int(str(x).replace(',', '').strip() or '0'):,}"
    except Exception:
        return str(x)

# ==== Metrics helpers (added) ==================================================
def _fmt_bytes(n: float) -> str:
    try:
        x = float(n or 0)
    except Exception:
        return str(n)
    units = ["B","KB","MB","GB","TB","PB"]
    k = 1000.0
    i = 0
    while x >= k and i < len(units)-1:
        x /= k
        i += 1
    if i == 0:
        return f"{int(x)} {units[i]}"
    return f"{x:.1f} {units[i]}"

async def _metrics_monthly_current_year(conn, customer_id: int):
    """
    Monthly Volume (Last 6 Months)
    Load = latest per (month, task) from v_metrics_events_clean
    CDC  = sum(cdc_bytes) within the month from v_metrics_events_clean
    """
    sql = f"""
        WITH base AS (
          SELECT
            date_trunc('month', ts)::date AS m,
            customer_id, server_id, task_id, task_uuid,
            ts, load_bytes, cdc_bytes
          FROM {SCHEMA}.v_metrics_events_clean
          WHERE customer_id = %s
        ),
        bounds AS (
          SELECT date_trunc('month', MAX(ts)) AS max_month
          FROM {SCHEMA}.v_metrics_events_clean
          WHERE customer_id = %s
        ),
        within6 AS (
          SELECT b.*
          FROM base b, bounds
          WHERE b.ts >= (bounds.max_month - INTERVAL '5 months')
            AND b.ts  < (bounds.max_month + INTERVAL '1 month')
        ),
        latest_per_task_month AS (
          SELECT m, COALESCE(task_id::text, task_uuid) AS tkey, load_bytes, ts,
                 ROW_NUMBER() OVER (PARTITION BY m, COALESCE(task_id::text, task_uuid) ORDER BY ts DESC) AS rn
          FROM within6
        ),
        load_month AS (
          SELECT m, SUM(load_bytes)::bigint AS load_b
          FROM latest_per_task_month
          WHERE rn = 1
          GROUP BY m
        ),
        cdc_month AS (
          SELECT m, SUM(cdc_bytes)::bigint AS cdc_b
          FROM within6
          GROUP BY m
        )
        SELECT COALESCE(l.m, c.m) AS m,
               COALESCE(l.load_b,0)::bigint AS load_b,
               COALESCE(c.cdc_b,0)::bigint  AS cdc_b
        FROM load_month l
        FULL OUTER JOIN cdc_month c USING (m)
        ORDER BY m;
    """
    rows = await _try_all(conn, sql, (customer_id, customer_id)) or []
    out = []
    for r in rows:
        m  = r["m"]
        lb = int(r.get("load_b") or 0)
        cb = int(r.get("cdc_b") or 0)
        out.append((str(m), _fmt_bytes(lb), _fmt_bytes(cb), _fmt_bytes(lb + cb)))
    return out
async def _metrics_yearly_last5(conn, customer_id: int):
    """
    Annual Volume (Last 5 Years)
    Load = latest per (year, task) from v_metrics_events_clean
    CDC  = sum(cdc_bytes) within the year from v_metrics_events_clean
    """
    sql = f"""
        WITH base AS (
          SELECT
            date_part('year', ts)::int AS y,
            customer_id, server_id, task_id, task_uuid,
            ts, load_bytes, cdc_bytes
          FROM {SCHEMA}.v_metrics_events_clean
          WHERE customer_id = %s
            AND date_part('year', ts) >= date_part('year', CURRENT_DATE) - 4
        ),
        latest_per_task_year AS (
          SELECT y, COALESCE(task_id::text, task_uuid) AS tkey, load_bytes, ts,
                 ROW_NUMBER() OVER (PARTITION BY y, COALESCE(task_id::text, task_uuid) ORDER BY ts DESC) AS rn
          FROM base
        ),
        load_year AS (
          SELECT y, SUM(load_bytes)::bigint AS load_b
          FROM latest_per_task_year
          WHERE rn = 1
          GROUP BY y
        ),
        cdc_year AS (
          SELECT y, SUM(cdc_bytes)::bigint AS cdc_b
          FROM base
          GROUP BY y
        )
        SELECT COALESCE(l.y, c.y) AS y,
               COALESCE(l.load_b,0)::bigint AS load_b,
               COALESCE(c.cdc_b,0)::bigint  AS cdc_b
        FROM load_year l
        FULL OUTER JOIN cdc_year c USING (y)
        ORDER BY y;
    """
    rows = await _try_all(conn, sql, (customer_id,)) or []
    out = []
    for r in rows:
        y  = int(r["y"])
        lb = int(r.get("load_b") or 0)
        cb = int(r.get("cdc_b") or 0)
        out.append((str(y), _fmt_bytes(lb), _fmt_bytes(cb), _fmt_bytes(lb + cb)))
    return out
async def _metrics_top_tasks(conn, customer_id: int, limit: int = 5):
    """
    Top 5 Tasks by Volume (Load + CDC)
    Load = load_bytes from v_metrics_task_latest_event (latest per task)
    CDC  = sum(cdc_bytes) from v_metrics_events_clean
    """
    sql = f"""
        WITH load_latest AS (
          SELECT customer_id, server_id, task_id, task_uuid, load_bytes
          FROM {SCHEMA}.v_metrics_task_latest_event
          WHERE customer_id = %s
        ),
        cdc_sum AS (
          SELECT customer_id, server_id, task_id, task_uuid, SUM(cdc_bytes)::bigint AS cdc_b
          FROM {SCHEMA}.v_metrics_events_clean
          WHERE customer_id = %s
          GROUP BY 1,2,3,4
        ),
        joined AS (
          SELECT
            COALESCE(l.customer_id, c.customer_id) AS customer_id,
            COALESCE(l.server_id,  c.server_id)  AS server_id,
            COALESCE(l.task_id,    c.task_id)    AS task_id,
            COALESCE(l.task_uuid,  c.task_uuid)  AS task_uuid,
            COALESCE(l.load_bytes, 0)::bigint    AS load_b,
            COALESCE(c.cdc_b,      0)::bigint    AS cdc_b
          FROM load_latest l
          FULL OUTER JOIN cdc_sum c
          ON l.customer_id = c.customer_id
         AND l.server_id  = c.server_id
         AND COALESCE(l.task_id::text, l.task_uuid) = COALESCE(c.task_id::text, c.task_uuid)
        ),
        name_resolved AS (
          SELECT j.*, ds.server_name,
                 (
                   SELECT rt.task_name
                   FROM {SCHEMA}.rep_task rt
                   JOIN {SCHEMA}.ingest_run ir2 ON ir2.run_id = rt.run_id
                   WHERE ir2.customer_id = j.customer_id
                     AND ( (j.task_id  IS NOT NULL AND rt.task_id  = j.task_id)
                        OR (j.task_uuid IS NOT NULL AND rt.task_uuid = j.task_uuid) )
                   ORDER BY rt.run_id DESC
                   LIMIT 1
                 ) AS task_name
          FROM joined j
          JOIN {SCHEMA}.dim_server ds ON ds.server_id = j.server_id
        )
        SELECT
          COALESCE(task_name, '(unknown) ' || LEFT(COALESCE(task_uuid::text,''),8)) AS task_label,
          server_name,
          load_b, cdc_b, (load_b + cdc_b) AS total_b
        FROM name_resolved
        ORDER BY total_b DESC NULLS LAST
        LIMIT {limit};
    """
    rows = await _try_all(conn, sql, (customer_id, customer_id)) or []
    out = []
    for r in rows:
        out.append((
            r["task_label"],
            r["server_name"] or "",
            _fmt_bytes(int(r["load_b"] or 0)),
            _fmt_bytes(int(r["cdc_b"]  or 0)),
            _fmt_bytes(int(r["total_b"] or 0)),
        ))
    return out
async def _metrics_top_endpoints(conn, customer_id: int, role: str, metric: str, limit: int = 5):
    """
    Top endpoints using only the latest metrics-log run per server.
    role   : "SOURCE" | "TARGET"
    metric : "load" | "cdc"
    For 'load', use v_metrics_task_latest_event (latest per task).
    For 'cdc', sum from v_metrics_events_clean.
    """
    assert role in ("SOURCE", "TARGET")
    assert metric in ("load", "cdc")

    fam_id_col = "source_family_id" if role == "SOURCE" else "target_family_id"
    type_col   = "source_type"      if role == "SOURCE" else "target_type"

    if metric == "load":
        sql = f"""
            WITH latest AS (
              SELECT customer_id, server_id, {fam_id_col} AS fam_id, {type_col} AS type_label, load_bytes
              FROM {SCHEMA}.v_metrics_task_latest_event
              WHERE customer_id = %s
            )
            SELECT COALESCE(f.family_name, l.type_label) AS endpoint_label,
                   SUM(l.load_bytes)::bigint AS vol_bytes
            FROM latest l
            LEFT JOIN {SCHEMA}.endpoint_family f ON f.family_id = l.fam_id
            GROUP BY 1
            ORDER BY vol_bytes DESC NULLS LAST
            LIMIT {limit};
        """
        rows = await _try_all(conn, sql, (customer_id,)) or []
        out = []
        for r in rows:
            lbl = r.get("endpoint_label")
            if not lbl:
                continue
            lbl = canonize_to_master(lbl, is_source=(role == "SOURCE"))
            out.append((lbl, _fmt_bytes(int(r.get("vol_bytes") or 0))))
        return out

    # metric == "cdc"
    sql = f"""
        SELECT COALESCE(f.family_name, e.{type_col}) AS endpoint_label,
               SUM(e.cdc_bytes)::bigint AS vol_bytes
        FROM {SCHEMA}.v_metrics_events_clean e
        LEFT JOIN {SCHEMA}.endpoint_family f ON f.family_id = e.{fam_id_col}
        WHERE e.customer_id = %s
        GROUP BY 1
        ORDER BY vol_bytes DESC NULLS LAST
        LIMIT {limit};
    """
    rows = await _try_all(conn, sql, (customer_id,)) or []
    out = []
    for r in rows:
        lbl = r.get("endpoint_label")
        if not lbl:
            continue
        lbl = canonize_to_master(lbl, is_source=(role == "SOURCE"))
        out.append((lbl, _fmt_bytes(int(r.get("vol_bytes") or 0))))
    return out
def _add_metrics_section(doc, title, rows, headers):
    _add_text(doc, title, size=12, bold=True)
    if not rows:
        _add_text(doc, "No Metrics Log data available yet.", size=10, italic=True)
        return
    _add_table(doc, headers=headers, rows=rows, style="Light Shading Accent 1")
# ==== end Metrics helpers ======================================================



def _pretty_type(raw: Any) -> str:
    if not raw:
        return "Unknown"
    s = str(raw)
    rep = {
        "Log Stream": "Log Stream",  # Added explicit mapping for Log Stream
        "Sqlserver": "Microsoft SQL Server",
        "Postgresql": "PostgreSQL",
        "Oracle": "Oracle (on-prem / Oracle Cloud)",
        "Snowflake": "Snowflake",
        "Bigquery": "Google BigQuery",
        "Kafka": "Kafka",
        "Filechannel": "File Channel endpoint",
        "Mysql": "MySQL",
        "Db2": "IBM DB2 for LUW",
        "Redshift": "Amazon Redshift",
        "S3": "Amazon S3",
        "Databricks": "Databricks (SQL Warehouse / Lakehouse)",
        "Gcs": "Google Cloud Storage",
        "Adls": "Microsoft Azure Data Lake / ADLS Gen2 / Blob",
    }
    # Check for exact match first (case-insensitive)
    for k, v in rep.items():
        if k.lower() == s.lower():
            return v
    # Then check for partial match
    low = s.lower()
    for k, v in rep.items():
        if k.lower() in low:
            return v
    s = s.replace("Settings", "").replace("Source", "").replace("Target", "")
    return s.strip().title() or "Unknown"

def _type_icon(pretty: str) -> str:
    m = {
        "Log Stream": "🌊",  # Added icon for Log Stream
        "Microsoft SQL Server": "🗄️",
        "PostgreSQL": "🐘",
        "Oracle (on-prem / Oracle Cloud)": "🟥",
        "Snowflake": "❄️",
        "Google BigQuery": "🔷",
        "Kafka": "🔷",
        "File Channel endpoint": "📁",
        "MySQL": "🐬",
        "IBM DB2 for LUW": "🟣",
        "Amazon Redshift": "📊",
        "Amazon S3": "🪣",
        "Databricks (SQL Warehouse / Lakehouse)": "🔥",
        "Google Cloud Storage": "☁️",
        "Microsoft Azure Data Lake / ADLS Gen2 / Blob": "🗂️",
        "Unknown": "🔹",
    }
    return m.get(pretty, "🔹")

def _version_badge(version: Optional[str]) -> str:
    if not version or version.strip() in ("", "-"):
        return "-  ⚠️"
    return version

def _age_badge(dt: Optional[datetime]) -> str:
    if not dt:
        return "-  ⚠️"
    now = datetime.now(timezone.utc)
    if dt.tzinfo is None:
        dt = dt.replace(tzinfo=timezone.utc)
    day_str = dt.astimezone(timezone.utc).date().isoformat()
    days = (now - dt).days
    if days > 60:
        return f"{day_str}  🔴 {days}d"
    if days > 30:
        return f"{day_str}  🟠 {days}d"
    return f"{day_str}  🟢 {days}d"

# ============================================================
# Release/train helpers
# ============================================================
TRAIN_RE = re.compile(r"^v?(\d{4})\.(\d{1,2})\.(\d+)$")  # vYYYY.M.SR

class Train(NamedTuple):
    year: int
    month_code: int  # 5 for May, 11 for Nov
    sr: int

def _parse_tag_to_train(tag: str) -> Optional[Train]:
    if not tag:
        return None
    m = TRAIN_RE.match(tag.strip())
    if not m:
        return None
    y, mcode, sr = map(int, m.groups())
    return Train(y, mcode, sr)

def _parse_replicate_version_to_train(version_str: Optional[str]) -> Optional[Train]:
    if not version_str:
        return None
    s = str(version_str).strip()
    nums = re.findall(r"\d+", s)
    if len(nums) < 3:
        return None
    y, mcode, sr = map(int, nums[:3])
    return Train(y, mcode, sr)

def _train_rank_key(t: Train) -> int:
    return t.year * 100 + t.month_code

def _trains_behind(latest: Train, current: Train) -> int:
    return max(0, _train_rank_key(latest) - _train_rank_key(current))

def _posture_label(delta: int) -> str:
    if delta == 0:
        return "🟢 Up-to-date"
    if delta == 1:
        return "🟡 1 train behind"
    if delta == 2:
        return "🟠 2 trains behind"
    if delta == 999:
        return "⚪ Unknown"
    return "🔴 >2 behind (Out-of-Support)"

async def _get_latest_ga_train(conn) -> Optional[Train]:
    row = await _one(
        conn,
        f"""
        SELECT tag, year, month_code, sr
        FROM {SCHEMA}.replicate_latest_release_cache
        ORDER BY fetched_at DESC
        LIMIT 1
        """,
        (),
    )
    if not row:
        return None
    return Train(int(row["year"]), int(row["month_code"]), int(row["sr"]))

async def _ensure_latest_cache(conn) -> Optional[Train]:
    latest_cached = await _get_latest_ga_train(conn)

    gh_api = "https://api.github.com/repos/qlik-download/replicate/releases"
    headers = {}
    tok = os.getenv("GITHUB_TOKEN")
    if tok:
        headers["Authorization"] = f"Bearer {tok}"

    try:
        async with httpx.AsyncClient(timeout=20) as cli:
            r = await cli.get(gh_api, headers=headers)
            r.raise_for_status()
            releases = r.json()
        newest: Optional[Train] = None
        newest_tag: Optional[str] = None

        for rel in releases:
            if rel.get("draft") or rel.get("prerelease"):
                continue
            tag = rel.get("tag_name") or ""
            t = _parse_tag_to_train(tag)
            if not t:
                continue
            if not newest or _train_rank_key(t) > _train_rank_key(newest):
                newest = t
                newest_tag = tag

        if newest:
            if (not latest_cached) or (_train_rank_key(newest) > _train_rank_key(latest_cached)):
                try:
                    await conn.execute(
                        f"""
                        INSERT INTO {SCHEMA}.replicate_latest_release_cache(tag, year, month_code, sr)
                        VALUES (%s,%s,%s,%s)
                        """,
                        (newest_tag, newest.year, newest.month_code, newest.sr),
                    )
                except Exception as e:
                    log.debug("cache insert failed; rolling back and continuing: %s", e)
                    try:
                        await conn.rollback()
                    except Exception:
                        pass
                latest_cached = newest
    except Exception as e:
        log.warning("GitHub latest GA fetch failed; using cache if present. err=%s", e)

    return await _get_latest_ga_train(conn)

# ============================================================
# License usage (modern layout, no "unlicensed in use" row)
# ============================================================
def _names_from_rows(rows: List[Dict[str, Any]], is_source: bool) -> set:
    out = set()
    for r in rows or []:
        raw = r.get("type")
        if not raw:
            continue
        s = str(raw).strip()
        if _filter_noise_token(s):
            continue
        pretty = canonize_to_master(s, is_source=is_source)
        if _filter_noise_token(pretty):
            continue
        out.add(pretty)
    return out

def _wrap_join(items: List[str]) -> str:
    return ", ".join(items) if items else "-"

def _license_pill_table(doc: Document, title: str,
                        used: List[str], licensed_not_used: List[str]):
    _add_text(doc, title, size=11, bold=True)
    t = _add_table(
        doc,
        headers=["", ""],
        rows=[
            ("Used", _wrap_join(used)),
            ("Licensed not used", _wrap_join(licensed_not_used)),
        ],
        style="Light Shading Accent 1",
    )
    colors = ["EEF2FF", "F1F8E9"]
    for i in range(2):
        _set_cell_shading(t.rows[i+1].cells[0], colors[i])
        _cell_bold(t.rows[i+1].cells[0], 10)

def _license_usage_section(doc: Document,
                           used_src: set, used_tgt: set,
                           lic_all_src: bool, lic_all_tgt: bool,
                           lic_src: set, lic_tgt: set):
    # Licensed universes
    lic_src_universe = set(MASTER_SOURCE_ENDPOINTS) if lic_all_src else set(lic_src)
    lic_tgt_universe = set(MASTER_TARGET_ENDPOINTS) if lic_all_tgt else set(lic_tgt)

    # Coverage math (ignore "unlicensed in use" – Replicate won’t allow it)
    src_used_ct = len(used_src if lic_all_src else (used_src & lic_src_universe))
    src_total = None if lic_all_src else len(lic_src_universe)
    src_not_used = [] if lic_all_src else sorted(lic_src_universe - used_src)

    tgt_used_ct = len(used_tgt if lic_all_tgt else (used_tgt & lic_tgt_universe))
    tgt_total = None if lic_all_tgt else len(lic_tgt_universe)
    tgt_not_used = [] if lic_all_tgt else sorted(lic_tgt_universe - used_tgt)

    # KPI tiles
    _add_text(doc, "License Usage", size=12, bold=True)
    src_kpi = f"{src_used_ct} (All Licensed)" if src_total is None else f"{src_used_ct} / {src_total}"
    tgt_kpi = f"{tgt_used_ct} (All Licensed)" if tgt_total is None else f"{tgt_used_ct} / {tgt_total}"
    _kpi_cards(doc, [
        ("Licensed Sources Used", src_kpi, "E8F5E9"),
        ("Licensed Targets Used", tgt_kpi, "E3F2FD"),
    ])
    doc.add_paragraph()

    # Panels
    _license_pill_table(doc, "Sources", sorted(used_src if lic_all_src else (used_src & lic_src_universe)), src_not_used)
    doc.add_paragraph()
    _license_pill_table(doc, "Targets", sorted(used_tgt if lic_all_tgt else (used_tgt & lic_tgt_universe)), tgt_not_used)


# ============================================================
# Flow diagram helpers (customer-level and server-level)
def _add_flow_pair_table(doc: Document, edges, title: str, max_rows: int = 500):
    """Render a simple three-column table sorted by Tasks desc: [Source | Target | Tasks]"""
    _add_text(doc, title, size=12, bold=True)
    if not edges:
        _add_text(doc, "No flow data available.", size=10, italic=True)
        return
    # Aggregate duplicate pairs and sanitize
    agg = {}
    for s, t, n in (edges or []):
        try:
            key = (str(s).strip(), str(t).strip())
            agg[key] = agg.get(key, 0) + int(n or 0)
        except Exception:
            continue
    rows_data = [(s, t, n) for (s, t), n in agg.items()]
    rows_data.sort(key=lambda e: (-int(e[2]), str(e[0]), str(e[1])))
    headers = ["Source", "Target", "Tasks"]
    rows = [(s, t, _fmt_int(n)) for s, t, n in rows_data[:max_rows]]
    _add_table(doc, headers=headers, rows=rows, style="Light Shading Accent 1")

# ============================================================
from tempfile import NamedTemporaryFile

_FLOW_NOISE = {"na", "n/a", "null", "null target", "unknown", "(unknown)"}

def _flow_is_noise(label: str) -> bool:
    return (not label) or (label.strip().lower() in _FLOW_NOISE)

def _flow_edges_from_coverage_rows(rows):
    """rows with keys s_type, t_type, n -> list of (src, tgt, n) after pretty + noise filter"""
    agg = {}
    for r in rows or []:
        s = _pretty_type(r.get("s_type"))
        t = _pretty_type(r.get("t_type"))
        if _flow_is_noise(s) or _flow_is_noise(t):
            continue
        n = int(r.get("n") or 0)
        if n <= 0:
            continue
        key = (s.strip(), t.strip())
        agg[key] = agg.get(key, 0) + n
    return [(s, t, n) for (s, t), n in agg.items()]

async def _gather_server_edges_map(conn, customer_id: int):
    """
    Returns {server_name: [(src, tgt, n), ...]}.
    Prefers a view v_coverage_matrix_by_server if present; falls back to TSV group-by.
    """
    rows = await _try_all(
        conn,
        f"SELECT server_name, s_type, t_type, n FROM {SCHEMA}.v_coverage_matrix_by_server WHERE customer_id=%s",
        (customer_id,),
    )
    if rows is None:
        rows = await _all(
            conn,
            f"""
            SELECT s.server_name,
                   COALESCE(q.source_type,'(unknown)') AS s_type,
                   COALESCE(q.target_type,'(unknown)') AS t_type,
                   COUNT(*) AS n
            FROM {SCHEMA}.qem_task_perf q
            JOIN {SCHEMA}.dim_server s USING (server_id)
            WHERE q.customer_id=%s
            GROUP BY s.server_name, COALESCE(q.source_type,'(unknown)'), COALESCE(q.target_type,'(unknown)')
            """,
            (customer_id,),
        )
    by_server = {}
    for r in rows or []:
        sname = r.get("server_name")
        if not sname:
            continue
        by_server.setdefault(sname, [])
        by_server[sname].append({"s_type": r.get("s_type"), "t_type": r.get("t_type"), "n": r.get("n")})
    # normalize + filter noise per server
    out = {}
    for sname, lst in by_server.items():
        out[sname] = _flow_edges_from_coverage_rows(lst)
    return out

def _render_flow_png(edges, max_edges=20, width_inches=6.5):
    """
    Try Graphviz first; fall back to Matplotlib. Return PNG bytes or None.
    """
    # Order and cap edges
    edges = sorted(edges or [], key=lambda e: (-int(e[2]), str(e[0]), str(e[1])))[:max_edges]
    if not edges:
        return None

    # 1) Graphviz
    try:
        from graphviz import Digraph
        g = Digraph("flows", format="png")
        g.attr(rankdir="LR", splines="spline", fontname="Calibri")
        g.attr("node", shape="box", style="rounded,filled", color="#4666A5", fillcolor="#EEF2FF", fontname="Calibri", fontsize="10")
        g.attr("edge", color="#4E5D78", fontname="Calibri", fontsize="9")

        sources = sorted({e[0] for e in edges})
        targets = sorted({e[1] for e in edges})
        max_n = max(int(e[2]) for e in edges) if edges else 1

        with g.subgraph(name="cluster_sources") as ssg:
            ssg.attr(rank="same", color="white")
            for s in sources:
                ssg.node(f"S::{s}", label=s)

        with g.subgraph(name="cluster_targets") as tsg:
            tsg.attr(rank="same", color="white")
            for t in targets:
                tsg.node(f"T::{t}", label=t)

        for s, t, n in edges:
            penw = 1 + (5 * (int(n) / max_n))
            g.edge(f"S::{s}", f"T::{t}", label=str(n), penwidth=str(penw))

        tmp = NamedTemporaryFile(suffix=".png", delete=False)
        tmp.close()
        out_path = tmp.name
        g.render(filename=out_path, cleanup=True)  # graphviz appends .png
        png_path = out_path + ".png"
        with open(png_path, "rb") as f:
            data = f.read()
        try:
            import os
            os.remove(png_path)
        except Exception:
            pass
        return data
    except Exception:
        pass

    # 2) Matplotlib fallback
    try:
        import matplotlib.pyplot as plt
        from matplotlib.patches import FancyBboxPatch, FancyArrowPatch

        # layout
        sources = sorted({e[0] for e in edges})
        targets = sorted({e[1] for e in edges})
        left_x, right_x = 0.05, 0.75
        src_y_gap = 0.8 / max(1, len(sources))
        tgt_y_gap = 0.8 / max(1, len(targets))
        src_pos = {s: (left_x, 0.9 - i * src_y_gap) for i, s in enumerate(sources)}
        tgt_pos = {t: (right_x, 0.9 - i * tgt_y_gap) for i, t in enumerate(targets)}

        fig, ax = plt.subplots(figsize=(width_inches, 4.0))

        def draw_node(ax, x, y, text):
            w, h = 0.18, 0.06
            box = FancyBboxPatch((x, y - h / 2), w, h, boxstyle="round,pad=0.02")
            ax.add_patch(box)
            ax.text(x + w / 2, y, text, ha="center", va="center", fontsize=9)

        for s, (x, y) in src_pos.items():
            draw_node(ax, x, y, s)
        for t, (x, y) in tgt_pos.items():
            draw_node(ax, x, y, t)

        max_count = max(int(c) for _, _, c in edges) if edges else 1

        for s, t, c in edges:
            (x1, y1) = src_pos[s]
            (x2, y2) = tgt_pos[t]
            x1r = x1 + 0.18
            x2l = x2
            lw = 1 + 5 * (int(c) / max_count)
            arrow = FancyArrowPatch((x1r, y1), (x2l, y2), arrowstyle="->", mutation_scale=12, linewidth=lw)
            ax.add_patch(arrow)
            xm = (x1r + x2l) / 2
            ym = (y1 + y2) / 2
            ax.text(xm, ym + 0.02, str(c), ha="center", va="bottom", fontsize=9)

        ax.set_xlim(0, 1)
        ax.set_ylim(0, 1)
        ax.axis("off")

        import io as _io
        buf = _io.BytesIO()
        fig.tight_layout()
        fig.savefig(buf, format="png", dpi=180, bbox_inches="tight")
        plt.close(fig)
        buf.seek(0)
        return buf.getvalue()
    except Exception:
        return None

def _add_flow_png_to_doc(doc: Document, png_bytes: bytes, width_inches: float = 6.5):
    from docx.shared import Inches
    import io as _io
    buf = _io.BytesIO(png_bytes)
    doc.add_picture(buf, width=Inches(width_inches))

# ============================================================
# Server-level report (retained)
# ============================================================
async def generate_summary_docx(customer_name: str, server_name: str) -> Tuple[bytes, str]:
    async with connection() as conn:
        await _set_row_factory(conn)
        await _load_master_and_alias_from_db(conn)

        run_row = await _one(
            conn,
            f"""
            SELECT r.run_id
            FROM {SCHEMA}.ingest_run r
            JOIN {SCHEMA}.dim_customer c ON c.customer_id = r.customer_id
            JOIN {SCHEMA}.dim_server   s ON s.server_id   = r.server_id
            WHERE c.customer_name=%s AND s.server_name=%s
            ORDER BY r.run_id DESC
            LIMIT 1
            """,
            (customer_name, server_name),
        )
        if not run_row:
            raise ValueError(
                f"No run found for customer='{customer_name}' server='{server_name}'. Ingest a repo JSON first."
            )
        run_id = run_row["run_id"]

        tasks_row = await _one(conn, f"SELECT COUNT(*) AS n FROM {SCHEMA}.rep_task WHERE run_id=%s", (run_id,))
        endpoints_row = await _one(conn, f"SELECT COUNT(*) AS n FROM {SCHEMA}.rep_database WHERE run_id=%s", (run_id,))
        tasks_count = int(tasks_row.get("n", 0))
        endpoints_count = int(endpoints_row.get("n", 0))

        role_counts = await _all(
            conn,
            f"SELECT role, COUNT(*) AS n FROM {SCHEMA}.rep_database WHERE run_id=%s GROUP BY role ORDER BY role",
            (run_id,),
        )
        role_map = {str(r.get("role")): int(r.get("n", 0)) for r in role_counts}
        src_n = role_map.get("SOURCE", 0)
        tgt_n = role_map.get("TARGET", 0)

        doc = Document()
        try:
            _apply_global_styles(doc)
            _apply_footer(doc)
        except Exception:
            pass
        try:
            _apply_branding(doc)
        except Exception:
            pass
    _add_title(doc, "Qlik Replicate - Server Review")
    _add_text(doc, f"Customer: {customer_name}", size=10)
    _add_text(doc, f"Server: {server_name}", size=10)
    _add_text(doc, f"Run ID: {run_id}", size=10)
    doc.add_paragraph()
    _add_toc(doc)
    doc.add_page_break()

    _add_heading(doc, "Executive Summary", 1)
    _add_note_panel(doc, "📝 Notes & Observations — Server Review")

    _kpi_cards(
        doc,
        cards=[
            ("Tasks", _fmt_int(tasks_count), "E8F5E9"),
            ("Endpoints", _fmt_int(endpoints_count), "E3F2FD"),
            ("Sources", _fmt_int(src_n), "FFF3E0"),
            ("Targets", _fmt_int(tgt_n), "F3E5F5"),
        ],
    )

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    filename = f"Replicate_Server_Review_{customer_name}_{server_name}.docx".replace(" ", "_")
    return buf.read(), filename

# ============================================================
# Customer Technical Overview
# ============================================================





async def _endpoint_mix_from_repo(conn, customer_id: int):
    """Compute Endpoint Mix using the SQL that correctly identifies Log Stream sources.
    Returns (source_mix, target_mix) as lists of {"type": str, "uses": int}.
    """
    import logging, os
    DBG = os.getenv("REPMETA_DEBUG_LOGSTREAM") in ("1", "true", "TRUE", "yes")
    log = logging.getLogger(__name__ + ".endpoint_mix_repo")
    
    def dbg(msg, *args):
        if DBG:
            try:
                log.info("[LMIX] " + str(msg), *args)
            except Exception:
                pass
    
    dbg("START customer_id=%s", customer_id)
    
    # Use the SQL that CORRECTLY identifies Log Stream sources
    # This SQL already handles the logstreamstagingtask detection properly
    sql = f"""
    -- Latest per server
    WITH latest AS (
      SELECT server_id, MAX(run_id) AS run_id
      FROM {SCHEMA}.ingest_run
      WHERE customer_id = %s
      GROUP BY server_id
    ),
    repo AS (
      SELECT d.endpoint_id, UPPER(d.role) AS role, 
             COALESCE(d.db_settings_type,'(unknown)') AS db_settings_type
      FROM {SCHEMA}.rep_database d 
      JOIN latest l ON l.run_id = d.run_id
    ),
    src_e AS (
      SELECT endpoint_id FROM repo WHERE role='SOURCE'
    ),
    -- Detect Log Stream sources via settings_json
    flagged AS (
      -- Check each source table for logstreamstagingtask
      SELECT s.endpoint_id
      FROM {SCHEMA}.rep_db_sqlserver_source s 
      JOIN src_e e USING(endpoint_id)
      WHERE CAST(s.settings_json AS text) ILIKE '%%logstreamstagingtask%%'
      
      UNION
      
      SELECT s.endpoint_id
      FROM {SCHEMA}.rep_db_db2_luw_source s 
      JOIN src_e e USING(endpoint_id)
      WHERE CAST(s.settings_json AS text) ILIKE '%%logstreamstagingtask%%'
      
      UNION
      
      SELECT s.endpoint_id
      FROM {SCHEMA}.rep_db_oracle_source s 
      JOIN src_e e USING(endpoint_id)
      WHERE CAST(s.settings_json AS text) ILIKE '%%logstreamstagingtask%%'
      
      UNION
      
      SELECT s.endpoint_id
      FROM {SCHEMA}.rep_db_mysql_source s 
      JOIN src_e e USING(endpoint_id)
      WHERE CAST(s.settings_json AS text) ILIKE '%%logstreamstagingtask%%'
      
      UNION
      
      SELECT s.endpoint_id
      FROM {SCHEMA}.rep_db_postgresql_source s 
      JOIN src_e e USING(endpoint_id)
      WHERE CAST(s.settings_json AS text) ILIKE '%%logstreamstagingtask%%'
      
      UNION
      
      SELECT s.endpoint_id
      FROM {SCHEMA}.rep_db_file_source s 
      JOIN src_e e USING(endpoint_id)
      WHERE CAST(s.settings_json AS text) ILIKE '%%logstreamstagingtask%%'
      
      UNION
      
      SELECT s.endpoint_id
      FROM {SCHEMA}.rep_db_odbc_source s 
      JOIN src_e e USING(endpoint_id)
      WHERE CAST(s.settings_json AS text) ILIKE '%%logstreamstagingtask%%'
      
      UNION
      
      SELECT s.endpoint_id
      FROM {SCHEMA}.rep_db_teradata_source s 
      JOIN src_e e USING(endpoint_id)
      WHERE CAST(s.settings_json AS text) ILIKE '%%logstreamstagingtask%%'
      
      UNION
      
      SELECT s.endpoint_id
      FROM {SCHEMA}.rep_db_vsam_source s 
      JOIN src_e e USING(endpoint_id)
      WHERE CAST(s.settings_json AS text) ILIKE '%%logstreamstagingtask%%'
      
      UNION
      
      SELECT s.endpoint_id
      FROM {SCHEMA}.rep_db_db2_iseries_source s 
      JOIN src_e e USING(endpoint_id)
      WHERE CAST(s.settings_json AS text) ILIKE '%%logstreamstagingtask%%'
      
      UNION
      
      SELECT s.endpoint_id
      FROM {SCHEMA}.rep_db_db2_zos_source s 
      JOIN src_e e USING(endpoint_id)
      WHERE CAST(s.settings_json AS text) ILIKE '%%logstreamstagingtask%%'
      
      UNION
      
      SELECT s.endpoint_id
      FROM {SCHEMA}.rep_db_ims_source s 
      JOIN src_e e USING(endpoint_id)
      WHERE CAST(s.settings_json AS text) ILIKE '%%logstreamstagingtask%%'
    )
    
    -- Get final counts: Log Stream for flagged sources, original type for others
    SELECT 
      'SOURCE' AS side, 
      'Log Stream' AS type, 
      COUNT(*) AS uses
    FROM flagged
    WHERE EXISTS (SELECT 1 FROM flagged)  -- Only include if there are any
    
    UNION ALL
    
    SELECT 
      'SOURCE' AS side, 
      db_settings_type AS type, 
      COUNT(*) AS uses
    FROM repo r
    WHERE role='SOURCE'
      AND NOT EXISTS (SELECT 1 FROM flagged f WHERE f.endpoint_id = r.endpoint_id)
    GROUP BY db_settings_type
    
    UNION ALL
    
    -- Target endpoints (no Log Stream carve-out)
    SELECT 
      'TARGET' AS side, 
      db_settings_type AS type, 
      COUNT(*) AS uses
    FROM repo
    WHERE role='TARGET'
    GROUP BY db_settings_type
    
    ORDER BY side, uses DESC, type;
    """
    
    try:
        rows = await _all(conn, sql, (customer_id,)) or []
        dbg("Query returned %s rows", len(rows))
    except Exception as e:
        log.error("Failed to execute endpoint mix query: %s", e)
        # If the query fails (maybe some tables don't exist), fall back to simple query
        sql_fallback = f"""
        WITH latest AS (
          SELECT server_id, MAX(run_id) AS run_id
          FROM {SCHEMA}.ingest_run
          WHERE customer_id = %s
          GROUP BY server_id
        )
        SELECT 
          UPPER(d.role) AS side,
          COALESCE(d.db_settings_type,'(unknown)') AS type,
          COUNT(*) AS uses
        FROM {SCHEMA}.rep_database d
        JOIN latest l ON l.run_id = d.run_id
        GROUP BY UPPER(d.role), d.db_settings_type
        ORDER BY side, uses DESC, type;
        """
        rows = await _all(conn, sql_fallback, (customer_id,)) or []
        dbg("Fallback query returned %s rows", len(rows))
    
    # Parse results into source and target mixes
    source_mix = []
    target_mix = []
    
    for row in rows:
        if isinstance(row, dict):
            side = str(row.get('side', '')).upper()
            type_name = str(row.get('type', '(unknown)'))
            uses = int(row.get('uses', 0))
        else:
            side = str(row[0] if len(row) > 0 else '').upper()
            type_name = str(row[1] if len(row) > 1 else '(unknown)')
            uses = int(row[2] if len(row) > 2 else 0)
        
        # Special handling for Log Stream - don't canonize it
        if type_name == 'Log Stream':
            display_type = 'Log Stream'
        else:
            # Canonize other type names for display
            display_type = canonize_to_master(type_name, is_source=(side=='SOURCE'))
        
        if side == 'SOURCE':
            source_mix.append({'type': display_type, 'uses': uses})
            dbg("Added SOURCE: %s = %s", display_type, uses)
        elif side == 'TARGET':
            target_mix.append({'type': display_type, 'uses': uses})
            dbg("Added TARGET: %s = %s", display_type, uses)
    
    # Sort by uses descending, then by type
    source_mix.sort(key=lambda x: (-x.get('uses', 0), x.get('type', '')))
    target_mix.sort(key=lambda x: (-x.get('uses', 0), x.get('type', '')))
    
    dbg("FINAL source_mix=%s", source_mix)
    dbg("FINAL target_mix=%s", target_mix)
    
    return source_mix, target_mix

def _endpoint_mix_cards(doc: Document,
                        src_rows: List[Dict[str, Any]],
                        tgt_rows: List[Dict[str, Any]],
                        title: str,
                        subtitle: Optional[str] = None):
    # Normalize into (label, count)
    src_items = [( _pretty_type(r.get("type")), int(r.get("uses", 0)) ) for r in src_rows]
    tgt_items = [( _pretty_type(r.get("type")), int(r.get("uses", 0)) ) for r in tgt_rows]
    src_items.sort(key=lambda x: (-x[1], x[0]))
    tgt_items.sort(key=lambda x: (-x[1], x[0]))
    TOP = 10
    src_items = src_items[:TOP]
    tgt_items = tgt_items[:TOP]
    max_len = max(len(src_items), len(tgt_items)) or 1

    _add_text(doc, title, size=12, bold=True)
    if subtitle:
        _add_text(doc, subtitle, size=9, italic=True)
    table = doc.add_table(rows=max_len + 1, cols=2)
    table.style = "Light Shading Accent 1"
    table.rows[0].cells[0].text = "Sources"
    table.rows[0].cells[1].text = "Targets"
    _cell_bold(table.rows[0].cells[0], 10)
    _cell_bold(table.rows[0].cells[1], 10)

    for i in range(max_len):
        if i < len(src_items):
            s_label, s_n = src_items[i]
            icon = _type_icon(s_label)
            table.rows[i + 1].cells[0].text = f"{(icon + '  ') if icon else ''}{s_label} - {_fmt_int(s_n)}"
        else:
            table.rows[i + 1].cells[0].text = ""
        if i < len(tgt_items):
            t_label, t_n = tgt_items[i]
            icon = _type_icon(t_label)
            table.rows[i + 1].cells[1].text = f"{(icon + '  ') if icon else ''}{t_label} - {_fmt_int(t_n)}"
        else:
            table.rows[i + 1].cells[1].text = ""

async def generate_customer_report_docx(customer_name: str, include_license: bool = True) -> Tuple[bytes, str]:
    async with connection() as conn:
        await _set_row_factory(conn)
        await _load_master_and_alias_from_db(conn)  # ensure masters/aliases ready

        c_row = await _one(
            conn, f"SELECT customer_id FROM {SCHEMA}.dim_customer WHERE customer_name=%s", (customer_name,)
        )
        if not c_row:
            raise ValueError(f"Customer '{customer_name}' not found. Add the customer and ingest data first.")
        customer_id = c_row["customer_id"]

        latest_train = await _ensure_latest_cache(conn)  # may be None

        # ---------- Servers ----------
        servers = await _all(
            conn,
            f"""
            SELECT server_id, server_name, COALESCE(environment,'') AS environment
            FROM {SCHEMA}.dim_server
            WHERE customer_id=%s
            ORDER BY (CASE WHEN LOWER(COALESCE(environment,'')) IN ('prod','production') THEN 0 ELSE 1 END),
                     server_name
            """,
            (customer_id,),
        )
        if not servers:
            raise ValueError(f"No servers found for customer '{customer_name}'. Ingest repository JSONs first.")

        # ---------- Latest Replicate version per server ----------
        version_rows = await _try_all(
            conn,
            f"SELECT server_name, replicate_version, last_repo FROM {SCHEMA}.v_customer_latest_runs WHERE customer_id=%s",
            (customer_id,),
        )
        if version_rows is None:
            version_rows = await _all(
                conn,
                f"""
                WITH latest AS (
                  SELECT server_id, MAX(created_at) AS last_ingest
                  FROM {SCHEMA}.ingest_run
                  WHERE customer_id=%s
                  GROUP BY server_id
                )
                SELECT s.server_name, r.replicate_version, r.created_at AS last_repo
                FROM latest l
                JOIN {SCHEMA}.ingest_run r
                  ON r.server_id=l.server_id AND r.created_at=l.last_ingest
                JOIN {SCHEMA}.dim_server s ON s.server_id=l.server_id
                ORDER BY s.server_name
                """,
                (customer_id,),
            )
        version_map: Dict[str, Tuple[Optional[str], Optional[datetime]]] = {
            r["server_name"]: (r.get("replicate_version"), r.get("last_repo")) for r in version_rows
        }

        # Posture map
        posture_map: Dict[str, Tuple[str, int, str]] = {}
        for sname, (ver_str, _last_repo_dt) in (version_map or {}).items():
            t = _parse_replicate_version_to_train(ver_str)
            if t and latest_train:
                delta = _trains_behind(latest_train, t)
                posture_map[sname] = (ver_str or "-", delta, _posture_label(delta))
            else:
                posture_map[sname] = (ver_str or "-", 999, _posture_label(999))

        # ---------- Repo totals ----------
        latest_repo_rows = await _all(
            conn,
            f"""
            WITH latest AS (
              SELECT server_id, MAX(created_at) AS last_ingest
              FROM {SCHEMA}.ingest_run
              WHERE customer_id=%s
              GROUP BY server_id
            ),
            runs AS (
              SELECT r.server_id, r.run_id
              FROM latest l
              JOIN {SCHEMA}.ingest_run r
                ON r.server_id=l.server_id AND r.created_at=l.last_ingest
            )
            SELECT
              (SELECT COUNT(*) FROM {SCHEMA}.rep_task     t JOIN runs u ON t.run_id=u.run_id) AS tasks,
              (SELECT COUNT(*) FROM {SCHEMA}.rep_database d JOIN runs u ON d.run_id=u.run_id) AS endpoints,
              (SELECT COUNT(*) FROM {SCHEMA}.rep_database d JOIN runs u ON d.run_id=u.run_id WHERE d.role='SOURCE') AS src,
              (SELECT COUNT(*) FROM {SCHEMA}.rep_database d JOIN runs u ON d.run_id=u.run_id WHERE d.role='TARGET') AS tgt
            """,
            (customer_id,),
        )
        repo_totals = latest_repo_rows[0] if latest_repo_rows else {"tasks": 0, "endpoints": 0, "src": 0, "tgt": 0}

        # ---------- Endpoint mix (prefer Repository JSON; fallback to QEM/TSV if empty) ----------
        src_rows_used, tgt_rows_used = await _endpoint_mix_from_repo(conn, customer_id)
        used_qem_view = False
        from_tsv = False
        if (not src_rows_used) and (not tgt_rows_used):
            mix_src = await _try_all(conn, f"""
            SELECT type, uses FROM {SCHEMA}.v_qem_endpoint_mix
            WHERE customer_id=%s AND role='SOURCE'
            ORDER BY uses DESC, type
            """, (customer_id,))
            mix_tgt = await _try_all(conn, f"""
            SELECT type, uses FROM {SCHEMA}.v_qem_endpoint_mix
            WHERE customer_id=%s AND role='TARGET'
            ORDER BY uses DESC, type
            """, (customer_id,))
            if mix_src is not None and mix_tgt is not None and (mix_src or mix_tgt):
                src_rows_used, tgt_rows_used = mix_src, mix_tgt
                used_qem_view = True
                from_tsv = True
            else:
                src_types_tsv = await _all(conn, f"""
                SELECT source_type AS type, COUNT(*) AS uses
                FROM {SCHEMA}.qem_task_perf
                WHERE customer_id=%s AND source_type IS NOT NULL
                GROUP BY source_type
                ORDER BY uses DESC, type
                """, (customer_id,))
                tgt_types_tsv = await _all(conn, f"""
                SELECT target_type AS type, COUNT(*) AS uses
                FROM {SCHEMA}.qem_task_perf
                WHERE customer_id=%s AND target_type IS NOT NULL
                GROUP BY target_type
                ORDER BY uses DESC, type
                """, (customer_id,))
                src_rows_used = src_types_tsv or []
                tgt_rows_used = tgt_types_tsv or []
                from_tsv = bool(src_rows_used or tgt_rows_used)
                used_qem_view = False

# ---------- Peak volumes (TSV) ----------
        peak_fl = await _one(
            conn,
            f"""
            SELECT s.server_name, q.task_name, q.fl_total_records
            FROM {SCHEMA}.qem_task_perf q
            JOIN {SCHEMA}.dim_server s USING (server_id)
            WHERE q.customer_id=%s AND q.fl_total_records IS NOT NULL
            ORDER BY q.fl_total_records DESC NULLS LAST
            LIMIT 1
            """,
            (customer_id,),
        )
        peak_cdc = await _one(
            conn,
            f"""
            SELECT s.server_name, q.task_name, q.cdc_commit_change_records
            FROM {SCHEMA}.qem_task_perf q
            JOIN {SCHEMA}.dim_server s USING (server_id)
            WHERE q.customer_id=%s AND q.cdc_commit_change_records IS NOT NULL
            ORDER BY q.cdc_commit_change_records DESC NULLS LAST
            LIMIT 1
            """,
            (customer_id,),
        )

        # ---------- Server rollup (prefer view) ----------
        rollup_rows = await _try_all(
            conn,
            f"SELECT * FROM {SCHEMA}.v_server_rollup WHERE customer_id=%s ORDER BY server_name",
            (customer_id,),
        )
        if rollup_rows is None:
            rollup_rows = await _all(
                conn,
                f"""
                WITH base AS (
                  SELECT server_id,
                         COUNT(DISTINCT task_name)   AS tasks,
                         COUNT(DISTINCT source_name) AS src_eps,
                         COUNT(DISTINCT target_name) AS tgt_eps
                  FROM {SCHEMA}.qem_task_perf
                  WHERE customer_id=%s
                  GROUP BY server_id
                ),
                last_repo AS (
                  SELECT server_id, MAX(created_at) AS last_repo
                  FROM {SCHEMA}.ingest_run
                  WHERE customer_id=%s
                  GROUP BY server_id
                ),
                last_qem AS (
                  SELECT server_id, MAX(created_at) AS last_qem
                  FROM {SCHEMA}.qem_ingest_run
                  WHERE customer_id=%s
                  GROUP BY server_id
                )
                SELECT s.server_name,
                       COALESCE(b.tasks,0)   AS tasks,
                       COALESCE(b.src_eps,0) AS src_eps,
                       COALESCE(b.tgt_eps,0) AS tgt_eps,
                       lr.last_repo,
                       lq.last_qem
                FROM {SCHEMA}.dim_server s
                LEFT JOIN base      b  USING (server_id)
                LEFT JOIN last_repo lr USING (server_id)
                LEFT JOIN last_qem  lq USING (server_id)
                WHERE s.customer_id=%s
                ORDER BY s.server_name
                """,
                (customer_id, customer_id, customer_id, customer_id),
            )

        # ---------- Primary pair per server (prefer view) ----------
        primary_pairs = await _try_all(
            conn,
            f"SELECT server_name, source_type, target_type, n FROM {SCHEMA}.v_primary_pairs WHERE customer_id=%s ORDER BY server_name",
            (customer_id,),
        )
        if primary_pairs is None:
            primary_pairs = await _all(
                conn,
                f"""
                WITH pairs AS (
                  SELECT server_id, source_type, target_type, COUNT(*) AS n
                  FROM {SCHEMA}.qem_task_perf
                  WHERE customer_id=%s
                  GROUP BY server_id, source_type, target_type
                ),
                ranked AS (
                  SELECT p.*, ROW_NUMBER() OVER (PARTITION BY p.server_id ORDER BY n DESC NULLS LAST, source_type, target_type) AS rn
                  FROM pairs p
                )
                SELECT s.server_name, r.source_type, r.target_type, r.n
                FROM ranked r
                JOIN {SCHEMA}.dim_server s USING (server_id)
                WHERE rn=1
                ORDER BY s.server_name
                """,
                (customer_id,),
            )
        primary_map = {r["server_name"]: (r.get("source_type"), r.get("target_type"), r.get("n")) for r in primary_pairs}

        # ---------- Coverage matrix (prefer view) ----------
        coverage = await _try_all(
            conn,
            f"SELECT s_type, t_type, n FROM {SCHEMA}.v_coverage_matrix WHERE customer_id=%s",
            (customer_id,),
        )
        if coverage is None:
            coverage = await _all(
                conn,
                f"""
                SELECT COALESCE(source_type,'(unknown)') AS s_type,
                       COALESCE(target_type,'(unknown)') AS t_type,
                       COUNT(*) AS n
                FROM {SCHEMA}.qem_task_perf
                WHERE customer_id=%s
                GROUP BY COALESCE(source_type,'(unknown)'), COALESCE(target_type,'(unknown)')
                """,
                (customer_id,),
            )
        src_types = sorted({_pretty_type(r["s_type"]) for r in coverage})
        tgt_types = sorted({_pretty_type(r["t_type"]) for r in coverage})
        cov_map: Dict[Tuple[str, str], int] = {}
        for r in coverage:
            cov_map[(_pretty_type(r["s_type"]), _pretty_type(r["t_type"]))] = int(r["n"])

        # Build customer-level edges from coverage
        customer_edges = _flow_edges_from_coverage_rows(coverage)

        # Build per-server edges map
        server_edges_map = await _gather_server_edges_map(conn, customer_id)

        # ---------- License coverage (use repmeta.v_license_vs_usage) ----------
        # It exposes: customer_id, ef_role ('SOURCE'/'TARGET'), family_name, is_licensed (bool),
        # and configured_count (endpoints configured in rep_database mapped to that family).
        lic_rows = await _all(
            conn,
            f"""
            SELECT ef_role, family_name, is_licensed, COALESCE(configured_count,0) AS configured_count
            FROM {SCHEMA}.v_license_vs_usage
            WHERE customer_id=%s
            """ ,
            (customer_id,),
        )

        # Always compute MIX/REPO-based 'used' as a fallback and merge later
        mix_used_src = _names_from_rows(src_rows_used, is_source=True)
        mix_used_tgt = _names_from_rows(tgt_rows_used, is_source=False)
        used_source_types, used_target_types = set(), set()
        lic_src, lic_tgt = set(), set()

        for r in (lic_rows or []):
            fam_raw = str(r.get("family_name") or "").strip()
            role = str(r.get("ef_role") or "").strip().upper()
            if not fam_raw or not role:
                continue
            fam = canonize_to_master(fam_raw, is_source=(role=="SOURCE"))
            # Collect licensed families
            is_lic = r.get("is_licensed")
            try:
                is_lic_bool = bool(is_lic) if isinstance(is_lic, (bool, int)) else str(is_lic).strip().lower() in {"true","t","yes","y","1"}
            except Exception:
                is_lic_bool = False
            if is_lic_bool:
                if role == "SOURCE":
                    lic_src.add(fam)
                elif role == "TARGET":
                    lic_tgt.add(fam)
            # Collect used from configured_count
            try:
                cfg_ct = int(r.get("configured_count") or 0)
            except Exception:
                cfg_ct = 0
            if cfg_ct > 0:
                if role == "SOURCE":
                    used_source_types.add(fam)
                elif role == "TARGET":
                    used_target_types.add(fam)

        # Merge MIX-based 'used' if LVU didn't report any
        if not used_source_types:
            used_source_types = set(mix_used_src)
        else:
            used_source_types |= set(mix_used_src)
        if not used_target_types:
            used_target_types = set(mix_used_tgt)
        else:
            used_target_types |= set(mix_used_tgt)

        # Determine whether the license effectively covers ALL families.
        lic_all_src = lic_src == set(MASTER_SOURCE_ENDPOINTS)
        lic_all_tgt = lic_tgt == set(MASTER_TARGET_ENDPOINTS)

        # Fallbacks when LVU data is missing or empty (avoid 0/0 & blank panels)
        if not lic_rows or (not lic_src and not lic_tgt):
            lic_row = await _one(
                conn,
                f"SELECT * FROM {SCHEMA}.v_latest_customer_license WHERE customer_id=%s",
                (customer_id,),
            )
            if lic_row:
                lic_all_src = bool(lic_row.get("licensed_all_sources"))
                lic_all_tgt = bool(lic_row.get("licensed_all_targets"))
                lic_src = set(canonize_to_master(s, True) for s in (lic_row.get("licensed_sources") or []))
                lic_tgt = set(canonize_to_master(t, False) for t in (lic_row.get("licensed_targets") or []))
            # If still empty, make the 'licensed universe' equal to what we can see in use
            if not lic_src and used_source_types:
                lic_src = set(used_source_types)
                lic_all_src = False
            if not lic_tgt and used_target_types:
                lic_tgt = set(used_target_types)
                lic_all_tgt = False

        # ---------- NEW: Top-5 tasks by #tables per server (precompute once) ----------
        top_tables_by_server: Dict[str, List[Tuple[str, int]]] = defaultdict(list)
        top_tables_rows = await _try_all(
            conn,
            f"""
            WITH latest AS (
              SELECT server_id, MAX(created_at) AS last_ingest
              FROM {SCHEMA}.ingest_run
              WHERE customer_id=%s
              GROUP BY server_id
            ),
            counts AS (
              SELECT r.server_id, t.task_name, COUNT(*) AS n_tables
              FROM {SCHEMA}.rep_task_table tt
              JOIN {SCHEMA}.rep_task t ON t.task_id = tt.task_id AND t.run_id = tt.run_id
              JOIN {SCHEMA}.ingest_run r ON r.run_id = tt.run_id
              JOIN latest l
                ON l.server_id = r.server_id AND r.created_at = l.last_ingest
              GROUP BY r.server_id, t.task_name
            ),
            ranked AS (
              SELECT server_id, task_name, n_tables,
                     ROW_NUMBER() OVER (PARTITION BY server_id ORDER BY n_tables DESC, task_name) AS rn
              FROM counts
            )
            SELECT s.server_name, task_name, n_tables
            FROM ranked rk
            JOIN {SCHEMA}.dim_server s ON s.server_id = rk.server_id
            WHERE rk.rn <= 5 AND s.customer_id = %s
            ORDER BY s.server_name, n_tables DESC, task_name
            """,
            (customer_id, customer_id),
        )
        if top_tables_rows:
            for r in top_tables_rows:
                try:
                    top_tables_by_server[r["server_name"]].append((r["task_name"], int(r["n_tables"])))
                except Exception:
                    pass

    # ---------------- DOCX BUILD ----------------
        doc = Document()
        try:
            _apply_global_styles(doc)
            _apply_footer(doc)
        except Exception:
            pass
        try:
            _apply_branding(doc)
        except Exception:
            pass

    _add_title(doc, "Customer Technical Overview")
    _add_text(doc, f"Customer: {customer_name}", size=10)
    doc.add_paragraph()
    _add_toc(doc)
    doc.add_page_break()

    # 1) Executive Summary
    _add_heading(doc, "1. Executive Summary", 1)

    oos_count = sum(1 for _, (_, delta, _) in posture_map.items() if delta > 2)
    cards = [
        ("Servers", _fmt_int(len(servers)), "E3F2FD"),
        ("Tasks", _fmt_int(repo_totals.get("tasks", 0)), "E8F5E9"),
        ("Endpoints", _fmt_int(repo_totals.get("endpoints", 0)), "F3E5F5"),
        ("Sources", _fmt_int(repo_totals.get("src", 0)), "FFF3E0"),
        ("Targets", _fmt_int(repo_totals.get("tgt", 0)), "FFF3E0"),
        ("Servers OOS", _fmt_int(oos_count), "FFE9E6"),
    ]
    _kpi_cards(doc, cards)
    _add_text(doc, "Totals reflect the inventory discovered from Repository JSON ingests.", size=9, italic=True)
    doc.add_paragraph()

    # Versions table with posture
    _add_text(doc, "Replicate Versions by Server", size=12, bold=True)
    headers = ["Server", "Replicate Version", "Last Repo Ingest", "Posture"]
    rows = []
    for s in sorted(version_map.keys()):
        ver, last_repo = version_map.get(s, (None, None))
        _, delta, label = posture_map.get(s, (ver or "-", 999, _posture_label(999)))
        rows.append((s, _version_badge(ver), _age_badge(last_repo), label))
    _add_table(doc, headers=headers, rows=rows, style="Light Shading Accent 3")
    doc.add_paragraph()

    # Endpoint Mix
    section_title = "Active Endpoint Mix (from QEM)" if used_qem_view else (f"Endpoint Mix{' (from TSV)' if from_tsv else ' (from repo)'}")
    section_subtitle = ("Endpoints tied to active tasks in the latest QEM runs." if used_qem_view else None)
    _endpoint_mix_cards(doc, src_rows_used, tgt_rows_used, section_title, section_subtitle)
    doc.add_paragraph()

    # License Usage - modern design (no "unlicensed in use")
    if include_license:
        _license_usage_section(
        doc,
        used_source_types, used_target_types,
        lic_all_src, lic_all_tgt,
        lic_src, lic_tgt,
    )
        # Flow Table: Customer-level (forced table)
    _add_flow_pair_table(doc, customer_edges, "Source → Target (Task Counts)")
    
    # ---- MetricsLog rollups (added) ----
    # Open a FRESH connection for metrics rollups to avoid "connection is closed"
    # if the earlier read-only block has been exited. We only read here.
    try:
        async with connection() as conn_metrics:
            await _set_row_factory(conn_metrics)

            try:
                monthly = await _metrics_monthly_current_year(conn_metrics, customer_id)
                _add_metrics_section(doc, f"Monthly Volume (Last 6 Months)",
                                     monthly, headers=["Month","Load","CDC","Total"])
            except Exception as e:
                _add_text(doc, f"⚠ MetricsLog monthly summary failed: {type(e).__name__}: {e}", size=10, italic=True)

            try:
                yearly  = await _metrics_yearly_last5(conn_metrics, customer_id)
                _add_metrics_section(doc, "Annual Volume (Last 5 Years)",
                                     yearly, headers=["Year","Load","CDC","Total"])
            except Exception as e:
                _add_text(doc, f"⚠ MetricsLog annual summary failed: {type(e).__name__}: {e}", size=10, italic=True)

            try:
                top_tasks = await _metrics_top_tasks(conn_metrics, customer_id, limit=5)
                _add_metrics_section(doc, "Top 5 Tasks by Volume (Load + CDC)",
                                     top_tasks, headers=["Task","Server","Load","CDC","Total"])
            except Exception as e:
                _add_text(doc, f"⚠ MetricsLog top tasks failed: {type(e).__name__}: {e}", size=10, italic=True)

            try:
                top_src_load = await _metrics_top_endpoints(conn_metrics, customer_id, role="SOURCE", metric="load", limit=5)
                top_src_cdc  = await _metrics_top_endpoints(conn_metrics, customer_id, role="SOURCE", metric="cdc",  limit=5)
                top_tgt_load = await _metrics_top_endpoints(conn_metrics, customer_id, role="TARGET", metric="load", limit=5)
                top_tgt_cdc  = await _metrics_top_endpoints(conn_metrics, customer_id, role="TARGET", metric="cdc",  limit=5)
                _add_metrics_section(doc, "Top 5 Sources by Load Volume",  top_src_load, headers=["Source","Load"])
                _add_metrics_section(doc, "Top 5 Sources by CDC Volume",   top_src_cdc,  headers=["Source","CDC"])
                _add_metrics_section(doc, "Top 5 Targets by Load Volume",  top_tgt_load, headers=["Target","Load"])
                _add_metrics_section(doc, "Top 5 Targets by CDC Volume",   top_tgt_cdc,  headers=["Target","CDC"])
            except Exception as e:
                _add_text(doc, f"⚠ MetricsLog endpoint leaders failed: {type(e).__name__}: {e}", size=10, italic=True)

    except Exception as outer_e:
        _add_text(doc, f"⚠ MetricsLog summary failed: {type(outer_e).__name__}: {outer_e}", size=10, italic=True)
    _add_note_panel(doc, "📝 Notes & Observations — 1. Executive Summary")


    doc.add_page_break()

    # 2) Customer Insights
    _add_heading(doc, "2. Customer Insights", 1)

    _add_text(doc, "Version posture vs latest GA train", size=11, bold=True)
    if not latest_train:
        _add_text(doc, "⚠ Latest GA train not available (no cache & GitHub fetch failed).", size=10, italic=True)
    else:
        behind2 = []
        beyond2 = []
        for sname, (ver_str, delta, label) in posture_map.items():
            if delta == 2:
                behind2.append((sname, ver_str, label))
            elif delta > 2:
                beyond2.append((sname, ver_str, label))
        headers = ["Server", "Replicate", "Posture"]
        if not behind2 and not beyond2:
            _add_text(doc, "All servers are within 0-1 train of latest GA. ✅", size=10, italic=True)
        else:
            if behind2:
                _add_text(doc, "🟠 2 trains behind", size=10, bold=True)
                _add_table(doc, headers, [(s, _version_badge(v), l) for s, v, l in sorted(behind2)], "Light Shading Accent 2")
            if beyond2:
                _add_text(doc, "🔴 >2 trains behind (Out-of-Support)", size=10, bold=True)
                _add_table(doc, headers, [(s, _version_badge(v), l) for s, v, l in sorted(beyond2)], "Light Shading Accent 2")

    doc.add_paragraph()

    # Insight #1: tasks with null target
    try:
        async with connection() as conn_ro:
            await _set_row_factory(conn_ro)
            tasks_null_tgt = await _all(
                conn_ro,
                f"""
                WITH latest_qem AS (
                  SELECT r.server_id, MAX(r.created_at) AS max_created
                  FROM {SCHEMA}.qem_ingest_run r
                  WHERE r.customer_id=%s
                  GROUP BY r.server_id
                ),
                rows AS (
                  SELECT q.*
                  FROM {SCHEMA}.qem_task_perf q
                  JOIN {SCHEMA}.qem_ingest_run r ON r.qem_run_id = q.qem_run_id
                  JOIN latest_qem l
                    ON l.server_id = r.server_id AND r.created_at = l.max_created
                  WHERE q.customer_id=%s
                )
                SELECT s.server_name,
                       q.task_name,
                       COALESCE(NULLIF(q.target_name,''), '(unknown)') AS target_name
                FROM rows q
                JOIN {SCHEMA}.dim_server s USING (server_id)
                WHERE (LOWER(BTRIM(q.target_type)) = 'null target' OR q.target_type IS NULL)
                GROUP BY s.server_name, q.task_name, COALESCE(NULLIF(q.target_name,''), '(unknown)')
                ORDER BY s.server_name, q.task_name
                """,
                (customer_id, customer_id),
            )
        if tasks_null_tgt:
            _add_text(doc, "Tasks with Null Target Type", size=11, bold=True)
            headers = ["Server", "Task", "Target Endpoint"]
            rows = [(r.get("server_name") or "-", r.get("task_name") or "-", r.get("target_name") or "-")
                    for r in tasks_null_tgt]
            _add_table(doc, headers=headers, rows=rows, style="Light Shading")
        else:
            _add_text(doc, "No tasks with Null Target Type found in latest QEM snapshot.", size=10, italic=True)
    except Exception as e:
        _add_text(doc, f"⚠ Insight #1 failed: {type(e).__name__}: {e}", size=10, italic=True)

    doc.add_paragraph()

    # Insight #2: duplicate endpoints (identical settings)
    try:
        async with connection() as conn_ro2:
            await _set_row_factory(conn_ro2)
            dup_eps = await _all(
                conn_ro2,
                f"""
                WITH latest AS (
                  SELECT server_id, MAX(created_at) AS last_ingest
                  FROM {SCHEMA}.ingest_run
                  WHERE customer_id=%s
                  GROUP BY server_id
                ),
                runs AS (
                  SELECT r.server_id, r.run_id
                  FROM latest l
                  JOIN {SCHEMA}.ingest_run r
                    ON r.server_id = l.server_id AND r.created_at = l.last_ingest
                ),

                /* Union all detailed endpoint tables + generic fallback */
                unioned AS (
                  SELECT endpoint_id, settings_json FROM {SCHEMA}.rep_db_postgresql_source
                  UNION ALL SELECT endpoint_id, settings_json FROM {SCHEMA}.rep_db_postgresql_target
                  UNION ALL SELECT endpoint_id, settings_json FROM {SCHEMA}.rep_db_sqlserver_source
                  UNION ALL SELECT endpoint_id, settings_json FROM {SCHEMA}.rep_db_sqlserver_target
                  UNION ALL SELECT endpoint_id, settings_json FROM {SCHEMA}.rep_db_mysql_source
                  UNION ALL SELECT endpoint_id, settings_json FROM {SCHEMA}.rep_db_mysql_target
                  UNION ALL SELECT endpoint_id, settings_json FROM {SCHEMA}.rep_db_oracle_source
                  UNION ALL SELECT endpoint_id, settings_json FROM {SCHEMA}.rep_db_oracle_target
                  UNION ALL SELECT endpoint_id, settings_json FROM {SCHEMA}.rep_db_snowflake_target
                  UNION ALL SELECT endpoint_id, settings_json FROM {SCHEMA}.rep_db_redshift_target
                  UNION ALL SELECT endpoint_id, settings_json FROM {SCHEMA}.rep_db_s3_target
                  UNION ALL SELECT endpoint_id, settings_json FROM {SCHEMA}.rep_db_databricks_delta_target
                  UNION ALL SELECT endpoint_id, settings_json FROM {SCHEMA}.rep_db_databricks_cloud_storage_target
                  UNION ALL SELECT endpoint_id, settings_json FROM {SCHEMA}.rep_db_azure_adls_target
                  UNION ALL SELECT endpoint_id, settings_json FROM {SCHEMA}.rep_db_hdinsight_target
                  UNION ALL SELECT endpoint_id, settings_json FROM {SCHEMA}.rep_db_hadoop_target
                  UNION ALL SELECT endpoint_id, settings_json FROM {SCHEMA}.rep_db_kafka_target
                  UNION ALL SELECT endpoint_id, settings_json FROM {SCHEMA}.rep_db_amazon_msk_target
                  UNION ALL SELECT endpoint_id, settings_json FROM {SCHEMA}.rep_db_confluent_cloud_target
                  UNION ALL SELECT endpoint_id, settings_json FROM {SCHEMA}.rep_db_eventhubs_target
                  UNION ALL SELECT endpoint_id, settings_json FROM {SCHEMA}.rep_db_pubsub_target
                  UNION ALL SELECT endpoint_id, settings_json FROM {SCHEMA}.rep_db_logstream_target
                  UNION ALL SELECT endpoint_id, settings_json FROM {SCHEMA}.rep_db_file_source
                  UNION ALL SELECT endpoint_id, settings_json FROM {SCHEMA}.rep_db_file_target
                  UNION ALL SELECT endpoint_id, settings_json FROM {SCHEMA}.rep_db_file_channel_target
                  UNION ALL SELECT endpoint_id, settings_json FROM {SCHEMA}.rep_db_db2_luw_source
                  UNION ALL SELECT endpoint_id, settings_json FROM {SCHEMA}.rep_db_db2_zos_source
                  UNION ALL SELECT endpoint_id, settings_json FROM {SCHEMA}.rep_db_db2_zos_target
                  UNION ALL SELECT endpoint_id, settings_json FROM {SCHEMA}.rep_db_db2_iseries_source
                  UNION ALL SELECT endpoint_id, settings_json FROM {SCHEMA}.rep_db_teradata_source
                  UNION ALL SELECT endpoint_id, settings_json FROM {SCHEMA}.rep_db_teradata_target
                  UNION ALL SELECT endpoint_id, settings_json FROM {SCHEMA}.rep_db_odbc_source
                  UNION ALL SELECT endpoint_id, settings_json FROM {SCHEMA}.rep_db_odbc_target
                  UNION ALL SELECT endpoint_id, settings_json FROM {SCHEMA}.rep_db_informix_source
                  UNION ALL SELECT endpoint_id, settings_json FROM {SCHEMA}.rep_db_vsam_source
                  UNION ALL SELECT endpoint_id, settings_json FROM {SCHEMA}.rep_db_ims_source
                  UNION ALL SELECT endpoint_id, settings_json FROM {SCHEMA}.rep_db_gcs_target
                  UNION ALL SELECT endpoint_id, settings_json FROM {SCHEMA}.rep_db_ms_fabric_dw_target
                  UNION ALL SELECT endpoint_id, settings_json FROM {SCHEMA}.rep_db_settings_json
                ),
                base AS (
                  SELECT
                    s.server_name,
                    d.role,
                    d.db_settings_type,
                    d.name AS endpoint_name,
                    u.settings_json
                  FROM {SCHEMA}.rep_database d
                  JOIN runs r ON d.run_id = r.run_id
                  JOIN {SCHEMA}.dim_server s ON s.server_id = r.server_id
                  JOIN unioned u ON u.endpoint_id = d.endpoint_id
                ),
                clean AS (
                  SELECT
                    server_name,
                    role,
                    db_settings_type,
                    endpoint_name,
                    jsonb_strip_nulls(
                      (settings_json
                        - 'Name' - 'EndpointName' - 'DisplayName' - 'Description'
                        - 'Id' - 'ID' - 'Guid' - 'GUID'
                        - 'CreatedTime' - 'ModifiedTime' - 'LastTestConnection'
                        - 'Password' - 'Pwd' - 'Secret' - 'AccessKey' - 'SecretKey' - 'Token'
                        - 'ProxyPassword' - 'SaslPassword' - 'OAuthToken'
                        - 'privateKey' - 'privateKeyFile'
                      )
                    ) AS cfg
                  FROM base
                )
                SELECT
                  server_name,
                  role,
                  db_settings_type,
                  md5(cfg::text) AS cfg_sig,
                  ARRAY_AGG(endpoint_name ORDER BY endpoint_name) AS endpoint_names,
                  COUNT(*) AS n
                FROM clean
                GROUP BY server_name, role, db_settings_type, md5(cfg::text)
                HAVING COUNT(*) > 1
                ORDER BY server_name, role, n DESC
                """,
                (customer_id,),
            )
        _add_text(doc, "Endpoints with identical configuration", size=11, bold=True)
        if dup_eps:
            headers = ["Server", "Role", "Type", "Endpoints (identical settings)", "Count"]
            rows = []
            for r in dup_eps:
                rows.append((
                    r.get("server_name") or "-",
                    r.get("role") or "-",
                    _pretty_type(r.get("db_settings_type") or "") or "-",
                    ", ".join(r.get("endpoint_names") or []) or "-",
                    _fmt_int(r.get("n")),
                ))
            _add_table(doc, headers=headers, rows=rows, style="Light Shading Accent 2")
        else:
            _add_text(doc, "No duplicate endpoint configurations detected in latest runs.", size=10, italic=True)
    except Exception as e:
        _add_text(doc, f"⚠ Insight #2 failed: {type(e).__name__}: {e}", size=10, italic=True)

    
    # Tasks with DEBUG loggers (latest repo ingest per server)
    _add_text(doc, "Tasks with DEBUG Loggers (latest repo ingest per server)", size=11, bold=True)
    try:
        async with connection() as conn_ro_debug:
            await _set_row_factory(conn_ro_debug)
            rows = await _all(
                conn_ro_debug,
                f"""
                WITH latest AS (
                  SELECT server_id, MAX(created_at) AS last_ingest
                  FROM {SCHEMA}.ingest_run
                  WHERE customer_id=%s
                  GROUP BY server_id
                ),
                runs AS (
                  SELECT r.server_id, r.run_id
                  FROM latest l
                  JOIN {SCHEMA}.ingest_run r
                    ON r.server_id=l.server_id AND r.created_at=l.last_ingest
                )
                SELECT s.server_name,
                       t.task_name,
                       STRING_AGG(l.logger_name, ', ' ORDER BY l.logger_name) AS debug_loggers
                FROM {SCHEMA}.rep_task_logger l
                JOIN {SCHEMA}.rep_task t
                  ON t.task_id=l.task_id AND t.run_id=l.run_id
                JOIN runs r ON r.run_id=t.run_id
                JOIN {SCHEMA}.dim_server s ON s.server_id=t.server_id
                WHERE t.customer_id=%s AND UPPER(l.level)='DEBUG'
                GROUP BY s.server_name, t.task_name
                ORDER BY s.server_name, t.task_name
                """,
                (customer_id, customer_id),
            )
        if rows:
            _add_table(
                doc,
                headers=["Server", "Task", "Logger(s) at DEBUG"],
                rows=[(r.get("server_name","-"), r.get("task_name","-"), r.get("debug_loggers","-")) for r in rows],
                style="Light Shading",
            )
        else:
            _add_text(doc, "No tasks found with DEBUG loggers in the latest ingest.", size=10, italic=True)
    except Exception as e:
        _add_text(doc, f"⚠ DEBUG logger insight failed: {type(e).__name__}: {e}", size=10, italic=True)
    _add_note_panel(doc, "📝 Notes & Observations — 2. Customer Insights")


    doc.add_page_break()

    # 3) Environment & Inventory
    _add_heading(doc, "3. Environment & Inventory", 1)
    _add_text(doc, "Servers Overview", size=12, bold=True)
    headers = ["Server", "Tasks", "Source EPs", "Target EPs", "Replicate", "Posture"]
    rows = []
    for r in rollup_rows:
        sname = r.get("server_name")
        ver_str, delta, label = posture_map.get(sname, ("-", 999, _posture_label(999)))
        rows.append([
            sname,
            _fmt_int(r.get("tasks")),
            _fmt_int(r.get("src_eps")),
            _fmt_int(r.get("tgt_eps")),
            _version_badge(ver_str),
            label,
        ])
    _add_table(doc, headers=headers, rows=rows, style="Light Shading Accent 1")
    doc.add_paragraph()

    # Coverage Matrix
    _add_text(doc, "Source × Target Coverage (TSV)", size=12, bold=True)
    headers = ["Source \\ Target"] + [(f"{_type_icon(t)}  {t}" if _type_icon(t) else t) for t in tgt_types]
    rows = []
    for s_type in src_types:
        row = [(f"{_type_icon(s_type)}  {s_type}" if _type_icon(s_type) else s_type)]
        for t_type in tgt_types:
            row.append(_fmt_int(cov_map.get((s_type, t_type), 0)))
        rows.append(tuple(row))
    _add_table(doc, headers=headers, rows=rows, style="Light Shading Accent 1")
    _add_note_panel(doc, "📝 Notes & Observations — 3. Environment & Inventory")

    doc.add_page_break()

    # 4) Server Deep Dives
    _add_heading(doc, "4. Server Deep Dives", 1)
    servers_by_name = {srv["server_name"]: srv for srv in servers}
    for s in [srv["server_name"] for srv in servers]:
        # Diagnostics: log server mapping used for T90
        _srv_row = servers_by_name.get(s)
        srv_id = _srv_row.get("server_id") if _srv_row else None
        try:
            log.info("Deep-dive mapping: server_name=%s -> server_id=%s (customer_id=%s)", s, srv_id, customer_id)
        except Exception:
            pass
        _add_heading(doc, s, 2)
        if _srv_row and "server_id" in _srv_row:
            try:
                await render_metricslog_90d_sections(doc, conn, customer_id, _srv_row["server_id"], s)
            except Exception as _e:
                log.exception("T90 MetricsLog section failed for %s: %s", s, _e)
                _add_text(doc, f"⚠ T90 MetricsLog section failed for {s}: {_e}", size=9, italic=True)
        pair = primary_map.get(s)
        if pair:
            _add_text(doc, f"Primary pair: {_pretty_type(pair[0])} → {_pretty_type(pair[1])} ({_fmt_int(pair[2])} tasks)",
                      size=10, italic=True)

        # NEW: Top-5 tasks by number of tables for this server
        top_rows = top_tables_by_server.get(s) or []
        if top_rows:
            last_repo_dt = (version_map.get(s) or (None, None))[1]
            when_txt = ""
            try:
                if last_repo_dt:
                    when_txt = f" (latest repo ingest: {last_repo_dt.date()})"
            except Exception:
                pass
            _add_text(doc, f"Top-5 tasks by number of tables{when_txt}", size=11, bold=True)
            _add_table(
                doc,
                headers=["Task", "# Tables"],
                rows=[(name, _fmt_int(n)) for (name, n) in top_rows],
                style="Light Shading Accent 1",
            )
        else:
            _add_text(doc, "Top-5 tasks by number of tables — no table data found for latest ingest.", size=10, italic=True)

                # Per-server flow table (forced table)
        edges = server_edges_map.get(s, [])
        _add_flow_pair_table(doc, edges, "Source → Target (Task Counts)")

        doc.add_paragraph()

    # Index
    _add_note_panel(doc, "📝 Notes & Observations — 4. Server Deep Dives")

    _add_heading(doc, "Index", 1)
    p = doc.add_paragraph("Update the index in Word: References → Update Table.")
    p.runs[0].italic = True


    # === Latest Release Fixes (from DB) ===
    try:
        latest_label, groups = await _load_and_group_latest_release_issues(conn)
        doc.add_page_break()
        _render_latest_release_fixes_section(doc, latest_label, groups)
        _add_note_panel(doc, "📝 Notes & Observations — Latest Release Fixes")
    except Exception as e:
        log.exception("Latest-release fixes section failed: %s: %s", type(e).__name__, e)
        _add_text(doc, f"\u26A0 Latest-release fixes section failed: {type(e).__name__}: {e}", size=10, italic=True)
    # === end Latest Release Fixes ===

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    filename = f"Customer_Technical_Overview_{customer_name}.docx".replace(" ", "_")
    return buf.read(), filename
# ========= METRICSLOG T90 ENHANCEMENTS (Dynamic per-server window) ==================
from dataclasses import dataclass
from typing import Any, Dict, List

from dataclasses import dataclass
from typing import Dict, List, Any

@dataclass
class _T90TaskHealth:
    tkey: str
    uptime_pct: float
    downtime_hours: float
    restarts_per_day: float
    error_stop_rate: float
    median_session_minutes: float
    throughput_rps: float
    rows_moved: int
    restarts_total: int | None = None
    window_start: object | None = None
    window_end: object | None = None

def _fmt_pct_t90(x: float) -> str:
    try:
        return f"{float(x):.2f}%"
    except Exception:
        return str(x)

# === T90 helpers (Endpoint Mix - Avg Uptime) ================================

def _uptime_avg_numeric(row):
    """
    Return a 0..100 *average* uptime from a v_endpoint_perf_t90 row.
    If the upstream view summed per-task percentages, normalize by task count.
    """
    try:
        get = row.get if hasattr(row, "get") else (lambda k: row[k])
        u = float(get("uptime_pct") or 0.0)
        tasks = float(get("tasks") or 0.0)
    except Exception:
        u, tasks = 0.0, 0.0

    # If someone accidentally aggregated by summing percentages, fix it.
    if u > 100.0 and tasks > 0:
        u = u / tasks

    # Clamp to [0, 100] for display
    if u < 0.0:
        u = 0.0
    if u > 100.0:
        u = 100.0
    return u


def _uptime_bar(pct: float, width: int = 12) -> str:
    """Tiny text progress bar (Word-safe) to make uptime scannable at a glance."""
    try:
        pct = max(0.0, min(float(pct), 100.0))
    except Exception:
        pct = 0.0
    filled = int(round((pct / 100.0) * width))
    if filled < 0: filled = 0
    if filled > width: filled = width
    return "█" * filled + "░" * (width - filled)


def _uptime_avg_display(row) -> str:
    """Human-friendly Avg Uptime cell text, e.g. '78.25%  ███████░░░░'."""
    pct = _uptime_avg_numeric(row)
    return f"{pct:.2f}%  {_uptime_bar(pct)}"
    # If you prefer without the bar, use:
    # return f"{pct:.2f}%"


def _fmt_float_t90(x: float) -> str:
    try:
        return f"{float(x):.2f}"
    except Exception:
        return str(x)

def _fmt_int_t90(x) -> str:
    try:
        return f"{int(x):,}"
    except Exception:
        return str(x)

async def _t90_fetch_task_health(conn, customer_id: int, server_id: int) -> List[_T90TaskHealth]:
    rows = await _try_all(conn, f"""
        SELECT
            t.tkey,
            t.uptime_pct,
            (t.downtime_sec/3600.0) AS downtime_hours,
            t.restarts_per_day,
            COALESCE(t.error_stop_rate,0.0) AS error_stop_rate,
            COALESCE(t.median_session_minutes,0.0) AS median_session_minutes,
            COALESCE(t.throughput_rps,0.0) AS throughput_rps,
            COALESCE(t.rows_moved,0) AS rows_moved,
            COALESCE(t.session_count,0) AS restarts_total,
            w.window_start,
            w.window_end
        FROM {SCHEMA}.v_task_health_t90 t
        JOIN {SCHEMA}.v_metrics_t90_window w USING (customer_id, server_id)
        WHERE t.customer_id=%s AND t.server_id=%s
    """, (customer_id, server_id)) or []
    out: List[_T90TaskHealth] = []
    for r in rows:
        out.append(_T90TaskHealth(
            tkey=str(r.get("tkey","")),
            uptime_pct=float(r.get("uptime_pct") or 0),
            downtime_hours=float(r.get("downtime_hours") or 0),
            restarts_per_day=float(r.get("restarts_per_day") or 0),
            error_stop_rate=float(r.get("error_stop_rate") or 0),
            median_session_minutes=float(r.get("median_session_minutes") or 0),
            throughput_rps=float(r.get("throughput_rps") or 0),
            rows_moved=int(r.get("rows_moved") or 0),
            restarts_total=int(r.get("restarts_total") or 0),
            window_start=r.get("window_start"),
            window_end=r.get("window_end"),
        ))
    return out

async def _diag_t90(conn, customer_id: int, server_id: int):
    try:
        win = await _one(conn, f"SELECT count(*) AS c FROM {SCHEMA}.v_metrics_t90_window WHERE customer_id=%s AND server_id=%s", (customer_id, server_id))
        th  = await _one(conn, f"SELECT count(*) AS c FROM {SCHEMA}.v_task_health_t90  WHERE customer_id=%s AND server_id=%s", (customer_id, server_id))
        ep  = await _one(conn, f"SELECT count(*) AS c FROM {SCHEMA}.v_endpoint_perf_t90 WHERE customer_id=%s AND server_id=%s", (customer_id, server_id))
        sids = await _try_all(conn, f"SELECT DISTINCT server_id FROM {SCHEMA}.v_task_health_t90 WHERE customer_id=%s ORDER BY 1", (customer_id,))
        log.info("T90 diag schema=%s cust=%s srv=%s | window=%s task_health=%s endpoint=%s | th_server_ids=%s",
                 SCHEMA, customer_id, server_id,
                 (win and win.get('c') if win else None),
                 (th and th.get('c') if th else None),
                 (ep and ep.get('c') if ep else None),
                 [r.get('server_id') if hasattr(r,'get') else r['server_id'] for r in (sids or [])])
    except Exception as e:
        log.exception("T90 diag failed for cust=%s srv=%s: %s", customer_id, server_id, e)

async def _t90_fetch_endpoint_perf(conn, customer_id: int, server_id: int):
    """Return endpoint performance rows for the 90-day window (v_endpoint_perf_t90)."""
    rows = await _try_all(conn, f"""
        SELECT role, family_id, tasks, rows_moved, uptime_pct, median_rps,
               COALESCE(err_stop_rate,0.0) AS err_stop_rate,
               COALESCE(median_session_minutes,0.0) AS median_session_minutes
        FROM {SCHEMA}.v_endpoint_perf_t90
        WHERE customer_id=%s AND server_id=%s
        ORDER BY role, rows_moved DESC
    """, (customer_id, server_id)) or []
    return rows



def _stable_score_t90(t):
    """Composite stability score; higher is better."""
    try:
        u = max(0.0, min((t.uptime_pct or 0.0) / 100.0, 1.0))
        r = 1.0 - max(0.0, min((t.restarts_per_day or 0.0) / 2.0, 1.0))
        e = 1.0 - max(0.0, min((t.error_stop_rate or 0.0), 1.0))
        d = max(0.0, min((t.median_session_minutes or 0.0) / 120.0, 1.0))
        th = 0.0
        if (t.throughput_rps or 0.0) >= 0.0:
            th = min((t.throughput_rps or 0.0) / ((t.throughput_rps or 0.0) + 100.0), 1.0)
        vol = min(max(float(t.rows_moved or 0), 0.0) / (max(float(t.rows_moved or 0), 0.0) + 10_000_000.0), 1.0)
        return 0.35*u + 0.20*r + 0.15*e + 0.15*d + 0.10*th + 0.05*vol
    except Exception:
        return 0.0


def _flapper_score_t90(t):
    """Flakiness score; higher is worse."""
    try:
        rest = min(max((t.restarts_per_day or 0.0), 0.0) / 2.0, 1.0)
        err  = min(max((t.error_stop_rate or 0.0), 0.0), 1.0)
        short= 1.0 - min(max((t.median_session_minutes or 0.0), 0.0) / 120.0, 1.0)
        up   = 1.0 - min(max((t.uptime_pct or 0.0)/100.0, 0.0), 1.0)
        return 0.45*rest + 0.25*short + 0.20*err + 0.10*up
    except Exception:
        return 0.0

def _add_table_t90(doc, headers, rows):
    """Lightweight table builder for T90 sections.
    - Aligns each row to the header length (pads/truncates)
    - Uses a simple docx table style if available
    """
    # Prefer existing generic helper if present and compatible
    try:
        if callable(globals().get("_add_table")):
            return _add_table(doc, headers, rows)
    except Exception:
        pass

    t = doc.add_table(rows=1, cols=len(headers))
    try:
        t.style = "Light Shading Accent 1"
    except Exception:
        pass
    hdr = t.rows[0].cells
    for i, h in enumerate(headers):
        r = hdr[i].paragraphs[0].add_run(str(h))
        r.bold = True
    for r in rows:
        vals = list(r) if isinstance(r, (list, tuple)) else [r]
        if len(vals) < len(headers):
            vals += [""] * (len(headers) - len(vals))
        elif len(vals) > len(headers):
            vals = vals[:len(headers)]
        cells = t.add_row().cells
        for i, v in enumerate(vals):
            cells[i].paragraphs[0].add_run("" if v is None else str(v))
    return t


async def render_metricslog_90d_sections(doc, conn, customer_id: int, server_id: int, server_name: str):
    # Use a fresh connection for this section (avoids closed-conn issues)
    async with connection() as conn_t90:
        await _set_row_factory(conn_t90)

        # Fetch data
        t90        = await _t90_fetch_task_health(conn_t90, customer_id, server_id)
        try:
            endpoints  = await _t90_fetch_endpoint_perf(conn_t90, customer_id, server_id)
        except NameError:
            endpoints  = await _try_all(conn_t90, f"""
                SELECT role, family_id, tasks, rows_moved, uptime_pct, median_rps,
                       COALESCE(err_stop_rate,0.0) AS err_stop_rate,
                       COALESCE(median_session_minutes,0.0) AS median_session_minutes
                FROM {SCHEMA}.v_endpoint_perf_t90
                WHERE customer_id=%s AND server_id=%s
                ORDER BY role, rows_moved DESC
            """, (customer_id, server_id)) or []
        task_map   = await _load_task_name_map(conn_t90, customer_id, server_id)
        family_map = await _load_family_name_map(conn_t90)

        # Heading + window
        _add_heading(doc, f"MetricsLog – 90-Day Window ({server_name})", level=3)
        try:
            if t90 and t90[0].window_start and t90[0].window_end:
                _add_text(doc, f"Data window: {t90[0].window_start:%Y-%m-%d} → {t90[0].window_end:%Y-%m-%d} (based on latest log event).", size=9)
        except Exception:
            pass

        # Summary and rankings
        if not t90:
            _add_text(doc, "No per-task health rows were found in the last 90 days for this server.", size=10, italic=True)
        else:
            avg_uptime = sum(t.uptime_pct for t in t90) / max(1, len(t90))
            _add_text(doc, f"Avg Uptime (tasks): {_fmt_pct_t90(avg_uptime)}", size=9)

            def _task_label(k: str) -> str:
                return task_map.get(k, task_map.get(str(k), k))

            stable_sorted  = sorted(t90, key=_stable_score_t90, reverse=True)
            flapper_sorted = sorted(t90, key=_flapper_score_t90, reverse=True)
            top_stable = [t for t in stable_sorted if t.rows_moved >= 10000][:3]
            top_flap   = [t for t in flapper_sorted if (t.restarts_per_day >= 0.5 or t.error_stop_rate > 0.0)][:2]

            if top_stable:
                _add_heading(doc, "Top 3 Stable Producers", level=4)
                _add_table_t90(doc,
                    ["Task", "Uptime", "Restarts (total)", "Restarts/Day", "Median Run"],
                    [[
                        _task_label(ts.tkey),
                        _fmt_pct_t90(ts.uptime_pct),
                        _fmt_int_t90(ts.restarts_total or 0),
                        _fmt_float_t90(ts.restarts_per_day),
                        _fmt_duration_t90(ts.median_session_minutes),
                    ] for ts in top_stable]
                )
            else:
                _add_text(doc, "No stable producers met the minimum activity threshold.", size=9, italic=True)

            if top_flap:
                _add_heading(doc, "Top 2 Flappers", level=4)
                _add_table_t90(doc,
                    ["Task", "Uptime", "Restarts (total)", "Restarts/Day", "Median Run"],
                    [[
                        _task_label(tf.tkey),
                        _fmt_pct_t90(tf.uptime_pct),
                        _fmt_int_t90(tf.restarts_total or 0),
                        _fmt_float_t90(tf.restarts_per_day),
                        _fmt_duration_t90(tf.median_session_minutes),
                    ] for tf in top_flap]
                )
            else:
                _add_text(doc, "No flappers detected in the last 90 days.", size=9, italic=True)

        # Endpoint mix (trimmed columns)
        _add_heading(doc, "Endpoint Mix – Performance (Last 90 Days)", level=4)
        if endpoints:
            src = [e for e in endpoints if (e.get("role") if hasattr(e, "get") else e["role"]) == "SOURCE"]
            tgt = [e for e in endpoints if (e.get("role") if hasattr(e, "get") else e["role"]) == "TARGET"]

            def _get(row, k): return row.get(k) if hasattr(row, "get") else row[k]

            if src:
                _add_heading(doc, "Sources", level=5)
                _add_table_t90(doc,
                    ["Family", "Avg Uptime", "Median Run"],
                    [[
                        family_map.get(int(_get(r, "family_id") or 0), str(_get(r, "family_id"))),
                        _uptime_avg_display(r),
                        _fmt_duration_t90(_get(r, 'median_session_minutes')),
                    ] for r in src]
                )
            if tgt:
                _add_heading(doc, "Targets", level=5)
                _add_table_t90(doc,
                    ["Family", "Avg Uptime", "Median Run"],
                    [[
                        family_map.get(int(_get(r, "family_id") or 0), str(_get(r, "family_id"))),
                        _uptime_avg_display(r),
                        _fmt_duration_t90(_get(r, 'median_session_minutes')),
                    ] for r in tgt]
                )
        else:
            _add_text(doc, "No endpoint activity in the last 90 days.", size=9, italic=True)

        # Server Top-5 Flappers
        if t90:
            _add_heading(doc, "Server Top-5 Flappers", level=4)
            flapper_sorted = sorted(t90, key=_flapper_score_t90, reverse=True)
            top5 = [t for t in flapper_sorted if (t.restarts_per_day >= 0.5 or t.error_stop_rate > 0.0)][:5]
            if top5:
                _add_table_t90(doc,
                    ["Task", "Uptime", "Downtime (h)", "Restarts (total)", "Restarts/Day", "Median Run"],
                    [[
                        _task_label(t.tkey),
                        _fmt_pct_t90(t.uptime_pct),
                        _fmt_float_t90(t.downtime_hours),
                        _fmt_int_t90(t.restarts_total or 0),
                        _fmt_float_t90(t.restarts_per_day),
                        _fmt_duration_t90(t.median_session_minutes),
                    ] for t in top5]
                )

async def _load_task_name_map(conn, customer_id: int, server_id: int) -> Dict[str, str]:
    """Resolve both task_id and task_uuid -> task_name using latest ingest on this server."""
    rows = await _try_all(conn, f"""
        WITH latest AS (
          SELECT MAX(created_at) AS ts
            FROM {SCHEMA}.ingest_run
           WHERE customer_id=%s AND server_id=%s
        )
        SELECT t.task_id::text AS tid, t.task_uuid::text AS uuid, t.task_name::text AS nm
          FROM {SCHEMA}.rep_task t
          JOIN {SCHEMA}.ingest_run r ON r.run_id = t.run_id
          JOIN latest L ON L.ts = r.created_at
         WHERE r.customer_id=%s AND r.server_id=%s
    """, (customer_id, server_id, customer_id, server_id)) or []
    m: Dict[str, str] = {}
    for r in rows:
        tid = r.get("tid"); uuid = r.get("uuid"); nm = (r.get("nm") or "").strip()
        if tid:  m[tid]  = nm or tid
        if uuid: m[uuid] = nm or uuid
    return m

async def _load_family_name_map(conn) -> Dict[int, str]:
    """Family id -> friendly name. Prefer endpoint_family; fallback to endpoint_alias_map."""
    rows = await _try_all(conn, f"""
        SELECT family_id::int AS fid, COALESCE(family_name, name)::text AS nm
          FROM {SCHEMA}.endpoint_family
        WHERE COALESCE(family_name, name) IS NOT NULL
    """, ()) or []
    m: Dict[int, str] = {int(r["fid"]): (r["nm"] or "").strip() for r in rows if r.get("fid") is not None and r.get("nm")}
    if m:
        return m
    rows = await _try_all(conn, f"""
        SELECT family_id::int AS fid, MIN(alias_value)::text AS nm
          FROM {SCHEMA}.endpoint_alias_map
         GROUP BY 1
    """, ()) or []
    for r in rows:
        fid = r.get("fid"); nm = (r.get("nm") or "").strip()
        if fid is not None:
            m[int(fid)] = nm or str(fid)
    return m

async def _t90_fetch_window(conn, customer_id: int, server_id: int):
    row = await _one(conn, f"""
        SELECT window_start, window_end
        FROM {SCHEMA}.v_metrics_t90_window
        WHERE customer_id=%s AND server_id=%s
    """, (customer_id, server_id))
    return row if row else {}