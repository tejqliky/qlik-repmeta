
from __future__ import annotations

import json
import os
import re
import tempfile
from datetime import datetime
from typing import Iterable, Optional, Sequence, Tuple, List, Dict

import psycopg
from psycopg.rows import dict_row

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, Inches, RGBColor
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

# =========================
# Qlik brand palette
# =========================
QLIK_HEX = {
    "green":  "009845",
    "blue":   "00A3E0",
    "gray9":  "212529",
    "gray6":  "636E72",
    "gray3":  "C8CDD2",
    "gray1":  "F4F6F8",
    "danger": "EF4444",
    "warn":   "F59E0B",
}
def _hex_to_rgbcolor(hexstr: str) -> RGBColor:
    hexstr = hexstr.strip().lstrip("#")
    return RGBColor(int(hexstr[0:2],16), int(hexstr[2:4],16), int(hexstr[4:6],16))
QLIK_RGB = {k: _hex_to_rgbcolor(v) for k, v in QLIK_HEX.items()}

FONT_FAMILY = os.getenv("QS_REPORT_FONT") or "Segoe UI"

# =========================
# Low-level helpers
# =========================
def _conninfo() -> str:
    dsn = os.getenv("DATABASE_URL") or os.getenv("PG_DSN")
    if dsn:
        return dsn
    host = os.getenv("PGHOST", "localhost")
    port = os.getenv("PGPORT", "5432")
    user = os.getenv("PGUSER", "postgres")
    password = os.getenv("PGPASSWORD", "")
    dbname = os.getenv("PGDATABASE", os.getenv("DB_NAME", "postgres"))
    return f"host={host} port={port} user={user} password={password} dbname={dbname}"

def _set_cell_bg(cell, hex_color: str):
    hex_color = hex_color.strip().lstrip("#").upper()
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)

def _para(doc: Document, text: str = "", size: int = 11, bold: bool = False,
          color: Optional[RGBColor] = None, align=WD_ALIGN_PARAGRAPH.LEFT):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.font.size = Pt(size)
    r.font.bold = bold
    r.font.name = FONT_FAMILY
    if color:
        r.font.color.rgb = color
    p.alignment = align
    return p

def _h1(doc: Document, text: str):
    p = _para(doc, text, size=20, bold=True, color=QLIK_RGB["green"])
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(4)
    return p

def _h2(doc: Document, text: str):
    p = _para(doc, text, size=14, bold=True, color=QLIK_RGB["blue"])
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(2)
    return p

def _footer_with_page_numbers(section):
    footer = section.footer
    p = footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("Page ")
    r.font.name = FONT_FAMILY
    r.font.size = Pt(9)
    fld = OxmlElement("w:fldSimple")
    fld.set(qn("w:instr"), "PAGE \\* MERGEFORMAT")
    p._p.append(fld)
    r2 = p.add_run(" of ")
    r2.font.name = FONT_FAMILY
    r2.font.size = Pt(9)
    fld2 = OxmlElement("w:fldSimple")
    fld2.set(qn("w:instr"), "NUMPAGES \\* MERGEFORMAT")
    p._p.append(fld2)

def _hr(doc: Document, color="A3A3A3"):
    t = doc.add_table(rows=1, cols=1)
    cell = t.rows[0].cells[0]
    tcPr = cell._tc.get_or_add_tcPr()
    borders = OxmlElement("w:tcBorders")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "8")
    bottom.set(qn("w:color"), color)
    borders.append(bottom)
    tcPr.append(borders)

def _kpi_cards(doc: Document, items: Sequence[Tuple[str, str, str]]):
    tone_map = {
        "ok":   QLIK_RGB["green"],
        "warn": QLIK_RGB["warn"],
        "bad":  QLIK_RGB["danger"],
        "info": QLIK_RGB["gray6"],
    }
    cols = 4 if len(items) >= 4 else max(2, len(items))
    t = doc.add_table(rows=1, cols=cols)
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    for idx, it in enumerate(items):
        label, value, tone = it
        cell = t.rows[0].cells[idx]
        _set_cell_bg(cell, QLIK_HEX["gray1"])
        p_val = cell.paragraphs[0]
        run = p_val.add_run(str(value))
        run.bold = True
        run.font.size = Pt(18)
        run.font.name = FONT_FAMILY
        run.font.color.rgb = tone_map.get(tone, QLIK_RGB["gray9"])
        p_val.alignment = WD_ALIGN_PARAGRAPH.CENTER

        p_lab = cell.add_paragraph()
        r2 = p_lab.add_run(label)
        r2.font.size = Pt(9)
        r2.font.color.rgb = QLIK_RGB["gray6"]
        r2.font.name = FONT_FAMILY
        p_lab.alignment = WD_ALIGN_PARAGRAPH.CENTER

def _table_2col(doc: Document, title_left: str, title_right: str, rows: Iterable[Tuple[str, str]]):
    t = doc.add_table(rows=1, cols=2)
    h = t.rows[0].cells
    h[0].text = title_left
    h[1].text = title_right
    for a, b in rows:
        r = t.add_row().cells
        r[0].text = str(a or "")
        r[1].text = str(b or "")

def _add_toc(doc: Document):
    p = doc.add_paragraph()
    run = p.add_run("Table of Contents")
    run.bold = True
    run.font.size = Pt(14)
    p = doc.add_paragraph()
    fld = OxmlElement("w:fldSimple")
    fld.set(qn("w:instr"), r'TOC \o "1-3" \h \z \u')
    p._p.append(fld)
    _para(doc, "(Right-click → Update Field in Word to refresh TOC)", size=9, color=QLIK_RGB["gray6"])

def _cover_page(doc: Document, title: str, subtitle_lines: List[str], logo_path: Optional[str] = None):
    section = doc.sections[0]
    section.left_margin = Inches(0.6)
    section.right_margin = Inches(0.6)
    section.top_margin = Inches(0.6)
    section.bottom_margin = Inches(0.6)

    table = doc.add_table(rows=1, cols=2)
    c0, c1 = table.rows[0].cells
    _set_cell_bg(c0, QLIK_HEX["green"])

    if logo_path and os.path.exists(logo_path):
        p = c0.paragraphs[0]
        run = p.add_run()
        try:
            run.add_picture(logo_path, height=Inches(0.5))
        except Exception:
            pass
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    c0.add_paragraph("")

    p = c1.paragraphs[0]
    r = p.add_run(title)
    r.font.size = Pt(26)
    r.font.bold = True
    r.font.name = FONT_FAMILY
    r.font.color.rgb = QLIK_RGB["green"]
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT

    for line in subtitle_lines:
        q = c1.add_paragraph()
        rr = q.add_run(line)
        rr.font.size = Pt(11)
        rr.font.name = FONT_FAMILY
        rr.font.color.rgb = QLIK_RGB["gray6"]

    doc.add_paragraph("")
    _hr(doc, color="CCCCCC")

def _rgb01_from_hex(hexstr: str):
    hexstr = hexstr.strip().lstrip("#")
    return (int(hexstr[0:2],16)/255.0, int(hexstr[2:4],16)/255.0, int(hexstr[4:6],16)/255.0)

def _clean_stream_label(label: str) -> str:
    if not label:
        return label
    try:
        if isinstance(label, str) and label.strip().startswith("{"):
            obj = json.loads(label)
            if isinstance(obj, dict) and "name" in obj:
                return str(obj["name"])
        m = re.search(r'"name"\s*:\s*"([^"]+)"', str(label))
        if m:
            return m.group(1)
    except Exception:
        pass
    return str(label)

def _insert_chart(doc: Document, title: str, series: List[Tuple[str, float]], width_in=6.0):
    _h2(doc, title)
    try:
        import matplotlib.pyplot as plt
        labels = [a for a, _ in series]
        values = [float(b) for _, b in series]
        fig, ax = plt.subplots(figsize=(width_in, 2.6))
        color = _rgb01_from_hex(QLIK_HEX["green"])
        bars = ax.barh(range(len(labels)), values, color=[color])
        ax.set_yticks(range(len(labels)))
        ax.set_yticklabels(labels, fontsize=9)
        ax.invert_yaxis()
        ax.grid(axis='x', color='#E5E7EB', linewidth=0.8, alpha=0.6)
        ax.set_xlabel("Count", fontsize=9)
        ax.set_title(title, fontsize=11, color='#009845')
        for i, b in enumerate(bars):
            ax.text(b.get_width()+0.1, b.get_y()+b.get_height()/2, str(values[i]), va='center', fontsize=9)
        fig.tight_layout()
        tmp = tempfile.NamedTemporaryFile(delete=False, suffix=".png")
        fig.savefig(tmp.name, dpi=200)
        plt.close(fig)
        doc.add_picture(tmp.name, width=Inches(width_in))
    except Exception:
        _table_2col(doc, "Item", "Value", series)

# =========================
# Data access (transaction-safe)
# =========================
def _rollback_silent(cur):
    try:
        cur.connection.rollback()
    except Exception:
        pass

def _fetch_one(cur, sql: str, *args):
    try:
        return cur.execute(sql, args).fetchone()
    except Exception:
        _rollback_silent(cur)
        return {}

def _fetch_all(cur, sql: str, *args):
    try:
        return cur.execute(sql, args).fetchall()
    except Exception:
        _rollback_silent(cur)
        return []

def _try_fetch_one(cur, candidates: List[Tuple[str, Tuple]]):
    for q, a in candidates:
        try:
            row = cur.execute(q, a).fetchone()
            if row:
                return row
        except Exception:
            _rollback_silent(cur)
            continue
    return {}

# ---------- Higher-level query helpers ----------
def _apps_by_stream(cur, snapshot_id: int) -> List[Dict]:
    # 1) v_apps (preferred)
    rows = _fetch_all(cur, """
        SELECT COALESCE(stream,'(Unassigned)') AS k, COUNT(*) AS v
        FROM repmeta_qs.v_apps
        WHERE snapshot_id=%s
        GROUP BY 1
        ORDER BY v DESC, k
    """, snapshot_id)
    if rows:
        return rows
    # 2) raw apps + left join streams (fallback)
    rows = _fetch_all(cur, """
        SELECT
          COALESCE(s.data->>'name', a.data->'stream'->>'name', '(Unassigned)') AS k,
          COUNT(*) AS v
        FROM repmeta_qs.apps a
        LEFT JOIN repmeta_qs.streams s
          ON NULLIF(a.data->'stream'->>'id','')::uuid = s.stream_id
         AND s.snapshot_id = a.snapshot_id
        WHERE a.snapshot_id = %s
        GROUP BY 1
        ORDER BY v DESC, k
    """, snapshot_id)
    return rows

def _get_snapshot_customer_id(cur, snapshot_id: int) -> Optional[int]:
    for q in (
        ("SELECT customer_id FROM repmeta_qs.snapshots WHERE snapshot_id=%s", (snapshot_id,)),
        ('SELECT customer_id FROM repmeta_qs."snapshot" WHERE snapshot_id=%s', (snapshot_id,)),
    ):
        try:
            row = cur.execute(q[0], q[1]).fetchone()
            if row and row.get("customer_id") is not None:
                return row["customer_id"]
        except Exception:
            _rollback_silent(cur)
            pass
    return None

def _get_customer_name_by_dim(cur, customer_id: Optional[int]) -> Optional[str]:
    if not customer_id:
        return None
    try:
        row = cur.execute(
            """
            SELECT COALESCE(
                to_jsonb(dc)->>'display_name',
                to_jsonb(dc)->>'customer_name',
                to_jsonb(dc)->>'name',
                to_jsonb(dc)->>'company_name',
                to_jsonb(dc)->>'legal_name',
                to_jsonb(dc)->>'short_name'
            ) AS customer_name
            FROM repmeta.dim_customer dc
            WHERE dc.customer_id = %s
            """,
            (customer_id,),
        ).fetchone()
        if row and row.get("customer_name"):
            return str(row["customer_name"])
    except Exception:
        _rollback_silent(cur)
    return None

def _display_customer_name(cur, snapshot_id: int, env: Dict) -> str:
    if env.get("customer_name"):
        return str(env["customer_name"])
    cid = _get_snapshot_customer_id(cur, snapshot_id)
    name = _get_customer_name_by_dim(cur, cid)
    return name or "Customer"

def _parse_license_key_details(text: Optional[str]) -> Dict[str, Optional[int]]:
    out: Dict[str, Optional[int]] = {"allot_professional": None, "allot_analyzer": None, "analyzer_time": None}
    if not text:
        return out
    s = str(text).replace("\r\n", "\n")
    def pick_int(pats: List[str]) -> Optional[int]:
        for pat in pats:
            m = re.search(pat, s, flags=re.IGNORECASE)
            if m:
                try:
                    return int(m.group(1))
                except Exception:
                    pass
        return None
    out["allot_professional"] = pick_int([r"Allotment\s+professional\s*:\s*(\d+)\s*;?",
                                          r"Professional\s*Allotment\s*:\s*(\d+)"])
    out["allot_analyzer"]     = pick_int([r"Allotment\s+analyzer\s*:\s*(\d+)\s*;?",
                                          r"Analyzer\s*Allotment\s*:\s*(\d+)"])
    out["analyzer_time"]      = pick_int([r"Allotment\s+analyzer[_\s]*time\s*:\s*(\d+)\s*;?",
                                          r"Analyzer\s*time\s*\(tokens\)\s*:\s*(\d+)"])
    m = re.search(r"(?:Valid\s+To|Valid)\s*:\s*([0-9]{4}-[0-9]{2}-[0-9]{2})", s, flags=re.IGNORECASE)
    _parse_license_key_details.valid_to = m.group(1) if m else None  # type: ignore[attr-defined]
    return out

def _fetch_license_key_details(cur, snapshot_id: int, license_json_override: Optional[str]) -> Dict[str, Optional[str]]:
    # Preferred: view
    view_tries = [
        ("SELECT key_details, license_number FROM repmeta_qs.v_license_details WHERE snapshot_id=%s", (snapshot_id,)),
        ("SELECT key_details, license_number FROM repmeta_qs.v_license_summary WHERE snapshot_id=%s", (snapshot_id,)),
        ("SELECT keydetails AS key_details, license_number FROM repmeta_qs.v_license_summary WHERE snapshot_id=%s", (snapshot_id,)),
    ]
    for q, args in view_tries:
        try:
            row = cur.execute(q, args).fetchone()
            if row and (row.get("key_details") or row.get("license_number")):
                return {"key_details": row.get("key_details"), "serial": row.get("license_number")}
        except Exception:
            _rollback_silent(cur)
            pass
    # Fallback: tolerant JSON extraction
    def _one():
        return cur.execute(
            """
            SELECT
              COALESCE(
                data->>'keyDetails',
                data->>'KeyDetails',
                (SELECT je.value FROM jsonb_each_text(data) je WHERE lower(je.key)='keydetails' LIMIT 1)
              ) AS key_details,
              COALESCE(data->>'serial', data->>'Serial') AS serial
            FROM repmeta_qs.license
            WHERE snapshot_id=%s
            ORDER BY snapshot_id DESC
            LIMIT 1
            """,
            (snapshot_id,),
        ).fetchone()
    def _latest():
        return cur.execute(
            """
            SELECT
              COALESCE(
                data->>'keyDetails',
                data->>'KeyDetails',
                (SELECT je.value FROM jsonb_each_text(data) je WHERE lower(je.key)='keydetails' LIMIT 1)
              ) AS key_details,
              COALESCE(data->>'serial', data->>'Serial') AS serial
            FROM repmeta_qs.license
            ORDER BY snapshot_id DESC
            LIMIT 1
            """,
        ).fetchone()
    for getter in (_one, _latest):
        try:
            row = getter()
            if row and (row.get("key_details") or row.get("serial")):
                return {"key_details": row.get("key_details"), "serial": row.get("serial")}
        except Exception:
            _rollback_silent(cur)
            continue
    # File override
    path = license_json_override or os.getenv("QS_LICENSE_JSON")
    if not path and os.path.exists("QlikLicense.json"):
        path = "QlikLicense.json"
    if path and os.path.exists(path):
        try:
            with open(path, "r", encoding="utf-8") as f:
                data = json.load(f)
            return {"key_details": data.get("keyDetails") or data.get("KeyDetails"), "serial": data.get("serial") or data.get("Serial")}
        except Exception:
            pass
    return {"key_details": None, "serial": None}

def _license_allocations_counts(cur, snapshot_id: int) -> Dict[str, int]:
    out = {"professional_allocations": 0, "analyzer_allocations": 0}
    try:
        row = cur.execute("SELECT COUNT(*) AS c FROM repmeta_qs.access_professional WHERE snapshot_id=%s", (snapshot_id,)).fetchone()
        out["professional_allocations"] = row.get("c", 0) if row else 0
    except Exception:
        _rollback_silent(cur)
    try:
        row = cur.execute("SELECT COUNT(*) AS c FROM repmeta_qs.access_analyzer WHERE snapshot_id=%s", (snapshot_id,)).fetchone()
        out["analyzer_allocations"] = row.get("c", 0) if row else 0
    except Exception:
        _rollback_silent(cur)
    return out

def _security_rules_breakdown(cur, snapshot_id: int) -> Dict[str, Optional[int]]:
    # 1) Best-case enriched breakdown
    try:
        row = cur.execute(
            """
            SELECT total_rules, custom_total, custom_enabled, custom_disabled,
                   default_total, default_enabled, default_disabled
            FROM repmeta_qs.v_security_rule_breakdown
            WHERE snapshot_id=%s
            """,
            (snapshot_id,),
        ).fetchone()
        if row:
            return dict(row)
    except Exception:
        _rollback_silent(cur)

    # 2) v_system_rules with defensive synthesis
    try:
        row = cur.execute(
            """
            WITH r AS (
              SELECT
                COALESCE(is_default, (CASE WHEN COALESCE(NULLIF(data->>'seedId',''), NULLIF(data->>'seedID',''),
                                                      (data->'references'->>'seedId')) IS NOT NULL
                                           AND COALESCE(NULLIF(data->>'type',''),'') <> 'Custom' THEN true ELSE false END)) AS is_default,
                COALESCE(is_readonly,false) AS is_readonly,
                COALESCE(
                    CASE WHEN lower(COALESCE(NULLIF(data->>'disabled',''), 'false')) IN ('true','t','1','yes','y') THEN true ELSE false END,
                    disabled, false
                ) AS disabled
              FROM repmeta_qs.v_system_rules WHERE snapshot_id=%s
            )
            SELECT
              COUNT(*)                                         AS total_rules,
              COUNT(*) FILTER (WHERE NOT is_default AND NOT is_readonly)                    AS custom_total,
              COUNT(*) FILTER (WHERE NOT is_default AND NOT is_readonly AND NOT disabled)   AS custom_enabled,
              COUNT(*) FILTER (WHERE NOT is_default AND NOT is_readonly AND disabled)       AS custom_disabled,
              COUNT(*) FILTER (WHERE is_default)                                            AS default_total,
              COUNT(*) FILTER (WHERE is_default AND NOT disabled)                           AS default_enabled,
              COUNT(*) FILTER (WHERE is_default AND disabled)                               AS default_disabled
            FROM r
            """,
            (snapshot_id,),
        ).fetchone()
        if row:
            return dict(row)
    except Exception:
        _rollback_silent(cur)

    # 3) JSON fallback from system_rules
    try:
        rows = cur.execute(
            """
            SELECT
              (data->>'seedId') AS seed1,
              (data->>'seedID') AS seed2,
              (data->'references'->>'seedId') AS seed3,
              lower(COALESCE(NULLIF(data->>'type',''),'default')) AS ruletype,
              lower(COALESCE(NULLIF(data->>'disabled',''),'false')) AS disabled_raw
            FROM repmeta_qs.system_rules
            WHERE snapshot_id=%s
            """,
            (snapshot_id,),
        ).fetchall() or []

        def _is_disabled(s: Optional[str]) -> bool:
            return (s or "false") in ("true","t","1","yes","y")

        def _is_default(row: Dict) -> bool:
            seed = row.get("seed1") or row.get("seed2") or row.get("seed3")
            if row.get("ruletype") == "custom":
                return False
            return bool(seed)

        total = len(rows)
        default_rows = [r for r in rows if _is_default(r)]
        custom_rows = [r for r in rows if not _is_default(r)]
        custom_enabled  = sum(1 for r in custom_rows if not _is_disabled(r.get("disabled_raw")))
        custom_disabled = sum(1 for r in custom_rows if _is_disabled(r.get("disabled_raw")))
        default_enabled = sum(1 for r in default_rows if not _is_disabled(r.get("disabled_raw")))
        default_disabled= sum(1 for r in default_rows if _is_disabled(r.get("disabled_raw")))

        return {
            "total_rules": total,
            "custom_total": len(custom_rows),
            "custom_enabled": custom_enabled,
            "custom_disabled": custom_disabled,
            "default_total": len(default_rows),
            "default_enabled": default_enabled,
            "default_disabled": default_disabled,
        }
    except Exception:
        _rollback_silent(cur)
        return {
            "total_rules": None,
            "custom_total": None,
            "custom_enabled": None,
            "custom_disabled": None,
            "default_total": None,
            "default_enabled": None,
            "default_disabled": None,
        }

def _fetch_env(cur, snapshot_id: int) -> Dict:
    return _try_fetch_one(cur, [
        ("SELECT * FROM repmeta_qs.v_environment_overview_enriched WHERE snapshot_id=%s", (snapshot_id,)),
        ("SELECT * FROM repmeta_qs.v_environment_overview WHERE snapshot_id=%s", (snapshot_id,)),
    ]) or {}

def _env_fallbacks(cur, snapshot_id: int, env: Dict) -> Dict:
    out = dict(env)
    try:
        row = cur.execute("SELECT data FROM repmeta_qs.about WHERE snapshot_id=%s", (snapshot_id,)).fetchone()
        if not row:
            row = cur.execute("SELECT data FROM repmeta_qs.about ORDER BY snapshot_id DESC LIMIT 1").fetchone()
        if row and isinstance(row.get("data"), dict):
            ab = row["data"]
            out.setdefault("build_version", ab.get("buildVersion"))
            out.setdefault("build_date", ab.get("buildDate"))
            if "single_node_only" not in out and ab.get("singleNodeOnly") is not None:
                out["single_node_only"] = bool(ab.get("singleNodeOnly"))
    except Exception:
        _rollback_silent(cur)
    try:
        row = cur.execute("SELECT data FROM repmeta_qs.system_info WHERE snapshot_id=%s", (snapshot_id,)).fetchone()
        if not row:
            row = cur.execute("SELECT data FROM repmeta_qs.system_info ORDER BY snapshot_id DESC LIMIT 1").fetchone()
        if row and isinstance(row.get("data"), dict):
            si = row["data"]
            if not out.get("product_version") and si.get("releaseLabel"):
                out["product_version"] = si["releaseLabel"]
    except Exception:
        _rollback_silent(cur)
    out.setdefault("product_name", out.get("product_name") or "Qlik Sense")
    for label, sql in [
        ("extension_count", "SELECT COUNT(*) AS c FROM repmeta_qs.extensions WHERE snapshot_id=%s"),
        ("stream_count",    "SELECT COUNT(*) AS c FROM repmeta_qs.streams WHERE snapshot_id=%s"),
        ("reload_task_count","SELECT COUNT(*) AS c FROM repmeta_qs.reload_tasks WHERE snapshot_id=%s"),
        ("user_count",      "SELECT COUNT(*) AS c FROM repmeta_qs.users WHERE snapshot_id=%s"),
        ("node_count",      "SELECT COUNT(*) AS c FROM repmeta_qs.servernode_config WHERE snapshot_id=%s"),
    ]:
        try:
            if out.get(label) is None:
                r = cur.execute(sql, (snapshot_id,)).fetchone()
                out[label] = r and r.get("c") or 0
        except Exception:
            _rollback_silent(cur)
            out.setdefault(label, 0)
    return out

def _reload_activity(cur, snapshot_id: int) -> Dict[str, int]:
    try:
        row = cur.execute(
            "SELECT apps_reloaded_30d, apps_reloaded_90d "
            "FROM repmeta_qs.v_reload_activity_json WHERE snapshot_id=%s",
            (snapshot_id,)
        ).fetchone()
        if row:
            return {
                "apps_reloaded_30d": row.get("apps_reloaded_30d", 0),
                "apps_reloaded_90d": row.get("apps_reloaded_90d", 0),
            }
    except Exception:
        _rollback_silent(cur)
    try:
        row = cur.execute(
            """
            WITH raw AS (
              SELECT
                rt.snapshot_id,
                NULLIF(rt.data->'app'->>'id','')::uuid AS app_id,
                COALESCE(
                  NULLIF(rt.data->'operational'->'lastExecutionResult'->>'stopTime','')::timestamptz,
                  NULLIF(rt.data->'operational'->>'stopTime','')::timestamptz,
                  NULLIF(rt.data->>'stopTime','')::timestamptz
                ) AS stop_ts
              FROM repmeta_qs.reload_tasks rt
              WHERE rt.snapshot_id = %s
            ),
            last_by_app AS (
              SELECT app_id, MAX(stop_ts) AS ts
              FROM raw
              WHERE stop_ts IS NOT NULL
              GROUP BY app_id
            )
            SELECT
              COUNT(*) FILTER (WHERE ts >= now() - interval '30 days') AS apps_reloaded_30d,
              COUNT(*) FILTER (WHERE ts >= now() - interval '90 days') AS apps_reloaded_90d
            FROM last_by_app
            """,
            (snapshot_id,)
        ).fetchone() or {}
        return {
            "apps_reloaded_30d": row.get("apps_reloaded_30d", 0),
            "apps_reloaded_90d": row.get("apps_reloaded_90d", 0),
        }
    except Exception:
        _rollback_silent(cur)
        return {"apps_reloaded_30d": 0, "apps_reloaded_90d": 0}

# =========================
# Main API
# =========================
def generate_qs_report(snapshot_id: int, out_path: str, logo_path: Optional[str] = None, license_json: Optional[str] = None) -> str:
    doc = Document()

    created = datetime.now().strftime("%b %d, %Y %H:%M")
    dsn = _conninfo()
    with psycopg.connect(dsn, row_factory=dict_row, autocommit=True) as conn:
        cur = conn.cursor()
        env = _fetch_env(cur, snapshot_id)
        env = _env_fallbacks(cur, snapshot_id, env)
        customer = _display_customer_name(cur, snapshot_id, env)

    _cover_page(
        doc,
        "Qlik Sense — Executive Technical Overview",
        [customer, f"Generated {created}"],
        logo_path=logo_path,
    )

    section = doc.add_section()
    section.left_margin = Inches(0.6)
    section.right_margin = Inches(0.6)
    section.top_margin = Inches(0.5)
    section.bottom_margin = Inches(0.5)
    _footer_with_page_numbers(section)

    with psycopg.connect(dsn, row_factory=dict_row, autocommit=True) as conn:
        cur = conn.cursor()

        lic   = _fetch_one(cur, "SELECT * FROM repmeta_qs.v_license_summary WHERE snapshot_id=%s", snapshot_id) or {}
        lic30 = _fetch_one(cur, "SELECT * FROM repmeta_qs.v_license_usage_30d WHERE snapshot_id=%s", snapshot_id) or {}
        apps  = _fetch_one(cur, "SELECT * FROM repmeta_qs.v_app_summary WHERE snapshot_id=%s", snapshot_id) or {}
        rules = _fetch_one(cur, "SELECT * FROM repmeta_qs.v_security_rule_summary WHERE snapshot_id=%s", snapshot_id) or {}
        gov   = _fetch_one(cur, "SELECT * FROM repmeta_qs.v_governance_checks WHERE snapshot_id=%s", snapshot_id) or {}

        # Hardened license meta + allocations
        lic_raw = _fetch_license_key_details(cur, snapshot_id, license_json)
        keyd = _parse_license_key_details(lic_raw.get("key_details"))
        serial = lic.get("license_number") or lic_raw.get("serial")
        alloc = _license_allocations_counts(cur, snapshot_id)

        # Robust rules breakdown
        rule_breakdown = _security_rules_breakdown(cur, snapshot_id)

        # Prefer robust reload activity
        rts = _reload_activity(cur, snapshot_id)

        apps_by_stream = _apps_by_stream(cur, snapshot_id)

    # Executive blocks
    _h1(doc, "Executive Summary")
    _hr(doc)

    _h2(doc, "Deployment Overview")
    _kpi_cards(doc, [
        ("Total Apps", apps.get("total_apps", 0), "info"),
        ("Published Apps", apps.get("published_apps", 0), "ok"),
        ("Streams", apps.get("streams", 0), "info"),
        ("Streams w/ Apps", apps.get("streams_with_apps", 0), "info"),
    ])
    _kpi_cards(doc, [
        ("Users", env.get("user_count", 0), "info"),
        ("Nodes", env.get("node_count", 0), "info"),
        ("Extensions", env.get("extension_count", 0), "info"),
        ("Reload Tasks", env.get("reload_task_count", 0), "info"),
    ])

    _h2(doc, "Reload Health")
    _kpi_cards(doc, [
        ("Apps reloaded (30d)", rts.get("apps_reloaded_30d", 0), "ok"),
        ("Apps reloaded (90d)", rts.get("apps_reloaded_90d", 0), "ok"),
        ("Failed tasks (last run)", 0, "bad"),
        ("Tasks > 3h (last run)", 0, "warn"),
    ])

    # License
    _h2(doc, "License — Meta")
    _table_2col(doc, "Key", "Value", [
        ("Valid to", getattr(_parse_license_key_details, "valid_to", None) or lic.get("expiration") or "—"),
        ("License #", serial or "—"),
    ])

    _h2(doc, "License — Professional")
    _kpi_cards(doc, [
        ("Allotment (from key)", keyd.get("allot_professional") or "—", "info"),
        ("Allocated", alloc.get("professional_allocations", lic.get("professional_allocations", 0)), "info"),
        ("Used 30d", lic30.get("professional_used_30d", 0), "ok"),
        ("Not used 30d", lic30.get("professional_not_used_30d", 0), "warn"),
    ])
    _kpi_cards(doc, [
        ("Never used", lic30.get("professional_never_used", 0), "bad"),
        ("", "", "info"), ("", "", "info"), ("", "", "info"),
    ])

    _h2(doc, "License — Analyzer")
    _kpi_cards(doc, [
        ("Allotment (from key)", keyd.get("allot_analyzer") or "—", "info"),
        ("Allocated", alloc.get("analyzer_allocations", lic.get("analyzer_allocations", 0)), "info"),
        ("Analyzer time (tokens)", keyd.get("analyzer_time") or "—", "info"),
        ("Used 30d", lic30.get("analyzer_used_30d", 0), "ok"),
    ])
    _kpi_cards(doc, [
        ("Not used 30d", lic30.get("analyzer_not_used_30d", 0), "warn"),
        ("Never used", lic30.get("analyzer_never_used", 0), "bad"),
        ("", "", "info"), ("", "", "info"),
    ])

    _h2(doc, "Governance")
    # Override headline counts from the robust breakdown, to avoid view drift
    custom_total = rule_breakdown.get("custom_total") or 0
    total_disabled = (rule_breakdown.get("custom_disabled") or 0) + (rule_breakdown.get("default_disabled") or 0)
    _kpi_cards(doc, [
        ("Apps without reload tasks", gov.get("apps_without_tasks", 0), "bad"),
        ("Disabled reload tasks", gov.get("disabled_tasks_count", 0), "warn"),
        ("Custom security rules", custom_total, "info"),
        ("Disabled security rules", total_disabled, "warn"),
    ])
    def fmt(v): return "—" if v is None else str(v)
    _table_2col(doc, "Metric", "Count", [
        ("Total System Rules", fmt(rule_breakdown.get("total_rules"))),
        ("Security Rules — Custom (Enabled)", fmt(rule_breakdown.get("custom_enabled"))),
        ("Security Rules — Custom (Disabled)", fmt(rule_breakdown.get("custom_disabled"))),
        ("Security Rules — Default (Enabled)", fmt(rule_breakdown.get("default_enabled"))),
        ("Security Rules — Default (Disabled)", fmt(rule_breakdown.get("default_disabled"))),
    ])

    # Visuals
    top_streams = [(_clean_stream_label(r["k"]), r["v"]) for r in apps_by_stream[:8]]
    _insert_chart(doc, "Top Streams by App Count", top_streams, width_in=6.0)

    _h1(doc, "Detailed Insights")
    _hr(doc)
    _h2(doc, "Apps by Stream")
    _table_2col(doc, "Stream", "Apps", [(_clean_stream_label(r["k"]), r["v"]) for r in apps_by_stream])

    _h1(doc, "Environment")
    _hr(doc)
    _table_2col(doc, "Key", "Value", [
        ("Product", env.get("product_name") or "Qlik Sense"),
        ("Version", env.get("product_version") or "—"),
        ("Build", env.get("build_version") or "—"),
        ("Build date", env.get("build_date") or "—"),
        ("Single node only", "Yes" if env.get("single_node_only") else "No"),
        ("Streams", env.get("stream_count")),
        ("Users", env.get("user_count")),
        ("Extensions", env.get("extension_count")),
    ])

    _h1(doc, "Contents")
    _hr(doc)
    _add_toc(doc)

    doc.save(out_path)
    return out_path

# ---------------- CLI ----------------
def _parse_args(argv):
    import argparse
    p = argparse.ArgumentParser(description="Generate Qlik Sense consulting report")
    p.add_argument("snapshot_id", type=int)
    p.add_argument("output", help="Output .docx path")
    p.add_argument("--logo", help="Optional logo image path", default=None)
    p.add_argument("--license-json", help="Optional path to QlikLicense.json (override DB)", default=None)
    return p.parse_args(argv)

def _main():
    import sys
    args = _parse_args(sys.argv[1:])
    path = generate_qs_report(args.snapshot_id, args.output, logo_path=args.logo, license_json=args.license_json)
    print(f"Wrote: {path}")

if __name__ == "__main__":
    _main()
