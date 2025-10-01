import io
import os
import re
import logging
from typing import List, Tuple, Dict, Optional, Iterable

# --- helpers: group Top-5 metrics by server ----------------------------------
def _group_top5_by_server_tables(rows: List[Tuple[str, str, int]]) -> Dict[str, List[Tuple[str, int]]]:
    """
    rows: [(server_name, task_name, tables_count), ...]
    returns: {server_name: [(task_name, tables_count), ... top 5 desc]}
    """
    top: Dict[str, List[Tuple[str, int]]] = defaultdict(list)
    for server, task, tables in rows:
        top[server].append((task, tables))
    for s in list(top.keys()):
        top[s].sort(key=lambda x: x[1], reverse=True)
        top[s] = top[s][:5]
    return top

def _group_top5_by_server_apply(rows: List[Tuple[str, str, float]]) -> Dict[str, List[Tuple[str, float]]]:
    """
    rows: [(server_name, task_name, avg_apply_rec_per_sec), ...]
    returns: {server_name: [(task_name, avg_apply), ... top 5 desc]}
    """
    top: Dict[str, List[Tuple[str, float]]] = defaultdict(list)
    for server, task, avg_apply in rows:
        top[server].append((task, avg_apply))
    for s in list(top.keys()):
        top[s].sort(key=lambda x: x[1], reverse=True)
        top[s] = top[s][:5]
    return top
# -----------------------------------------------------------------------------
from collections import defaultdict
from datetime import datetime, timezone
from typing import Any, Dict, List, Tuple, Optional, NamedTuple, Dict, List, Tuple

import httpx  # used to fetch latest GA train from GitHub

try:
    from docx import Document
    from docx.shared import Pt
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
except Exception as e:  # pragma: no cover
    raise RuntimeError(
        "python-docx is required. Add `python-docx` to requirements.txt and pip install."
    ) from e

from .db import connection

SCHEMA = os.getenv("REPMETA_SCHEMA", "repmeta")
log = logging.getLogger("export_report")

# ============================================================
# Built-in master lists (fallback if DB tables are absent/empty)
# ============================================================
BUILTIN_MASTER_SOURCE_ENDPOINTS = [
    "Amazon Aurora (MySQL)",
    "Amazon Aurora (PostgreSQL)",
    "Amazon RDS for MySQL",
    "Amazon RDS for MariaDB",
    "Amazon RDS for PostgreSQL",
    "Amazon RDS for SQL Server",
    "Amazon RDS for Oracle",
    "Google Cloud SQL for MySQL",
    "Google Cloud SQL for PostgreSQL",
    "Google Cloud SQL for SQL Server",
    "Google Cloud AlloyDB for PostgreSQL",
    "Microsoft Azure SQL Database (MS-CDC)",
    "Microsoft Azure SQL Managed Instance",
    "Microsoft Azure Database for MySQL",
    "Microsoft Azure Database for PostgreSQL",
    "MongoDB Atlas",
    "Oracle (on-prem / Oracle Cloud)",
    "Teradata Vantage",
    "IBM DB2 for LUW",
    "IBM DB2 for z/OS",
    "IBM DB2 for iSeries",
    "IBM Informix",
    "Microsoft SQL Server",
    "MySQL",
    "MariaDB",
    "Percona (via MySQL endpoint)",
    "SAP Sybase ASE",
    "SAP HANA 2.0",
    "File endpoint",
    "File Channel endpoint",
    "IBM IMS (ARC)",
    "IBM VSAM Batch (ARC)",
    "Salesforce (Streaming CDC / Incremental Load)",
]

BUILTIN_MASTER_TARGET_ENDPOINTS = [
    "Amazon Aurora (MySQL)",
    "Amazon Aurora (PostgreSQL)",
    "Amazon RDS for MySQL",
    "Amazon RDS for MariaDB",
    "Amazon RDS for PostgreSQL",
    "Amazon RDS for SQL Server",
    "Amazon RDS for Oracle",
    "Google Cloud SQL for MySQL",
    "Google Cloud SQL for PostgreSQL",
    "Google Cloud SQL for SQL Server",
    "Google Cloud AlloyDB for PostgreSQL",
    "Microsoft Azure SQL Database",
    "Microsoft Azure SQL Managed Instance",
    "Microsoft Azure Database for MySQL",
    "Microsoft Azure Database for PostgreSQL",
    "Oracle (on-prem / Oracle Cloud)",
    "IBM DB2 for LUW",
    "IBM DB2 for z/OS",
    "IBM DB2 for iSeries",
    "Microsoft SQL Server",
    "MySQL",
    "PostgreSQL",                    # added – appears in your data
    "MariaDB",
    "SAP Sybase ASE",
    "File endpoint",
    "File Channel endpoint",
    "Amazon Redshift",
    "Amazon S3",
    "Amazon MSK",
    "Kafka",
    "Log Stream",                    # added – license ticker exists
    "Google BigQuery",
    "Google Cloud Storage",
    "Google Dataproc",
    "Microsoft Azure Data Lake / ADLS Gen2 / Blob",
    "Databricks (SQL Warehouse / Lakehouse)",
    "Snowflake",
    "Oracle Autonomous Data Warehouse",
]

# Will be filled at runtime (from DB if available, else fall back to BUILTIN_*)
MASTER_SOURCE_ENDPOINTS: List[str] = []
MASTER_TARGET_ENDPOINTS: List[str] = []
ALIAS_TO_CANON: Dict[str, str] = {}

# ============================================================
# Alias map (fallback) — built from Replicate license tickers
# and common variants; all keys are normalized lower/alnum.
# ============================================================
def _n(s: str) -> str:
    return re.sub(r"[^a-z0-9]+", "", s.lower())

DEFAULT_ALIAS_TO_CANON = {
    # ---------- Sources (License "Source Type" tickers) ----------
    _n("IMS"): "IBM IMS (ARC)",
    _n("VSAM"): "IBM VSAM Batch (ARC)",
    _n("MySQL"): "MySQL",
    _n("RDSMySQL"): "Amazon RDS for MySQL",
    _n("Oracle"): "Oracle (on-prem / Oracle Cloud)",
    _n("RDSPostgreSQL"): "Amazon RDS for PostgreSQL",
    _n("RDSSQLServer"): "Amazon RDS for SQL Server",
    _n("AWSAuroraPostgreSQL"): "Amazon Aurora (PostgreSQL)",
    _n("AzureSQLServerM"): "Microsoft Azure SQL Database (MS-CDC)",
    _n("File"): "File endpoint",
    _n("FileChannel"): "File Channel endpoint",
    _n("GoogleAlloyDBPostgreSQL"): "Google Cloud AlloyDB for PostgreSQL",
    _n("GoogleCloudMySQL"): "Google Cloud SQL for MySQL",
    _n("GoogleCloudPostgreSQL"): "Google Cloud SQL for PostgreSQL",
    _n("GoogleCloudSQLServer"): "Google Cloud SQL for SQL Server",
    _n("DB2LUW"): "IBM DB2 for LUW",
    _n("DB2iSeries"): "IBM DB2 for iSeries",
    _n("DB2zOS"): "IBM DB2 for z/OS",
    _n("Informix"): "IBM Informix",
    _n("SQLServer"): "Microsoft SQL Server",
    _n("AzureMySQL"): "Microsoft Azure Database for MySQL",
    _n("AzureSQL"): "Microsoft Azure SQL Database (MS-CDC)",
    _n("MongoDBAtlas"): "MongoDB Atlas",
    _n("MongoDB"): "MongoDB Atlas",
    _n("PostgreSQL"): "PostgreSQL",
    _n("SAP HANA"): "SAP HANA 2.0",
    _n("SybaseASE"): "SAP Sybase ASE",
    _n("Teradata"): "Teradata Vantage",
    _n("Salesforce"): "Salesforce (Streaming CDC / Incremental Load)",

    # Common free-form variants seen in QEM/JSON:
    _n("mssql"): "Microsoft SQL Server",
    _n("sql server"): "Microsoft SQL Server",
    _n("azuresqldatabase"): "Microsoft Azure SQL Database",
    _n("alloydb"): "Google Cloud AlloyDB for PostgreSQL",
    _n("gcp alloydb"): "Google Cloud AlloyDB for PostgreSQL",
    _n("percona"): "Percona (via MySQL endpoint)",
    _n("mariadb"): "MariaDB",

    # ---------- Targets (License "Target Type" tickers) ----------
    _n("MySQL"): "MySQL",
    _n("PostgreSQL"): "PostgreSQL",
    _n("Oracle"): "Oracle (on-prem / Oracle Cloud)",
    _n("SQLServer"): "Microsoft SQL Server",
    _n("S3"): "Amazon S3",
    _n("AmazonMSK"): "Amazon MSK",
    _n("Kafka"): "Kafka",
    _n("LogStream"): "Log Stream",
    _n("File"): "File endpoint",
    _n("FileChannel"): "File Channel endpoint",
    _n("GoogleAlloyDBPostgreSQL"): "Google Cloud AlloyDB for PostgreSQL",
    _n("BigQuery"): "Google BigQuery",
    _n("GoogleStorage"): "Google Cloud Storage",
    _n("Dataproc"): "Google Dataproc",
    _n("ADLS"): "Microsoft Azure Data Lake / ADLS Gen2 / Blob",
    _n("AzureSQLServer"): "Microsoft Azure SQL Database",
    _n("AzureMySQL"): "Microsoft Azure Database for MySQL",
    _n("AzurePostgreSQL"): "Microsoft Azure Database for PostgreSQL",
    _n("Redshift"): "Amazon Redshift",
    _n("DB2zOS"): "IBM DB2 for z/OS",
    _n("GCPDatabricksDelta"): "Databricks (SQL Warehouse / Lakehouse)",
    _n("AWSDatabricksDelta"): "Databricks (SQL Warehouse / Lakehouse)",
    _n("DatabricksAWS"): "Databricks (SQL Warehouse / Lakehouse)",
    _n("DatabricksAzure"): "Databricks (SQL Warehouse / Lakehouse)",
    _n("DatabricksGoogleCloud"): "Databricks (SQL Warehouse / Lakehouse)",
    _n("AzureDatabricksDelta"): "Databricks (SQL Warehouse / Lakehouse)",
    _n("SnowflakeAWS"): "Snowflake",
    _n("SnowflakeAzure"): "Snowflake",
    _n("SnowflakeGoogle"): "Snowflake",
    _n("Snowflake"): "Snowflake",

    # more forgiving common strings
    _n("googlegcs"): "Google Cloud Storage",
    _n("gcs"): "Google Cloud Storage",
    _n("ms adls"): "Microsoft Azure Data Lake / ADLS Gen2 / Blob",
    _n("databricks"): "Databricks (SQL Warehouse / Lakehouse)",
}

# Extra permissive synonyms used across both roles
DEFAULT_ALIAS_TO_CANON.update({
    _n("microsoftsqlserver"): "Microsoft SQL Server",
    _n("sqlserver_mscdc"): "Microsoft SQL Server",
    _n("postgres"): "PostgreSQL",
    _n("postgresql"): "PostgreSQL",
    _n("googlebigquery"): "Google BigQuery",
})

# ============================================================
# Utilities: normalization & noise filtering
# ============================================================
def _normalize_token(s: str) -> str:
    return re.sub(r"[^a-z0-9]+", "", (s or "").lower())

_STRIP_TAILS = [" (ms-cdc)"]

def _filter_noise_token(s: str) -> bool:
    return _normalize_token(s) in {"na", "n/a", "null", "nulltarget", "unknown", "(unknown)"}

def _build_master_norm() -> Dict[str, str]:
    return {
        _normalize_token(name): name
        for name in (MASTER_SOURCE_ENDPOINTS + MASTER_TARGET_ENDPOINTS)
    }

MASTER_NORM: Dict[str, str] = {}  # set at runtime after loading masters

def canonize_to_master(name: str, is_source: bool) -> str:
    if not name:
        return "Unknown"
    n = str(name).strip()

    # strip tails
    low = n.lower()
    for t in _STRIP_TAILS:
        if low.endswith(t) and _normalize_token(n) not in MASTER_NORM:
            n = n[: -len(t)]
            break

    key = _normalize_token(n)

    # direct master hit
    if key in MASTER_NORM:
        return MASTER_NORM[key]

    # alias map hit
    if key in ALIAS_TO_CANON:
        return ALIAS_TO_CANON[key]

    # last-ditch: case-insensitive compare within relevant universe
    master = MASTER_SOURCE_ENDPOINTS if is_source else MASTER_TARGET_ENDPOINTS
    for m in master:
        if m.lower() == n.lower():
            return m

    return n

# ============================================================
# DB helpers & dynamic config loaders
# ============================================================
async def _set_row_factory(conn):
    try:
        from psycopg.rows import dict_row
        await conn.set_row_factory(dict_row)
    except Exception:
        try:
            conn.row_factory = dict_row  # type: ignore[attr-defined]
        except Exception:
            pass

async def _one(conn, sql: str, params: Tuple[Any, ...]) -> Dict[str, Any]:
    cur = await conn.execute(sql, params)
    row = await cur.fetchone()
    return row or {}

async def _all(conn, sql: str, params: Tuple[Any, ...]) -> List[Dict[str, Any]]:
    cur = await conn.execute(sql, params)
    rows = await cur.fetchall()
    return list(rows or [])

async def _try_all(conn, sql: str, params: Tuple[Any, ...]) -> Optional[List[Dict[str, Any]]]:
    try:
        return await _all(conn, sql, params)
    except Exception as e:
        log.debug("optional query failed; rolling back: %s", e)
        try:
            await conn.rollback()
        except Exception:
            pass
        return None

async def _load_master_and_alias_from_db(conn):
    """Populate MASTER_* and ALIAS_TO_CANON from DB if the config tables exist; else fall back."""
    global MASTER_SOURCE_ENDPOINTS, MASTER_TARGET_ENDPOINTS, ALIAS_TO_CANON, MASTER_NORM

    sources = await _try_all(conn, f"SELECT name FROM {SCHEMA}.endpoint_master_sources ORDER BY name", ())
    targets = await _try_all(conn, f"SELECT name FROM {SCHEMA}.endpoint_master_targets ORDER BY name", ())
    MASTER_SOURCE_ENDPOINTS = [r["name"] for r in sources] if sources else BUILTIN_MASTER_SOURCE_ENDPOINTS
    MASTER_TARGET_ENDPOINTS = [r["name"] for r in targets] if targets else BUILTIN_MASTER_TARGET_ENDPOINTS

    alias_rows = await _try_all(
        conn,
        f"SELECT alias, canonical FROM {SCHEMA}.endpoint_alias_map",
        (),
    )
    if alias_rows:
        ALIAS_TO_CANON = { _normalize_token(r["alias"]): r["canonical"] for r in alias_rows }
    else:
        ALIAS_TO_CANON = dict(DEFAULT_ALIAS_TO_CANON)

    MASTER_NORM = _build_master_norm()

# ============================================================
# docx helpers
# ============================================================
def _add_title(doc: Document, text: str):
    p = doc.add_paragraph()
    p.style = doc.styles["Title"]
    r = p.add_run(text)
    r.font.size = Pt(24)
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    return p

def _add_heading(doc: Document, text: str, level: int = 1):
    return doc.add_heading(text, level=level)

def _add_text(doc: Document, text: str, size: int = 11, bold: bool = False, italic: bool = False):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.font.size = Pt(size)
    r.bold = bold
    r.italic = italic
    return p

def _add_toc(doc: Document):
    p = doc.add_paragraph()
    fld = OxmlElement("w:fldSimple")
    fld.set(qn("w:instr"), r'TOC \o "1-3" \h \z \u')
    p._p.append(fld)  # type: ignore[attr-defined]

def _set_cell_shading(cell, fill_hex: str = "EDF2FF"):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), fill_hex)
    tc_pr.append(shd)

def _cell_bold(cell, size: int = 12):
    for p in cell.paragraphs:
        for r in p.runs:
            r.bold = True
            r.font.size = Pt(size)

def _add_table(doc: Document, headers: List[str], rows: List[Tuple[Any, ...]], style: str = "Light Shading Accent 1"):
    t = doc.add_table(rows=1, cols=len(headers))
    t.style = style
    hdr = t.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = str(h)
        _cell_bold(hdr[i], 10)
    for r in rows:
        row = t.add_row().cells
        for i, v in enumerate(r):
            row[i].text = "" if v is None else str(v)
    return t

def _kpi_cards(doc: Document, cards: List[Tuple[str, Any, str]]):
    t = doc.add_table(rows=1, cols=len(cards))
    for i, (label, value, color) in enumerate(cards):
        c = t.rows[0].cells[i]
        _set_cell_shading(c, color)
        p1 = c.paragraphs[0]
        p1.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r1 = p1.add_run(str(value))
        r1.font.size = Pt(18)
        r1.bold = True
        p2 = c.add_paragraph()
        p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r2 = p2.add_run(label)
        r2.font.size = Pt(9)
        r2.italic = True
    return t

def _fmt_int(x: Any) -> str:
    try:
        return f"{int(str(x).replace(',', '').strip() or '0'):,}"
    except Exception:
        return str(x)

def _pretty_type(raw: Any) -> str:
    if not raw:
        return "Unknown"
    s = str(raw)
    rep = {
        "Sqlserver": "Microsoft SQL Server",
        "Postgresql": "PostgreSQL",
        "Oracle": "Oracle (on-prem / Oracle Cloud)",
        "Snowflake": "Snowflake",
        "Bigquery": "Google BigQuery",
        "Kafka": "",
        "Filechannel": "File Channel",
        "Mysql": "MySQL",
        "Db2": "IBM DB2 for LUW",
        "Redshift": "Amazon Redshift",
        "S3": "Amazon S3",
        "Databricks": "Databricks (SQL Warehouse / Lakehouse)",
        "Gcs": "Google Cloud Storage",
        "Adls": "Microsoft Azure Data Lake / ADLS Gen2 / Blob",
    }
    low = s.lower()
    for k, v in rep.items():
        if k.lower() in low:
            return v
    s = s.replace("Settings", "").replace("Source", "").replace("Target", "")
    return s.strip().title() or "Unknown"

def _type_icon(pretty: str) -> str:
    m = {
        "Microsoft SQL Server": "🗄️",
        "PostgreSQL": "🐘",
        "Oracle (on-prem / Oracle Cloud)": "🟥",
        "Snowflake": "❄️",
        "Google BigQuery": "🔷",
        "Kafka": "Kafka",
        "File Channel endpoint": "📁",
        "MySQL": "🐬",
        "IBM DB2 for LUW": "🟣",
        "Amazon Redshift": "📊",
        "Amazon S3": "🪣",
        "Databricks (SQL Warehouse / Lakehouse)": "🔥",
        "Google Cloud Storage": "☁️",
        "Microsoft Azure Data Lake / ADLS Gen2 / Blob": "🗂️",
        "Unknown": "🔹",
    }
    return m.get(pretty, "🔹")

def _version_badge(version: Optional[str]) -> str:
    if not version or version.strip() in ("", "-"):
        return "-  ⚠️"
    return version

def _age_badge(dt: Optional[datetime]) -> str:
    if not dt:
        return "-  ⚠️"
    now = datetime.now(timezone.utc)
    if dt.tzinfo is None:
        dt = dt.replace(tzinfo=timezone.utc)
    day_str = dt.astimezone(timezone.utc).date().isoformat()
    days = (now - dt).days
    if days > 60:
        return f"{day_str}  🔴 {days}d"
    if days > 30:
        return f"{day_str}  🟠 {days}d"
    return f"{day_str}  🟢 {days}d"

# ============================================================
# Release/train helpers
# ============================================================
TRAIN_RE = re.compile(r"^v?(\d{4})\.(\d{1,2})\.(\d+)$")  # vYYYY.M.SR

class Train(NamedTuple):
    year: int
    month_code: int  # 5 for May, 11 for Nov
    sr: int

def _parse_tag_to_train(tag: str) -> Optional[Train]:
    if not tag:
        return None
    m = TRAIN_RE.match(tag.strip())
    if not m:
        return None
    y, mcode, sr = map(int, m.groups())
    return Train(y, mcode, sr)

def _parse_replicate_version_to_train(version_str: Optional[str]) -> Optional[Train]:
    if not version_str:
        return None
    s = str(version_str).strip()
    nums = re.findall(r"\d+", s)
    if len(nums) < 3:
        return None
    y, mcode, sr = map(int, nums[:3])
    return Train(y, mcode, sr)

def _train_rank_key(t: Train) -> int:
    return t.year * 100 + t.month_code

def _trains_behind(latest: Train, current: Train) -> int:
    return max(0, _train_rank_key(latest) - _train_rank_key(current))

def _posture_label(delta: int) -> str:
    if delta == 0:
        return "🟢 Up-to-date"
    if delta == 1:
        return "🟡 1 train behind"
    if delta == 2:
        return "🟠 2 trains behind"
    if delta == 999:
        return "⚪ Unknown"
    return "🔴 >2 behind (Out-of-Support)"

async def _get_latest_ga_train(conn) -> Optional[Train]:
    row = await _one(
        conn,
        f"""
        SELECT tag, year, month_code, sr
        FROM {SCHEMA}.replicate_latest_release_cache
        ORDER BY fetched_at DESC
        LIMIT 1
        """,
        (),
    )
    if not row:
        return None
    return Train(int(row["year"]), int(row["month_code"]), int(row["sr"]))

async def _ensure_latest_cache(conn) -> Optional[Train]:
    latest_cached = await _get_latest_ga_train(conn)

    gh_api = "https://api.github.com/repos/qlik-download/replicate/releases"
    headers = {}
    tok = os.getenv("GITHUB_TOKEN")
    if tok:
        headers["Authorization"] = f"Bearer {tok}"

    try:
        async with httpx.AsyncClient(timeout=20) as cli:
            r = await cli.get(gh_api, headers=headers)
            r.raise_for_status()
            releases = r.json()
        newest: Optional[Train] = None
        newest_tag: Optional[str] = None

        for rel in releases:
            if rel.get("draft") or rel.get("prerelease"):
                continue
            tag = rel.get("tag_name") or ""
            t = _parse_tag_to_train(tag)
            if not t:
                continue
            if not newest or _train_rank_key(t) > _train_rank_key(newest):
                newest = t
                newest_tag = tag

        if newest:
            if (not latest_cached) or (_train_rank_key(newest) > _train_rank_key(latest_cached)):
                try:
                    await conn.execute(
                        f"""
                        INSERT INTO {SCHEMA}.replicate_latest_release_cache(tag, year, month_code, sr)
                        VALUES (%s,%s,%s,%s)
                        """,
                        (newest_tag, newest.year, newest.month_code, newest.sr),
                    )
                except Exception as e:
                    log.debug("cache insert failed; rolling back and continuing: %s", e)
                    try:
                        await conn.rollback()
                    except Exception:
                        pass
                latest_cached = newest
    except Exception as e:
        log.warning("GitHub latest GA fetch failed; using cache if present. err=%s", e)

    return await _get_latest_ga_train(conn)

# ============================================================
# License usage (modern layout, no "unlicensed in use" row)
# ============================================================
def _names_from_rows(rows: List[Dict[str, Any]], is_source: bool) -> set:
    out = set()
    for r in rows or []:
        raw = r.get("type")
        if not raw:
            continue
        s = str(raw).strip()
        if _filter_noise_token(s):
            continue
        pretty = canonize_to_master(s, is_source=is_source)
        if _filter_noise_token(pretty):
            continue
        out.add(pretty)
    return out

def _wrap_join(items: List[str]) -> str:
    return ", ".join(items) if items else "-"

def _license_pill_table(doc: Document, title: str,
                        used: List[str], licensed_not_used: List[str]):
    _add_text(doc, title, size=11, bold=True)
    t = _add_table(
        doc,
        headers=["", ""],
        rows=[
            ("Used", _wrap_join(used)),
            ("Licensed not used", _wrap_join(licensed_not_used)),
        ],
        style="Light Shading Accent 1",
    )
    colors = ["EEF2FF", "F1F8E9"]
    for i in range(2):
        _set_cell_shading(t.rows[i+1].cells[0], colors[i])
        _cell_bold(t.rows[i+1].cells[0], 10)

def _license_usage_section(doc: Document,
                           used_src: set, used_tgt: set,
                           lic_all_src: bool, lic_all_tgt: bool,
                           lic_src: set, lic_tgt: set):
    # Licensed universes
    lic_src_universe = set(MASTER_SOURCE_ENDPOINTS) if lic_all_src else set(lic_src)
    lic_tgt_universe = set(MASTER_TARGET_ENDPOINTS) if lic_all_tgt else set(lic_tgt)

    # Coverage math (ignore "unlicensed in use" – Replicate won’t allow it)
    src_used_ct = len(used_src if lic_all_src else (used_src & lic_src_universe))
    src_total = None if lic_all_src else len(lic_src_universe)
    src_not_used = [] if lic_all_src else sorted(lic_src_universe - used_src)

    tgt_used_ct = len(used_tgt if lic_all_tgt else (used_tgt & lic_tgt_universe))
    tgt_total = None if lic_all_tgt else len(lic_tgt_universe)
    tgt_not_used = [] if lic_all_tgt else sorted(lic_tgt_universe - used_tgt)

    # KPI tiles
    _add_text(doc, "License Usage", size=12, bold=True)
    src_kpi = f"{src_used_ct} (All Licensed)" if src_total is None else f"{src_used_ct} / {src_total}"
    tgt_kpi = f"{tgt_used_ct} (All Licensed)" if tgt_total is None else f"{tgt_used_ct} / {tgt_total}"
    _kpi_cards(doc, [
        ("Licensed Sources Used", src_kpi, "E8F5E9"),
        ("Licensed Targets Used", tgt_kpi, "E3F2FD"),
    ])
    doc.add_paragraph()

    # Panels
    _license_pill_table(doc, "Sources", sorted(used_src if lic_all_src else (used_src & lic_src_universe)), src_not_used)
    doc.add_paragraph()
    _license_pill_table(doc, "Targets", sorted(used_tgt if lic_all_tgt else (used_tgt & lic_tgt_universe)), tgt_not_used)

# ============================================================
# Server-level report (retained)
# ============================================================
async def generate_summary_docx(customer_name: str, server_name: str) -> Tuple[bytes, str]:
    async with connection() as conn:
        await _set_row_factory(conn)
        await _load_master_and_alias_from_db(conn)

        run_row = await _one(
            conn,
            f"""
            SELECT r.run_id
            FROM {SCHEMA}.ingest_run r
            JOIN {SCHEMA}.dim_customer c ON c.customer_id = r.customer_id
            JOIN {SCHEMA}.dim_server   s ON s.server_id   = r.server_id
            WHERE c.customer_name=%s AND s.server_name=%s
            ORDER BY r.run_id DESC
            LIMIT 1
            """,
            (customer_name, server_name),
        )
        if not run_row:
            raise ValueError(
                f"No run found for customer='{customer_name}' server='{server_name}'. Ingest a repo JSON first."
            )
        run_id = run_row["run_id"]

        tasks_row = await _one(conn, f"SELECT COUNT(*) AS n FROM {SCHEMA}.rep_task WHERE run_id=%s", (run_id,))
        endpoints_row = await _one(conn, f"SELECT COUNT(*) AS n FROM {SCHEMA}.rep_database WHERE run_id=%s", (run_id,))
        tasks_count = int(tasks_row.get("n", 0))
        endpoints_count = int(endpoints_row.get("n", 0))

        role_counts = await _all(
            conn,
            f"SELECT role, COUNT(*) AS n FROM {SCHEMA}.rep_database WHERE run_id=%s GROUP BY role ORDER BY role",
            (run_id,),
        )
        role_map = {str(r.get("role")): int(r.get("n", 0)) for r in role_counts}
        src_n = role_map.get("SOURCE", 0)
        tgt_n = role_map.get("TARGET", 0)

    doc = Document()
    _add_title(doc, "Qlik Replicate - Server Review")
    _add_text(doc, f"Customer: {customer_name}", size=10)
    _add_text(doc, f"Server: {server_name}", size=10)
    _add_text(doc, f"Run ID: {run_id}", size=10)
    doc.add_paragraph()
    _add_toc(doc)
    doc.add_page_break()

    _add_heading(doc, "Executive Summary", 1)
    _kpi_cards(
        doc,
        cards=[
            ("Tasks", _fmt_int(tasks_count), "E8F5E9"),
            ("Endpoints", _fmt_int(endpoints_count), "E3F2FD"),
            ("Sources", _fmt_int(src_n), "FFF3E0"),
            ("Targets", _fmt_int(tgt_n), "F3E5F5"),
        ],
    )

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    filename = f"Replicate_Server_Review_{customer_name}_{server_name}.docx".replace(" ", "_")
    return buf.read(), filename

# ============================================================
# Customer Technical Overview
# ============================================================
def _endpoint_mix_cards(doc: Document,
                        src_rows: List[Dict[str, Any]],
                        tgt_rows: List[Dict[str, Any]],
                        from_tsv: bool):
    # Normalize into (label, count)
    src_items = [( _pretty_type(r.get("type")), int(r.get("uses", 0)) ) for r in src_rows]
    tgt_items = [( _pretty_type(r.get("type")), int(r.get("uses", 0)) ) for r in tgt_rows]
    src_items.sort(key=lambda x: (-x[1], x[0]))
    tgt_items.sort(key=lambda x: (-x[1], x[0]))
    TOP = 10
    src_items = src_items[:TOP]
    tgt_items = tgt_items[:TOP]
    max_len = max(len(src_items), len(tgt_items)) or 1

    _add_text(doc, f"Endpoint Mix{' (from TSV)' if from_tsv else ' (from repo)'}", size=12, bold=True)
    table = doc.add_table(rows=max_len + 1, cols=2)
    table.style = "Light Shading Accent 1"
    table.rows[0].cells[0].text = "Sources"
    table.rows[0].cells[1].text = "Targets"
    _cell_bold(table.rows[0].cells[0], 10)
    _cell_bold(table.rows[0].cells[1], 10)

    for i in range(max_len):
        if i < len(src_items):
            s_label, s_n = src_items[i]
            icon = _type_icon(s_label)
            table.rows[i + 1].cells[0].text = f"{(icon + '  ') if icon else ''}{s_label} - {_fmt_int(s_n)}"
        else:
            table.rows[i + 1].cells[0].text = ""
        if i < len(tgt_items):
            t_label, t_n = tgt_items[i]
            icon = _type_icon(t_label)
            table.rows[i + 1].cells[1].text = f"{(icon + '  ') if icon else ''}{t_label} - {_fmt_int(t_n)}"
        else:
            table.rows[i + 1].cells[1].text = ""

async def generate_customer_report_docx(customer_name: str) -> Tuple[bytes, str]:
    async with connection() as conn:
        await _set_row_factory(conn)
        await _load_master_and_alias_from_db(conn)  # ensure masters/aliases ready

        c_row = await _one(
            conn, f"SELECT customer_id FROM {SCHEMA}.dim_customer WHERE customer_name=%s", (customer_name,)
        )
        if not c_row:
            raise ValueError(f"Customer '{customer_name}' not found. Add the customer and ingest data first.")
        customer_id = c_row["customer_id"]

        latest_train = await _ensure_latest_cache(conn)  # may be None

        # ---------- Servers ----------
        servers = await _all(
            conn,
            f"""
            SELECT server_id, server_name, COALESCE(environment,'') AS environment
            FROM {SCHEMA}.dim_server
            WHERE customer_id=%s
            ORDER BY (CASE WHEN LOWER(COALESCE(environment,'')) IN ('prod','production') THEN 0 ELSE 1 END),
                     server_name
            """,
            (customer_id,),
        )
        if not servers:
            raise ValueError(f"No servers found for customer '{customer_name}'. Ingest repository JSONs first.")

        # ---------- Latest Replicate version per server ----------
        version_rows = await _try_all(
            conn,
            f"SELECT server_name, replicate_version, last_repo FROM {SCHEMA}.v_customer_latest_runs WHERE customer_id=%s",
            (customer_id,),
        )
        if version_rows is None:
            version_rows = await _all(
                conn,
                f"""
                WITH latest AS (
                  SELECT server_id, MAX(created_at) AS last_ingest
                  FROM {SCHEMA}.ingest_run
                  WHERE customer_id=%s
                  GROUP BY server_id
                )
                SELECT s.server_name, r.replicate_version, r.created_at AS last_repo
                FROM latest l
                JOIN {SCHEMA}.ingest_run r
                  ON r.server_id=l.server_id AND r.created_at=l.last_ingest
                JOIN {SCHEMA}.dim_server s ON s.server_id=l.server_id
                ORDER BY s.server_name
                """,
                (customer_id,),
            )
        version_map = {r["server_name"]: (r.get("replicate_version"), r.get("last_repo")) for r in version_rows}

        # Posture map
        posture_map: Dict[str, Tuple[str, int, str]] = {}
        for sname, (ver_str, _last_repo_dt) in (version_map or {}).items():
            t = _parse_replicate_version_to_train(ver_str)
            if t and latest_train:
                delta = _trains_behind(latest_train, t)
                posture_map[sname] = (ver_str or "-", delta, _posture_label(delta))
            else:
                posture_map[sname] = (ver_str or "-", 999, _posture_label(999))

        # ---------- Repo totals ----------
        latest_repo_rows = await _all(
            conn,
            f"""
            WITH latest AS (
              SELECT server_id, MAX(created_at) AS last_ingest
              FROM {SCHEMA}.ingest_run
              WHERE customer_id=%s
              GROUP BY server_id
            ),
            runs AS (
              SELECT r.server_id, r.run_id
              FROM latest l
              JOIN {SCHEMA}.ingest_run r
                ON r.server_id=l.server_id AND r.created_at=l.last_ingest
            )
            SELECT
              (SELECT COUNT(*) FROM {SCHEMA}.rep_task     t JOIN runs u ON t.run_id=u.run_id) AS tasks,
              (SELECT COUNT(*) FROM {SCHEMA}.rep_database d JOIN runs u ON d.run_id=u.run_id) AS endpoints,
              (SELECT COUNT(*) FROM {SCHEMA}.rep_database d JOIN runs u ON d.run_id=u.run_id WHERE d.role='SOURCE') AS src,
              (SELECT COUNT(*) FROM {SCHEMA}.rep_database d JOIN runs u ON d.run_id=u.run_id WHERE d.role='TARGET') AS tgt
            """,
            (customer_id,),
        )
        repo_totals = latest_repo_rows[0] if latest_repo_rows else {"tasks": 0, "endpoints": 0, "src": 0, "tgt": 0}

        # ---------- Endpoint mix (prefer view) ----------
        mix_src = await _try_all(
            conn,
            f"""
            SELECT type, uses FROM {SCHEMA}.v_qem_endpoint_mix
            WHERE customer_id=%s AND role='SOURCE'
            ORDER BY uses DESC, type
            """,
            (customer_id,),
        )
        mix_tgt = await _try_all(
            conn,
            f"""
            SELECT type, uses FROM {SCHEMA}.v_qem_endpoint_mix
            WHERE customer_id=%s AND role='TARGET'
            ORDER BY uses DESC, type
            """,
            (customer_id,),
        )
        if mix_src is None or mix_tgt is None:
            src_types_tsv = await _all(
                conn,
                f"""
                SELECT source_type AS type, COUNT(*) AS uses
                FROM {SCHEMA}.qem_task_perf
                WHERE customer_id=%s AND source_type IS NOT NULL
                GROUP BY source_type
                ORDER BY uses DESC, type
                """,
                (customer_id,),
            )
            tgt_types_tsv = await _all(
                conn,
                f"""
                SELECT target_type AS type, COUNT(*) AS uses
                FROM {SCHEMA}.qem_task_perf
                WHERE customer_id=%s AND target_type IS NOT NULL
                GROUP BY target_type
                ORDER BY uses DESC, type
                """,
                (customer_id,),
            )
            src_types_repo: List[Dict[str, Any]] = []
            tgt_types_repo: List[Dict[str, Any]] = []
            if not src_types_tsv or not tgt_types_tsv:
                types_repo = await _all(
                    conn,
                    f"""
                    WITH latest AS (
                      SELECT server_id, MAX(created_at) AS last_ingest
                      FROM {SCHEMA}.ingest_run
                      WHERE customer_id=%s
                      GROUP BY server_id
                    ),
                    runs AS (
                      SELECT r.server_id, r.run_id
                      FROM latest l
                      JOIN {SCHEMA}.ingest_run r
                        ON r.server_id=l.server_id AND r.created_at=l.last_ingest
                    )
                    SELECT role, db_settings_type AS type, COUNT(*) AS uses
                    FROM {SCHEMA}.rep_database d
                    JOIN runs u ON d.run_id=u.run_id
                    GROUP BY role, db_settings_type
                    """,
                    (customer_id,),
                )
                src_types_repo = [r for r in types_repo if str(r.get("role")) == "SOURCE"]
                tgt_types_repo = [r for r in types_repo if str(r.get("role")) == "TARGET"]
            src_rows_used = (src_types_tsv or src_types_repo)
            tgt_rows_used = (tgt_types_tsv or tgt_types_repo)
        else:
            src_rows_used = mix_src
            tgt_rows_used = mix_tgt

        # ---------- Peak volumes (TSV) ----------
        peak_fl = await _one(
            conn,
            f"""
            SELECT s.server_name, q.task_name, q.fl_total_records
            FROM {SCHEMA}.qem_task_perf q
            JOIN {SCHEMA}.dim_server s USING (server_id)
            WHERE q.customer_id=%s AND q.fl_total_records IS NOT NULL
            ORDER BY q.fl_total_records DESC NULLS LAST
            LIMIT 1
            """,
            (customer_id,),
        )
        peak_cdc = await _one(
            conn,
            f"""
            SELECT s.server_name, q.task_name, q.cdc_commit_change_records
            FROM {SCHEMA}.qem_task_perf q
            JOIN {SCHEMA}.dim_server s USING (server_id)
            WHERE q.customer_id=%s AND q.cdc_commit_change_records IS NOT NULL
            ORDER BY q.cdc_commit_change_records DESC NULLS LAST
            LIMIT 1
            """,
            (customer_id,),
        )

        # ---------- Server rollup (prefer view) ----------
        rollup_rows = await _try_all(
            conn,
            f"SELECT * FROM {SCHEMA}.v_server_rollup WHERE customer_id=%s ORDER BY server_name",
            (customer_id,),
        )
        if rollup_rows is None:
            rollup_rows = await _all(
                conn,
                f"""
                WITH base AS (
                  SELECT server_id,
                         COUNT(DISTINCT task_name)   AS tasks,
                         COUNT(DISTINCT source_name) AS src_eps,
                         COUNT(DISTINCT target_name) AS tgt_eps
                  FROM {SCHEMA}.qem_task_perf
                  WHERE customer_id=%s
                  GROUP BY server_id
                ),
                last_repo AS (
                  SELECT server_id, MAX(created_at) AS last_repo
                  FROM {SCHEMA}.ingest_run
                  WHERE customer_id=%s
                  GROUP BY server_id
                ),
                last_qem AS (
                  SELECT server_id, MAX(created_at) AS last_qem
                  FROM {SCHEMA}.qem_ingest_run
                  WHERE customer_id=%s
                  GROUP BY server_id
                )
                SELECT s.server_name,
                       COALESCE(b.tasks,0)   AS tasks,
                       COALESCE(b.src_eps,0) AS src_eps,
                       COALESCE(b.tgt_eps,0) AS tgt_eps,
                       lr.last_repo,
                       lq.last_qem
                FROM {SCHEMA}.dim_server s
                LEFT JOIN base     b  USING (server_id)
                LEFT JOIN last_repo lr USING (server_id)
                LEFT JOIN last_qem  lq USING (server_id)
                WHERE s.customer_id=%s
                ORDER BY s.server_name
                """,
                (customer_id, customer_id, customer_id, customer_id),
            )

        # ---------- Primary pair per server (prefer view) ----------
        primary_pairs = await _try_all(
            conn,
            f"SELECT server_name, source_type, target_type, n FROM {SCHEMA}.v_primary_pairs WHERE customer_id=%s ORDER BY server_name",
            (customer_id,),
        )
        if primary_pairs is None:
            primary_pairs = await _all(
                conn,
                f"""
                WITH pairs AS (
                  SELECT server_id, source_type, target_type, COUNT(*) AS n
                  FROM {SCHEMA}.qem_task_perf
                  WHERE customer_id=%s
                  GROUP BY server_id, source_type, target_type
                ),
                ranked AS (
                  SELECT p.*, ROW_NUMBER() OVER (PARTITION BY p.server_id ORDER BY n DESC NULLS LAST, source_type, target_type) AS rn
                  FROM pairs p
                )
                SELECT s.server_name, r.source_type, r.target_type, r.n
                FROM ranked r
                JOIN {SCHEMA}.dim_server s USING (server_id)
                WHERE rn=1
                ORDER BY s.server_name
                """,
                (customer_id,),
            )
        primary_map = {r["server_name"]: (r.get("source_type"), r.get("target_type"), r.get("n")) for r in primary_pairs}

        # ---------- Coverage matrix (prefer view) ----------
        coverage = await _try_all(
            conn,
            f"SELECT s_type, t_type, n FROM {SCHEMA}.v_coverage_matrix WHERE customer_id=%s",
            (customer_id,),
        )
        if coverage is None:
            coverage = await _all(
                conn,
                f"""
                SELECT COALESCE(source_type,'(unknown)') AS s_type,
                       COALESCE(target_type,'(unknown)') AS t_type,
                       COUNT(*) AS n
                FROM {SCHEMA}.qem_task_perf
                WHERE customer_id=%s
                GROUP BY COALESCE(source_type,'(unknown)'), COALESCE(target_type,'(unknown)')
                """,
                (customer_id,),
            )
        src_types = sorted({_pretty_type(r["s_type"]) for r in coverage})
        tgt_types = sorted({_pretty_type(r["t_type"]) for r in coverage})
        cov_map: Dict[Tuple[str, str], int] = {}
        for r in coverage:
            cov_map[(_pretty_type(r["s_type"]), _pretty_type(r["t_type"]))] = int(r["n"])

        # ---------- License coverage (use repmeta.v_license_vs_usage) ----------
        # Prefer the DB view that already expands explicit tickers + wildcard "ALL" into families.
        # It exposes: customer_id, ef_role ('SOURCE'/'TARGET'), family_name, is_licensed (bool),
        # and configured_count (endpoints configured in rep_database mapped to that family).
        lic_rows = await _all(
            conn,
            f"""
            SELECT ef_role, family_name, is_licensed, COALESCE(configured_count,0) AS configured_count
            FROM {SCHEMA}.v_license_vs_usage
            WHERE customer_id=%s
            """ ,
            (customer_id,),
        )

        used_source_types, used_target_types = set(), set()
        lic_src, lic_tgt = set(), set()

        for r in (lic_rows or []):
            fam = str(r.get("family_name") or "").strip()
            role = str(r.get("ef_role") or "").strip().upper()
            if not fam or not role:
                continue
            if bool(r.get("is_licensed")):
                if role == "SOURCE":
                    lic_src.add(fam)
                elif role == "TARGET":
                    lic_tgt.add(fam)
            if int(r.get("configured_count") or 0) > 0:
                if role == "SOURCE":
                    used_source_types.add(fam)
                elif role == "TARGET":
                    used_target_types.add(fam)

        # Determine whether the license effectively covers ALL families.
        lic_all_src = lic_src == set(MASTER_SOURCE_ENDPOINTS)
        lic_all_tgt = lic_tgt == set(MASTER_TARGET_ENDPOINTS)

        # Fallback: if the view is missing/empty, fall back to previous behavior
        if (not lic_rows):
            lic_row = await _one(
                conn,
                f"SELECT * FROM {SCHEMA}.v_latest_customer_license WHERE customer_id=%s",
                (customer_id,),
            )
            lic_all_src = bool(lic_row.get("licensed_all_sources")) if lic_row else False
            lic_all_tgt = bool(lic_row.get("licensed_all_targets")) if lic_row else False
            lic_src, lic_tgt = set(), set()
            if lic_row:
                for s in (lic_row.get("licensed_sources") or []):
                    lic_src.add(canonize_to_master(s, is_source=True))
                for t in (lic_row.get("licensed_targets") or []):
                    lic_tgt.add(canonize_to_master(t, is_source=False))

            # Prior approach for "used" based on endpoint mix tables (already queried above)
            used_source_types = _names_from_rows(src_rows_used, is_source=True)
            used_target_types = _names_from_rows(tgt_rows_used, is_source=False)
            top_tables_by_server: Dict[str, List[Tuple[str, int]]] = defaultdict(list)
            # ---------- NEW: Top-5 tasks by #tables per server (precompute once) ----------
            top_tables_rows = await _try_all(
                conn,
                f"""
                WITH latest AS (
                  SELECT server_id, MAX(created_at) AS last_ingest
                  FROM {SCHEMA}.ingest_run
                  WHERE customer_id=%s
                  GROUP BY server_id
                ),
                counts AS (
                  SELECT r.server_id, t.task_name, COUNT(*) AS n_tables
                  FROM {SCHEMA}.rep_task_table tt
                  JOIN {SCHEMA}.rep_task t ON t.task_id = tt.task_id AND t.run_id = tt.run_id
                  JOIN {SCHEMA}.ingest_run r ON r.run_id = tt.run_id
                  JOIN latest l
                    ON l.server_id = r.server_id AND r.created_at = l.last_ingest
                  GROUP BY r.server_id, t.task_name
                ),
                ranked AS (
                  SELECT server_id, task_name, n_tables,
                         ROW_NUMBER() OVER (PARTITION BY server_id ORDER BY n_tables DESC, task_name) AS rn
                  FROM counts
                )
                SELECT s.server_name, task_name, n_tables
                FROM ranked rk
                JOIN {SCHEMA}.dim_server s ON s.server_id = rk.server_id
                WHERE rk.rn <= 5 AND s.customer_id = %s
                ORDER BY s.server_name, n_tables DESC, task_name
                """,
                (customer_id, customer_id),
            )
            if top_tables_rows:
                for r in top_tables_rows:
                    top_tables_by_server[r["server_name"]].append((r["task_name"], int(r["n_tables"])))

    # ---------------- DOCX BUILD ----------------
    doc = Document()
    _add_title(doc, "Customer Technical Overview")
    _add_text(doc, f"Customer: {customer_name}", size=10)
    doc.add_paragraph()
    _add_toc(doc)
    doc.add_page_break()

    # 1) Executive Summary
    _add_heading(doc, "1. Executive Summary", 1)

    oos_count = sum(1 for _, (_, delta, _) in posture_map.items() if delta > 2)
    cards = [
        ("Servers", _fmt_int(len(servers)), "E3F2FD"),
        ("Tasks", _fmt_int(repo_totals.get("tasks", 0)), "E8F5E9"),
        ("Endpoints", _fmt_int(repo_totals.get("endpoints", 0)), "F3E5F5"),
        ("Sources", _fmt_int(repo_totals.get("src", 0)), "FFF3E0"),
        ("Targets", _fmt_int(repo_totals.get("tgt", 0)), "FFF3E0"),
        ("Servers OOS", _fmt_int(oos_count), "FFE9E6"),
    ]
    _kpi_cards(doc, cards)
    doc.add_paragraph()

    # Versions table with posture
    _add_text(doc, "Replicate Versions by Server", size=12, bold=True)
    headers = ["Server", "Replicate Version", "Last Repo Ingest", "Posture"]
    rows = []
    for s in sorted(version_map.keys()):
        ver, last_repo = version_map.get(s, (None, None))
        _, delta, label = posture_map.get(s, (ver or "-", 999, _posture_label(999)))
        rows.append((s, _version_badge(ver), _age_badge(last_repo), label))
    _add_table(doc, headers=headers, rows=rows, style="Light Shading Accent 3")
    doc.add_paragraph()

    # Endpoint Mix
    _endpoint_mix_cards(doc, src_rows_used, tgt_rows_used, bool(src_rows_used and tgt_rows_used))
    doc.add_paragraph()

    # License Usage - modern design (no "unlicensed in use")
    _license_usage_section(
        doc,
        used_source_types, used_target_types,
        lic_all_src, lic_all_tgt,
        lic_src, lic_tgt,
    )
    doc.add_page_break()

    # 2) Customer Insights
    _add_heading(doc, "2. Customer Insights", 1)

    _add_text(doc, "Version posture vs latest GA train", size=11, bold=True)
    if not latest_train:
        _add_text(doc, "⚠ Latest GA train not available (no cache & GitHub fetch failed).", size=10, italic=True)
    else:
        behind2 = []
        beyond2 = []
        for sname, (ver_str, delta, label) in posture_map.items():
            if delta == 2:
                behind2.append((sname, ver_str, label))
            elif delta > 2:
                beyond2.append((sname, ver_str, label))
        headers = ["Server", "Replicate", "Posture"]
        if not behind2 and not beyond2:
            _add_text(doc, "All servers are within 0-1 train of latest GA. ✅", size=10, italic=True)
        else:
            if behind2:
                _add_text(doc, "🟠 2 trains behind", size=10, bold=True)
                _add_table(doc, headers, [(s, _version_badge(v), l) for s, v, l in sorted(behind2)], "Light Shading Accent 2")
            if beyond2:
                _add_text(doc, "🔴 >2 trains behind (Out-of-Support)", size=10, bold=True)
                _add_table(doc, headers, [(s, _version_badge(v), l) for s, v, l in sorted(beyond2)], "Light Shading Accent 2")

    doc.add_paragraph()

    # Insight #1: tasks with null target
    try:
        async with connection() as conn_ro:
            await _set_row_factory(conn_ro)
            tasks_null_tgt = await _all(
                conn_ro,
                f"""
                WITH latest_qem AS (
                  SELECT r.server_id, MAX(r.created_at) AS max_created
                  FROM {SCHEMA}.qem_ingest_run r
                  WHERE r.customer_id=%s
                  GROUP BY r.server_id
                ),
                rows AS (
                  SELECT q.*
                  FROM {SCHEMA}.qem_task_perf q
                  JOIN {SCHEMA}.qem_ingest_run r ON r.qem_run_id = q.qem_run_id
                  JOIN latest_qem l
                    ON l.server_id = r.server_id AND r.created_at = l.max_created
                  WHERE q.customer_id=%s
                )
                SELECT s.server_name,
                       q.task_name,
                       COALESCE(NULLIF(q.target_name,''), '(unknown)') AS target_name
                FROM rows q
                JOIN {SCHEMA}.dim_server s USING (server_id)
                WHERE (LOWER(BTRIM(q.target_type)) = 'null target' OR q.target_type IS NULL)
                GROUP BY s.server_name, q.task_name, COALESCE(NULLIF(q.target_name,''), '(unknown)')
                ORDER BY s.server_name, q.task_name
                """,
                (customer_id, customer_id),
            )
        if tasks_null_tgt:
            _add_text(doc, "Tasks with missing Target Type", size=11, bold=True)
            headers = ["Server", "Task", "Target Endpoint"]
            rows = [(r.get("server_name") or "-", r.get("task_name") or "-", r.get("target_name") or "-")
                    for r in tasks_null_tgt]
            _add_table(doc, headers=headers, rows=rows, style="Light Shading")
        else:
            _add_text(doc, "No tasks with missing Target Type found in latest QEM snapshot.", size=10, italic=True)
    except Exception as e:
        _add_text(doc, f"⚠ Insight #1 failed: {type(e).__name__}: {e}", size=10, italic=True)

    doc.add_paragraph()

    # Insight #2: duplicate endpoints (identical settings)
    try:
        async with connection() as conn_ro2:
            await _set_row_factory(conn_ro2)
            dup_eps = await _all(
                conn_ro2,
                f"""
                WITH latest AS (
                  SELECT server_id, MAX(created_at) AS last_ingest
                  FROM {SCHEMA}.ingest_run
                  WHERE customer_id=%s
                  GROUP BY server_id
                ),
                runs AS (
                  SELECT r.server_id, r.run_id
                  FROM latest l
                  JOIN {SCHEMA}.ingest_run r
                    ON r.server_id = l.server_id AND r.created_at = l.last_ingest
                ),

                /* Union all detailed endpoint tables + generic fallback */
                unioned AS (
                  SELECT endpoint_id, settings_json FROM {SCHEMA}.rep_db_postgresql_source
                  UNION ALL SELECT endpoint_id, settings_json FROM {SCHEMA}.rep_db_postgresql_target
                  UNION ALL SELECT endpoint_id, settings_json FROM {SCHEMA}.rep_db_sqlserver_source
                  UNION ALL SELECT endpoint_id, settings_json FROM {SCHEMA}.rep_db_sqlserver_target
                  UNION ALL SELECT endpoint_id, settings_json FROM {SCHEMA}.rep_db_mysql_source
                  UNION ALL SELECT endpoint_id, settings_json FROM {SCHEMA}.rep_db_mysql_target
                  UNION ALL SELECT endpoint_id, settings_json FROM {SCHEMA}.rep_db_oracle_source
                  UNION ALL SELECT endpoint_id, settings_json FROM {SCHEMA}.rep_db_oracle_target
                  UNION ALL SELECT endpoint_id, settings_json FROM {SCHEMA}.rep_db_snowflake_target
                  UNION ALL SELECT endpoint_id, settings_json FROM {SCHEMA}.rep_db_redshift_target
                  UNION ALL SELECT endpoint_id, settings_json FROM {SCHEMA}.rep_db_s3_target
                  UNION ALL SELECT endpoint_id, settings_json FROM {SCHEMA}.rep_db_databricks_delta_target
                  UNION ALL SELECT endpoint_id, settings_json FROM {SCHEMA}.rep_db_databricks_cloud_storage_target
                  UNION ALL SELECT endpoint_id, settings_json FROM {SCHEMA}.rep_db_azure_adls_target
                  UNION ALL SELECT endpoint_id, settings_json FROM {SCHEMA}.rep_db_hdinsight_target
                  UNION ALL SELECT endpoint_id, settings_json FROM {SCHEMA}.rep_db_hadoop_target
                  UNION ALL SELECT endpoint_id, settings_json FROM {SCHEMA}.rep_db_kafka_target
                  UNION ALL SELECT endpoint_id, settings_json FROM {SCHEMA}.rep_db_amazon_msk_target
                  UNION ALL SELECT endpoint_id, settings_json FROM {SCHEMA}.rep_db_confluent_cloud_target
                  UNION ALL SELECT endpoint_id, settings_json FROM {SCHEMA}.rep_db_eventhubs_target
                  UNION ALL SELECT endpoint_id, settings_json FROM {SCHEMA}.rep_db_pubsub_target
                  UNION ALL SELECT endpoint_id, settings_json FROM {SCHEMA}.rep_db_logstream_target
                  UNION ALL SELECT endpoint_id, settings_json FROM {SCHEMA}.rep_db_file_source
                  UNION ALL SELECT endpoint_id, settings_json FROM {SCHEMA}.rep_db_file_target
                  UNION ALL SELECT endpoint_id, settings_json FROM {SCHEMA}.rep_db_file_channel_target
                  UNION ALL SELECT endpoint_id, settings_json FROM {SCHEMA}.rep_db_db2_luw_source
                  UNION ALL SELECT endpoint_id, settings_json FROM {SCHEMA}.rep_db_db2_zos_source
                  UNION ALL SELECT endpoint_id, settings_json FROM {SCHEMA}.rep_db_db2_zos_target
                  UNION ALL SELECT endpoint_id, settings_json FROM {SCHEMA}.rep_db_db2_iseries_source
                  UNION ALL SELECT endpoint_id, settings_json FROM {SCHEMA}.rep_db_teradata_source
                  UNION ALL SELECT endpoint_id, settings_json FROM {SCHEMA}.rep_db_teradata_target
                  UNION ALL SELECT endpoint_id, settings_json FROM {SCHEMA}.rep_db_odbc_source
                  UNION ALL SELECT endpoint_id, settings_json FROM {SCHEMA}.rep_db_odbc_target
                  UNION ALL SELECT endpoint_id, settings_json FROM {SCHEMA}.rep_db_informix_source
                  UNION ALL SELECT endpoint_id, settings_json FROM {SCHEMA}.rep_db_vsam_source
                  UNION ALL SELECT endpoint_id, settings_json FROM {SCHEMA}.rep_db_ims_source
                  UNION ALL SELECT endpoint_id, settings_json FROM {SCHEMA}.rep_db_gcs_target
                  UNION ALL SELECT endpoint_id, settings_json FROM {SCHEMA}.rep_db_ms_fabric_dw_target
                  UNION ALL SELECT endpoint_id, settings_json FROM {SCHEMA}.rep_db_settings_json
                ),
                base AS (
                  SELECT
                    s.server_name,
                    d.role,
                    d.db_settings_type,
                    d.name AS endpoint_name,
                    u.settings_json
                  FROM {SCHEMA}.rep_database d
                  JOIN runs r ON d.run_id = r.run_id
                  JOIN {SCHEMA}.dim_server s ON s.server_id = r.server_id
                  JOIN unioned u ON u.endpoint_id = d.endpoint_id
                ),
                clean AS (
                  SELECT
                    server_name,
                    role,
                    db_settings_type,
                    endpoint_name,
                    jsonb_strip_nulls(
                      (settings_json
                        - 'Name' - 'EndpointName' - 'DisplayName' - 'Description'
                        - 'Id' - 'ID' - 'Guid' - 'GUID'
                        - 'CreatedTime' - 'ModifiedTime' - 'LastTestConnection'
                        - 'Password' - 'Pwd' - 'Secret' - 'AccessKey' - 'SecretKey' - 'Token'
                        - 'ProxyPassword' - 'SaslPassword' - 'OAuthToken'
                        - 'privateKey' - 'privateKeyFile'
                      )
                    ) AS cfg
                  FROM base
                )
                SELECT
                  server_name,
                  role,
                  db_settings_type,
                  md5(cfg::text) AS cfg_sig,
                  ARRAY_AGG(endpoint_name ORDER BY endpoint_name) AS endpoint_names,
                  COUNT(*) AS n
                FROM clean
                GROUP BY server_name, role, db_settings_type, md5(cfg::text)
                HAVING COUNT(*) > 1
                ORDER BY server_name, role, n DESC
                """,
                (customer_id,),
            )
        _add_text(doc, "Endpoints with identical configuration", size=11, bold=True)
        if dup_eps:
            headers = ["Server", "Role", "Type", "Endpoints (identical settings)", "Count"]
            rows = []
            for r in dup_eps:
                rows.append((
                    r.get("server_name") or "-",
                    r.get("role") or "-",
                    _pretty_type(r.get("db_settings_type") or "") or "-",
                    ", ".join(r.get("endpoint_names") or []) or "-",
                    _fmt_int(r.get("n")),
                ))
            _add_table(doc, headers=headers, rows=rows, style="Light Shading Accent 2")
        else:
            _add_text(doc, "No duplicate endpoint configurations detected in latest runs.", size=10, italic=True)
    except Exception as e:
        _add_text(doc, f"⚠ Insight #2 failed: {type(e).__name__}: {e}", size=10, italic=True)

    doc.add_page_break()

    # 3) Environment & Inventory
    _add_heading(doc, "3. Environment & Inventory", 1)
    _add_text(doc, "Servers Overview", size=12, bold=True)
    headers = ["Server", "Tasks", "Source EPs", "Target EPs", "Last Repo Ingest", "Last QEM Ingest", "Replicate", "Posture"]
    rows = []
    for r in rollup_rows:
        sname = r.get("server_name")
        ver_str, delta, label = posture_map.get(sname, ("-", 999, _posture_label(999)))
        rows.append([
            sname,
            _fmt_int(r.get("tasks")),
            _fmt_int(r.get("src_eps")),
            _fmt_int(r.get("tgt_eps")),
            _age_badge(r.get("last_repo")),
            _age_badge(r.get("last_qem")),
            _version_badge(ver_str),
            label,
        ])
    _add_table(doc, headers=headers, rows=rows, style="Light Shading Accent 1")
    doc.add_paragraph()

    # Coverage Matrix
    _add_text(doc, "Source × Target Coverage (TSV)", size=12, bold=True)
    headers = ["Source \\ Target"] + [f"{_type_icon(t)}  {t}" for t in tgt_types]
    rows = []
    for s_type in src_types:
        row = [(f"{_type_icon(s_type)}  {s_type}" if _type_icon(s_type) else s_type)]
        for t_type in tgt_types:
            row.append(_fmt_int(cov_map.get((s_type, t_type), 0)))
        rows.append(tuple(row))
    _add_table(doc, headers=headers, rows=rows, style="Light Shading Accent 1")
    doc.add_page_break()

    # 4) Server Deep Dives
    _add_heading(doc, "4. Server Deep Dives", 1)
    for s in [srv["server_name"] for srv in servers]:
        _add_heading(doc, s, 2)
        pair = primary_map.get(s)
        if pair:
            _add_text(doc, f"Primary pair: {_pretty_type(pair[0])} → {_pretty_type(pair[1])} ({_fmt_int(pair[2])} tasks)",
                      size=10, italic=True)

        # NEW: Top-5 tasks by number of tables for this server
        top_rows = top_tables_by_server.get(s) or []
        if top_rows:
            last_repo_dt = (version_map.get(s) or (None, None))[1]
            when_txt = ""
            try:
                if last_repo_dt:
                    when_txt = f" (latest repo ingest: {last_repo_dt.date()})"
            except Exception:
                pass
            _add_text(doc, f"Top-5 tasks by number of tables{when_txt}", size=11, bold=True)
            _add_table(
                doc,
                headers=["Task", "# Tables"],
                rows=[(name, _fmt_int(n)) for (name, n) in top_rows],
                style="Light Shading Accent 1",
            )
        else:
            _add_text(doc, "Top-5 tasks by number of tables — no table data found for latest ingest.", size=10, italic=True)

        doc.add_paragraph()

    # Index
    _add_heading(doc, "Index", 1)
    p = doc.add_paragraph("Update the index in Word: References → Update Table.")
    p.runs[0].italic = True

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    filename = f"Customer_Technical_Overview_{customer_name}.docx".replace(" ", "_")
    return buf.read(), filename