# patch_export_report.py
# Usage:
#   python patch_export_report.py
#
# Produces: export_report.enhanced.py next to export_report.py
# Keeps all existing behavior; adds page numbers (not on cover), number alignment, and charts.

from pathlib import Path

SRC = Path("export_report.py")
DST = Path("export_report.enhanced.py")

text = SRC.read_text(encoding="utf-8")

def insert_after_first(haystack: str, needle: str, insertion: str) -> str:
    i = haystack.find(needle)
    if i == -1:
        return haystack
    j = i + len(needle)
    return haystack[:j] + insertion + haystack[j:]

def insert_after_anchor_regex_free(haystack: str, anchor: str, insertion: str) -> str:
    # Simpler: look for an anchor string and insert right after it
    idx = haystack.find(anchor)
    if idx == -1:
        return haystack
    return haystack[:idx+len(anchor)] + insertion + haystack[idx+len(anchor):]

def replace_add_table_body(haystack: str) -> str:
    start = haystack.find("def _add_table(")
    if start == -1:  # nothing to patch
        return haystack
    # naive func end: next "\ndef " after start or EOF
    nxt = haystack.find("\ndef ", start + 1)
    end = len(haystack) if nxt == -1 else nxt
    func = haystack[start:end]
    # Replace whole function with our aligned version
    new_func = (
        "def _add_table(doc: Document, headers: List[str], rows: List[Tuple[Any, ...]], style: str = \"Light Shading Accent 1\"):\n"
        "    t = doc.add_table(rows=1, cols=len(headers))\n"
        "    t.style = style\n"
        "    hdr = t.rows[0].cells\n"
        "    for i, h in enumerate(headers):\n"
        "        hdr[i].text = str(h)\n"
        "        _cell_bold(hdr[i], 10)\n"
        "    def _is_num(s: str) -> bool:\n"
        "        s = (s or '').strip().replace(',', '').replace('_', '').upper()\n"
        "        if ' ' in s:\n"
        "            head, unit = s.split(' ', 1)\n"
        "            if unit in {'B','KB','MB','GB','TB','PB'}:\n"
        "                try:\n"
        "                    float(head); return True\n"
        "                except Exception:\n"
        "                    return False\n"
        "        try:\n"
        "            float(s.replace('%','')); return True\n"
        "        except Exception:\n"
        "            return False\n"
        "    for r in rows:\n"
        "        row = t.add_row().cells\n"
        "        for i, v in enumerate(r):\n"
        "            s = '' if v is None else str(v)\n"
        "            row[i].text = s\n"
        "            if _is_num(s) and row[i].paragraphs:\n"
        "                row[i].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT\n"
        "    return t\n"
    )
    return haystack[:start] + new_func + haystack[end:]

def add_page_numbers_call(haystack: str, fn_name: str) -> str:
    # Find function block by simple scanning
    start = haystack.find(f"def {fn_name}(")
    if start == -1:
        return haystack
    # Block end = next top-level "def " or EOF
    nxt = haystack.find("\ndef ", start + 1)
    end = len(haystack) if nxt == -1 else nxt
    block = haystack[start:end]

    # Find "doc = Document()" inside the block and insert the call on the next line
    k = block.find("doc = Document()")
    if k == -1:
        return haystack  # nothing to do
    # Determine indent by reading backwards to line start
    line_start = block.rfind("\n", 0, k) + 1
    line = block[line_start:block.find("\n", k)]
    indent = ""
    for ch in line:
        if ch in (" ", "\t"):
            indent += ch
        else:
            break

    insertion = f"\n{indent}_install_page_numbers(doc)"
    # insert after "doc = Document()"
    k2 = k + len("doc = Document()")
    block2 = block[:k2] + insertion + block[k2:]

    return haystack[:start] + block2 + haystack[end:]

def add_cover_spacing_after_title(haystack: str, fn_name: str) -> str:
    start = haystack.find(f"def {fn_name}(")
    if start == -1:
        return haystack
    nxt = haystack.find("\ndef ", start + 1)
    end = len(haystack) if nxt == -1 else nxt
    block = haystack[start:end]
    tpos = block.find("_add_title(")
    if tpos == -1:
        return haystack
    # find end of that line
    line_end = block.find("\n", tpos)
    if line_end == -1:
        return haystack
    # determine indent from next line
    following_line_start = line_end + 1
    following_line_end = block.find("\n", following_line_start)
    following_line = block[following_line_start:following_line_end if following_line_end!=-1 else len(block)]
    indent = ""
    for ch in following_line:
        if ch in (" ", "\t"):
            indent += ch
        else:
            break
    insertion = f"\n{indent}doc.add_paragraph()"
    block2 = block[:line_end] + insertion + block[line_end:]
    return haystack[:start] + block2 + haystack[end:]

def insert_chart_helpers(haystack: str) -> str:
    anchor = "def _add_metrics_section("
    idx = haystack.find(anchor)
    if idx == -1:
        return haystack
    insert_at = idx + len(anchor)
    snippet = (
        "\n\ndef _add_bar_chart(doc: Document, title: str, labels: List[str], series: List[Tuple[str, List[float]]], legend: bool = True):\n"
        "    if plt is None or not labels or not series:\n"
        "        return\n"
        "    try:\n"
        "        import matplotlib.pyplot as _plt\n"
        "        fig, ax = _plt.subplots(figsize=(6.5, 2.2))\n"
        "        x = list(range(len(labels)))\n"
        "        n = max(1, len(series))\n"
        "        width = 0.8 / n\n"
        "        offs = [(i - (n - 1) / 2) * width for i in range(n)]\n"
        "        for (name, vals), off in zip(series, offs):\n"
        "            ax.bar([xi + off for xi in x], vals, width, label=name)\n"
        "        ax.set_xticks(x)\n"
        "        ax.set_xticklabels(labels, rotation=0)\n"
        "        ax.set_title(title)\n"
        "        if legend and len(series) > 1:\n"
        "            ax.legend(loc='upper right', frameon=False)\n"
        "        fig.tight_layout()\n"
        "        import tempfile\n"
        "        with tempfile.NamedTemporaryFile(suffix='.png', delete=False) as tmp:\n"
        "            fig.savefig(tmp.name, dpi=120)\n"
        "        _plt.close(fig)\n"
        "        doc.add_picture(tmp.name)\n"
        "    except Exception:\n"
        "        pass\n"
        "\n"
        "def _bytes_str_to_mb(s: str) -> float:\n"
        "    try:\n"
        "        parts = (s or '').split()\n"
        "        if not parts:\n"
        "            return 0.0\n"
        "        v = float(parts[0])\n"
        "        unit = parts[1].upper() if len(parts) > 1 else 'B'\n"
        "        mul = {'B':1/1e6,'KB':1e-3,'MB':1.0,'GB':1e3,'TB':1e6,'PB':1e9}.get(unit, 0.0)\n"
        "        return v * mul\n"
        "    except Exception:\n"
        "        return 0.0\n"
    )
    return haystack[:insert_at] + snippet + haystack[insert_at:]

def insert_monthly_yearly_charts(haystack: str) -> str:
    # Insert charts right AFTER the existing monthly table block
    monthly_marker = ' _add_metrics_section(doc, f"Volume by Month'
    i = haystack.find(monthly_marker)
    if i != -1:
        # find end of that call (the closing parenthesis of the function call)
        close = haystack.find(")\n", i)
        if close != -1:
            insertion = (
                "\n            if monthly:\n"
                "                labels = [r[0] for r in monthly]\n"
                "                load = [_bytes_str_to_mb(r[1]) for r in monthly]\n"
                "                cdc  = [_bytes_str_to_mb(r[2]) for r in monthly]\n"
                "                _add_bar_chart(doc, 'Monthly Volume (MB)', labels, [('Load', load), ('CDC', cdc)], legend=True)\n"
            )
            haystack = haystack[:close+2] + insertion + haystack[close+2:]

    yearly_marker = '_add_metrics_section(doc, "Annual Volume (Last 5 Years)"'
    j = haystack.find(yearly_marker)
    if j != -1:
        close2 = haystack.find(")\n", j)
        if close2 != -1:
            insertion2 = (
                "\n            if yearly:\n"
                "                labels = [r[0] for r in yearly]\n"
                "                load = [_bytes_str_to_mb(r[1]) for r in yearly]\n"
                "                cdc  = [_bytes_str_to_mb(r[2]) for r in yearly]\n"
                "                _add_bar_chart(doc, 'Yearly Volume (MB)', labels, [('Load', load), ('CDC', cdc)], legend=True)\n"
            )
            haystack = haystack[:close2+2] + insertion2 + haystack[close2+2:]

    return haystack

# 1) optional matplotlib import after docx.oxml.ns import qn
text = insert_after_first(
    text,
    "from docx.oxml.ns import qn\n",
    "\n# --- optional charting (skip if unavailable) ---\ntry:\n    import matplotlib.pyplot as plt\nexcept Exception:\n    plt = None\n"
)

# 2) page number helper after "# docx helpers"
text = insert_after_first(
    text,
    "# docx helpers",
    "\n\ndef _install_page_numbers(doc: Document):\n"
    "    \"\"\"Add page numbers; suppress on cover page.\"\"\"\n"
    "    try:\n"
    "        sect = doc.sections[0]\n"
    "        sect.different_first_page_header_footer = True\n"
    "        footer = sect.footer\n"
    "        p = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()\n"
    "        fld = OxmlElement('w:fldSimple')\n"
    "        fld.set(qn('w:instr'), 'PAGE')\n"
    "        p._p.append(fld)\n"
    "        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT\n"
    "    except Exception:\n"
    "        pass\n"
)

# 3) replace _add_table for numeric alignment
text = replace_add_table_body(text)

# 4) chart helpers appended after _add_metrics_section
text = insert_chart_helpers(text)

# 5) call page numbers and add cover spacing in both generators
for fn in ("generate_customer_report_docx", "generate_summary_docx"):
    text = add_page_numbers_call(text, fn)
    text = add_cover_spacing_after_title(text, fn)

# 6) insert charts next to monthly/yearly sections
text = insert_monthly_yearly_charts(text)

DST.write_text(text, encoding="utf-8")
print("Patched OK ->", DST)
